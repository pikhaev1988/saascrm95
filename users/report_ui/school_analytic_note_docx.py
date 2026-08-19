"""
BI-оформление DOCX аналитической справки по итогам ГИА.

Только Presentation Layer поверх готового analytic_note_ui.
"""

from __future__ import annotations

from typing import Any


_C = {
    "ink": "102A43",
    "muted": "627D98",
    "bg": "F0F4F8",
    "blue": "2B6CB0",
    "blue_bg": "EBF8FF",
    "good": "2F855A",
    "good_bg": "F0FFF4",
    "mid": "B7791F",
    "mid_bg": "FFFBEB",
    "warn": "C05621",
    "warn_bg": "FFFAF0",
    "low": "C53030",
    "low_bg": "FFF5F5",
    "title": "1A365D",
}

KIND_LABEL_DOCX = {
    "fact": "Подтверждённый факт",
    "conclusion": "Аналитический вывод",
    "hypothesis": "Гипотеза",
    "decision": "Управленческое решение",
}


def render_analytic_note_docx(doc, ui: dict[str, Any]) -> None:
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Cm

    section = doc.sections[0]
    section.top_margin = Cm(1.4)
    section.bottom_margin = Cm(1.4)
    section.left_margin = Cm(1.5)
    section.right_margin = Cm(1.5)

    year = ui.get("year")
    year_bit = f" · {year}" if year else ""
    _banner(
        doc,
        f"АНАЛИТИЧЕСКАЯ СПРАВКА ПО ИТОГАМ ГИА  ·  {ui.get('exam_label')}{year_bit}",
        "Управленческий документ образовательной организации",
    )

    avg_label = ui.get("avg_label") or "Средний балл"
    kpi = ui.get("kpi") or {}
    tones = kpi.get("tones") or {}

    # 01 Общая аналитика
    _module(doc, "01", "Общая аналитика")
    items = [
        ("👥", str(kpi.get("participants") or 0), "Количество участников", tones.get("participants", "neutral")),
        ("📄", str(kpi.get("total_results") or 0), "Результатов по предметам", "neutral"),
        ("📚", str(kpi.get("subjects_count") or 0), "Количество предметов", tones.get("subjects", "neutral")),
        ("📈", str(kpi.get("avg_score") if kpi.get("avg_score") is not None else "—"), avg_label, tones.get("avg", "mid")),
        ("🎯", f"{kpi.get('quality_rate')}%", "Качество знаний", tones.get("quality", "mid")),
        ("✅", f"{kpi.get('pass_rate')}%", "Успеваемость", tones.get("pass", "mid")),
        ("🏅", str(kpi.get("high_count") or 0), "Высокобалльники", tones.get("high", "neutral")),
        ("⚠", str(kpi.get("risk_count") or 0), "Группа риска", "warn"),
    ]
    kpi_t = doc.add_table(rows=2, cols=4)
    for idx, (icon, value, label, tone) in enumerate(items[:8]):
        r, c = divmod(idx, 4)
        _kpi_cell(kpi_t.rows[r].cells[c], icon, value, label, tone)
    _nofit(kpi_t)
    doc.add_paragraph()

    # 02 Анализ по предметам
    subjects = ui.get("subjects") or []
    if subjects:
        _module(doc, "02", "Анализ по предметам")
        st = doc.add_table(rows=1, cols=7)
        headers = ["Предмет", "Участ.", avg_label, "Усп., %", "Мин", "Макс", "Статус"]
        for i, h in enumerate(headers):
            p = st.rows[0].cells[i].paragraphs[0]
            run = p.add_run(h)
            _font(run, 8, True, _C["muted"])
            _shade(st.rows[0].cells[i], _C["bg"])
        for row in subjects:
            cells = st.add_row().cells
            vals = [
                str(row.get("name") or ""),
                str(row.get("participants") or ""),
                str(row.get("avg") if row.get("avg") is not None else ""),
                str(row.get("pass_rate") if row.get("pass_rate") is not None else ""),
                _fmt_num(row.get("min_v")),
                _fmt_num(row.get("max_v")),
                f"{row.get('icon', '')} {row.get('status', '')}".strip(),
            ]
            for i, val in enumerate(vals):
                p = cells[i].paragraphs[0]
                run = p.add_run(val)
                _font(run, 9, False, _C["ink"])
            tone = row.get("tone") or "mid"
            _shade(cells[6], _tone_bg(tone))
        _nofit(st)
        doc.add_paragraph()

    # 03 Поклассная статистика
    classes = ui.get("classes") or []
    if classes:
        _module(doc, "03", "Поклассная статистика")
        ct = doc.add_table(rows=1, cols=4)
        for i, h in enumerate(["Класс / параллель", "Участ.", avg_label, "Успеваемость, %"]):
            p = ct.rows[0].cells[i].paragraphs[0]
            run = p.add_run(h)
            _font(run, 8, True, _C["muted"])
            _shade(ct.rows[0].cells[i], _C["bg"])
        for row in classes:
            cells = ct.add_row().cells
            vals = [
                str(row.get("name") or ""),
                str(row.get("participants") or ""),
                str(row.get("avg") if row.get("avg") is not None else ""),
                f"{row.get('pass_rate')}%",
            ]
            for i, val in enumerate(vals):
                p = cells[i].paragraphs[0]
                run = p.add_run(val)
                _font(run, 9, False, _C["ink"])
            _shade(cells[3], _tone_bg(row.get("tone") or "mid"))
        _nofit(ct)
        doc.add_paragraph()

    # 04 Проблемные зоны
    _module(doc, "04", "Проблемные зоны")
    zones = ui.get("weak_zones") or []
    if zones:
        for i in range(0, len(zones), 2):
            pair = zones[i : i + 2]
            t = doc.add_table(rows=1, cols=len(pair))
            for j, card in enumerate(pair):
                cell = t.rows[0].cells[j]
                tone = card.get("tone") or "warn"
                _shade(cell, _tone_bg(tone))
                _pad(cell)
                p = cell.paragraphs[0]
                r = p.add_run(str(card.get("name") or ""))
                _font(r, 11, True, _C["ink"])
                p2 = cell.add_paragraph()
                r2 = p2.add_run(
                    f"Риск: {card.get('risk')}  ·  {avg_label}: {card.get('avg')}  ·  усп. {card.get('pass_rate')}%"
                )
                _font(r2, 9, False, _tone_fg(tone))
                p3 = cell.add_paragraph()
                r3 = p3.add_run(str(card.get("text") or ""))
                _font(r3, 9, False, _C["ink"])
            _nofit(t)
            doc.add_paragraph()
    else:
        box = doc.add_table(rows=1, cols=1)
        cell = box.rows[0].cells[0]
        _shade(cell, _C["good_bg"])
        _pad(cell)
        p = cell.paragraphs[0]
        r = p.add_run("🟢  По заданным критериям проблемные зоны не выявлены.")
        _font(r, 11, True, _C["good"])
        _nofit(box)
        doc.add_paragraph()

    # 05 Анализ высокобалльников
    _module(doc, "05", "Анализ высокобалльников")
    hs = ui.get("high_scorers") or {}
    box = doc.add_table(rows=1, cols=1)
    cell = box.rows[0].cells[0]
    _shade(cell, _tone_bg(hs.get("tone") or "neutral"))
    _pad(cell)
    p = cell.paragraphs[0]
    r = p.add_run(str(hs.get("count") if hs.get("count") is not None else 0))
    _font(r, 28, True, _C["ink"])
    p2 = cell.add_paragraph()
    r2 = p2.add_run(f"Высокобалльных результатов  ·  критерий: {hs.get('threshold_hint') or '—'}")
    _font(r2, 9, True, _C["muted"])
    statements = hs.get("statements") or []
    if statements:
        for line in statements:
            p3 = cell.add_paragraph()
            r3 = p3.add_run(f"• {_stmt_line(line)}")
            _font(r3, 10, False, _C["ink"])
    else:
        p3 = cell.add_paragraph()
        r3 = p3.add_run(str(hs.get("insight") or ""))
        _font(r3, 10, False, _C["ink"])
    _nofit(box)
    doc.add_paragraph()

    # 06 Сравнительный анализ и динамика
    dynamics = ui.get("dynamics") or []
    _module(doc, "06", "Сравнительный анализ и динамика")
    if dynamics:
        hint = doc.add_paragraph()
        hr = hint.add_run("Динамика среднего результата и успеваемости по годам (таблица — приложение).")
        _font(hr, 9, False, _C["muted"])
        for row in dynamics:
            _hbar(
                doc,
                label=f"{row.get('year')}  ·  участников={row.get('participants')}  ·  результатов={row.get('results') or row.get('participants')}  ·  ср. {row.get('avg')}",
                meta=f"{row.get('pass_rate')}%",
                pct=float(row.get("pass_bar") or row.get("pass_rate") or 0),
                color=_C["blue"],
            )
        dt = doc.add_table(rows=1, cols=5)
        for i, h in enumerate(["Год", "Участников", "Результатов", avg_label, "Успеваемость, %"]):
            p = dt.rows[0].cells[i].paragraphs[0]
            run = p.add_run(h)
            _font(run, 8, True, _C["muted"])
            _shade(dt.rows[0].cells[i], _C["bg"])
        for row in dynamics:
            cells = dt.add_row().cells
            for i, val in enumerate(
                [
                    str(row.get("year") or ""),
                    str(row.get("participants") or ""),
                    str(row.get("results") or row.get("participants") or ""),
                    str(row.get("avg") if row.get("avg") is not None else ""),
                    f"{row.get('pass_rate')}%",
                ]
            ):
                p = cells[i].paragraphs[0]
                run = p.add_run(val)
                _font(run, 9, True, _C["ink"])
        _nofit(dt)
    else:
        p = doc.add_paragraph()
        r = p.add_run("Недостаточно данных для построения динамики.")
        _font(r, 10, False, _C["muted"])
    doc.add_paragraph()

    # 07 Выводы
    _module(doc, "07", "Выводы")
    concl = ui.get("conclusions") or {}
    blocks = [
        ("Сильные стороны", concl.get("strengths") or [], "high"),
        ("Что требует внимания", concl.get("attention") or [], "warn"),
        ("Общая оценка", concl.get("overall") or [], "mid"),
    ]
    for title, lines, tone in blocks:
        box = doc.add_table(rows=1, cols=1)
        cell = box.rows[0].cells[0]
        _shade(cell, _tone_bg(tone))
        _pad(cell)
        p = cell.paragraphs[0]
        r = p.add_run(title)
        _font(r, 11, True, _tone_fg(tone))
        if lines:
            for line in lines:
                p2 = cell.add_paragraph()
                r2 = p2.add_run(f"• {line}")
                _font(r2, 10, False, _C["ink"])
        else:
            p2 = cell.add_paragraph()
            r2 = p2.add_run("—")
            _font(r2, 9, False, _C["muted"])
        _nofit(box)
        doc.add_paragraph()

    # 08 Рекомендации
    groups = ui.get("reco_groups") or []
    if groups:
        _module(doc, "08", "Рекомендации")
        for g in groups:
            box = doc.add_table(rows=1, cols=1)
            cell = box.rows[0].cells[0]
            _shade(cell, _C["bg"])
            _pad(cell)
            p = cell.paragraphs[0]
            r = p.add_run(f"{g.get('icon', '')}  {g.get('title', '')}".strip())
            _font(r, 11, True, _C["title"])
            for item in g.get("items") or []:
                p2 = cell.add_paragraph()
                r2 = p2.add_run(f"• {item}")
                _font(r2, 10, False, _C["ink"])
            _nofit(box)
            doc.add_paragraph()

    expert = ui.get("expert") or {}
    if expert.get("factor_analysis"):
        _module(doc, "09", "Факторный анализ")
        ft = doc.add_table(rows=1, cols=3)
        for i, h in enumerate(["Фактор", "Влияние", "Интерпретация"]):
            rr = ft.rows[0].cells[i].paragraphs[0].add_run(h)
            _font(rr, 8, True, _C["muted"])
            _shade(ft.rows[0].cells[i], _C["bg"])
        for row in expert.get("factor_analysis") or []:
            c = ft.add_row().cells
            why = row.get("why")
            kind = row.get("kind")
            if kind:
                why = f"[{KIND_LABEL_DOCX.get(kind, kind)}] {why}"
            vals = [row.get("factor"), row.get("impact"), why]
            for i, val in enumerate(vals):
                rr = c[i].paragraphs[0].add_run(str(val or ""))
                _font(rr, 9, False, _C["ink"])
        _nofit(ft)
        doc.add_paragraph()

    if expert.get("management_risks"):
        _module(doc, "10", "Управленческие риски")
        rt = doc.add_table(rows=1, cols=4)
        for i, h in enumerate(["Риск", "Вероятность", "Влияние", "Приоритет"]):
            rr = rt.rows[0].cells[i].paragraphs[0].add_run(h)
            _font(rr, 8, True, _C["muted"])
            _shade(rt.rows[0].cells[i], _C["bg"])
        for row in expert.get("management_risks") or []:
            c = rt.add_row().cells
            vals = [row.get("type"), row.get("probability"), row.get("impact"), row.get("priority")]
            for i, val in enumerate(vals):
                rr = c[i].paragraphs[0].add_run(str(val or ""))
                _font(rr, 9, False, _C["ink"])
        _nofit(rt)
        doc.add_paragraph()

    if expert.get("forecast"):
        _module(doc, "11", "Прогноз")
        for line in expert.get("forecast") or []:
            p = doc.add_paragraph()
            rr = p.add_run(f"• {_stmt_line(line)}")
            _font(rr, 10, False, _C["ink"])

    if expert.get("management_decisions"):
        _module(doc, "12", "Управленческие решения")
        for block in expert.get("management_decisions") or []:
            b = doc.add_table(rows=1, cols=1)
            cell = b.rows[0].cells[0]
            _shade(cell, _C["bg"])
            _pad(cell)
            rr = cell.paragraphs[0].add_run(str(block.get("role") or ""))
            _font(rr, 11, True, _C["title"])
            for action in block.get("actions") or []:
                p = cell.add_paragraph()
                r2 = p.add_run(f"• {action}")
                _font(r2, 10, False, _C["ink"])
            _nofit(b)
            doc.add_paragraph()

    if expert.get("plan_rows"):
        _module(doc, "13", "План мероприятий")
        pt = doc.add_table(rows=1, cols=6)
        headers = [
            "Проблема (факт риска)",
            "Гипотеза (требует проверки)",
            "Мероприятие (решение)",
            "Ответственный",
            "Ожидаемый результат",
            "Приоритет",
        ]
        for i, h in enumerate(headers):
            rr = pt.rows[0].cells[i].paragraphs[0].add_run(h)
            _font(rr, 8, True, _C["muted"])
            _shade(pt.rows[0].cells[i], _C["bg"])
        for row in expert.get("plan_rows") or []:
            c = pt.add_row().cells
            vals = [
                row.get("problem"),
                row.get("cause"),
                row.get("action"),
                row.get("owner"),
                row.get("effect"),
                row.get("priority"),
            ]
            for i, val in enumerate(vals):
                rr = c[i].paragraphs[0].add_run(str(val or ""))
                _font(rr, 9, False, _C["ink"])
        _nofit(pt)
        doc.add_paragraph()

    if expert.get("final_expert_conclusion"):
        _module(doc, "14", "Итоговое экспертное заключение")
        for line in expert.get("final_expert_conclusion") or []:
            p = doc.add_paragraph()
            rr = p.add_run(f"• {_stmt_line(line)}")
            _font(rr, 10, False, _C["ink"])

    m = ui.get("methodological") or {}
    _module(doc, "15–21", "Методический анализ ФИПИ: темы и приоритеты")
    p = doc.add_paragraph()
    rr = p.add_run(str(m.get("methodology_note") or ""))
    _font(rr, 9, False, _C["muted"])
    p = doc.add_paragraph()
    rr = p.add_run(str(m.get("appendix_note") or ""))
    _font(rr, 9, False, _C["muted"])
    if not m.get("has_any_task_data"):
        p = doc.add_paragraph()
        rr = p.add_run(str(m.get("empty_message") or "Недостаточно данных по заданиям КИМ."))
        _font(rr, 10, False, _C["ink"])
    else:
        for line in m.get("aggregation_overview") or []:
            p = doc.add_paragraph()
            rr = p.add_run(f"• {line}")
            _font(rr, 10, False, _C["ink"])
        for block in m.get("subject_blocks") or []:
            if not block.get("has_task_data"):
                continue
            p = doc.add_paragraph()
            rr = p.add_run(str(block.get("subject") or ""))
            _font(rr, 11, True, _C["title"])
            if block.get("analysis_text"):
                p2 = doc.add_paragraph()
                r2 = p2.add_run(str(block.get("analysis_text")))
                _font(r2, 10, False, _C["ink"])
            for t in (block.get("topics") or [])[:3]:
                p3 = doc.add_paragraph()
                r3 = p3.add_run(
                    f"{t.get('topic')}: задания {t.get('tasks_label')}, "
                    f"{t.get('success_rate')}%, {t.get('risk_label')}, {t.get('grade_label')}."
                )
                _font(r3, 9, False, _C["ink"])

        p = doc.add_paragraph()
        rr = p.add_run("Приоритетные методические дефициты (ТОП-10)")
        _font(rr, 11, True, _C["title"])
        for idx, row in enumerate(m.get("priority_topics") or [], start=1):
            p = doc.add_paragraph()
            rr = p.add_run(
                f"{idx}. {row.get('topic')} · {row.get('subject')} · {row.get('grade_label')} · "
                f"{row.get('success_rate')}% · заданий {row.get('task_count')} ({row.get('tasks_label')}) · "
                f"{row.get('risk_label')}. {row.get('why_critical')} {row.get('exam_impact')}"
            )
            _font(rr, 9, False, _C["ink"])

        p = doc.add_paragraph()
        rr = p.add_run("Методическое заключение")
        _font(rr, 11, True, _C["title"])
        for line in m.get("methodological_conclusion") or []:
            p = doc.add_paragraph()
            rr = p.add_run(f"• {line}")
            _font(rr, 10, False, _C["ink"])

        if m.get("correction_plan"):
            p = doc.add_paragraph()
            rr = p.add_run("План коррекционной работы (кратко)")
            _font(rr, 11, True, _C["title"])
            for row in (m.get("correction_plan") or [])[:8]:
                p = doc.add_paragraph()
                rr = p.add_run(
                    f"{row.get('subject')} · {row.get('topic')} · {row.get('grade')} — {row.get('reason')}"
                )
                _font(rr, 9, False, _C["ink"])

    if expert.get("quality_self_check"):
        _module(doc, "QC", "Самопроверка качества анализа")
        ready = expert.get("ready")
        if ready is not None:
            p0 = doc.add_paragraph()
            r0 = p0.add_run(
                "✔ Справка соответствует критериям экспертного анализа."
                if ready
                else "✖ Требуется переработка: не все критерии качества выполнены."
            )
            _font(r0, 10, True, _C["good"] if ready else _C["low"])
        for check in expert.get("quality_self_check") or []:
            mark = "✔" if check.get("ok") else "✖"
            p = doc.add_paragraph()
            rr = p.add_run(f"{mark} {check.get('label')}")
            _font(rr, 10, False, _C["ink"])


def _stmt_line(item) -> str:
    if isinstance(item, dict):
        label = item.get("label") or item.get("kind") or ""
        text = item.get("text") or ""
        return f"[{label}] {text}".strip() if label else str(text)
    return str(item or "")


def _banner(doc, title: str, subtitle: str) -> None:
    t = doc.add_table(rows=1, cols=1)
    cell = t.rows[0].cells[0]
    _shade(cell, _C["blue_bg"])
    _pad(cell, 140)
    p = cell.paragraphs[0]
    r = p.add_run(title)
    _font(r, 14, True, _C["title"])
    p2 = cell.add_paragraph()
    r2 = p2.add_run(subtitle)
    _font(r2, 9, False, _C["muted"])
    _nofit(t)
    doc.add_paragraph()


def _module(doc, num: str, title: str) -> None:
    p = doc.add_paragraph()
    r = p.add_run(f"{num}  ·  {title}")
    _font(r, 13, True, _C["title"])
    p.paragraph_format.space_before = 200
    p.paragraph_format.space_after = 120


def _kpi_cell(cell, icon: str, value: str, label: str, tone: str) -> None:
    _shade(cell, _tone_bg(tone))
    _pad(cell)
    p = cell.paragraphs[0]
    r = p.add_run(f"{icon}  {label}")
    _font(r, 8, True, _C["muted"])
    p2 = cell.add_paragraph()
    r2 = p2.add_run(value)
    _font(r2, 16, True, _C["ink"])


def _hbar(doc, *, label: str, meta: str, pct: float, color: str) -> None:
    t = doc.add_table(rows=1, cols=1)
    cell = t.rows[0].cells[0]
    _pad(cell, 80)
    p = cell.paragraphs[0]
    r = p.add_run(label)
    _font(r, 9, True, _C["ink"])
    p2 = cell.add_paragraph()
    bar = _bar(pct)
    r2 = p2.add_run(f"{bar}  {meta}")
    _font(r2, 9, False, color)
    _nofit(t)


def _bar(pct: float) -> str:
    filled = max(0, min(10, int(round(float(pct or 0) / 10.0))))
    return "█" * filled + "░" * (10 - filled)


def _fmt_num(v) -> str:
    if v is None:
        return "—"
    try:
        return f"{float(v):.2f}"
    except (TypeError, ValueError):
        return str(v)


def _tone_bg(tone: str) -> str:
    return {
        "high": _C["good_bg"],
        "mid": _C["mid_bg"],
        "warn": _C["warn_bg"],
        "low": _C["low_bg"],
        "neutral": _C["blue_bg"],
    }.get(tone, _C["bg"])


def _tone_fg(tone: str) -> str:
    return {
        "high": _C["good"],
        "mid": _C["mid"],
        "warn": _C["warn"],
        "low": _C["low"],
        "neutral": _C["blue"],
    }.get(tone, _C["ink"])


def _shade(cell, hex_color: str) -> None:
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), hex_color)
    shd.set(qn("w:val"), "clear")
    tcPr.append(shd)


def _pad(cell, twips: int = 100) -> None:
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcMar = OxmlElement("w:tcMar")
    for edge in ("top", "left", "bottom", "right"):
        node = OxmlElement(f"w:{edge}")
        node.set(qn("w:w"), str(twips))
        node.set(qn("w:type"), "dxa")
        tcMar.append(node)
    tcPr.append(tcMar)


def _font(run, size_pt: int, bold: bool, color_hex: str) -> None:
    from docx.shared import Pt, RGBColor

    run.bold = bold
    run.font.size = Pt(size_pt)
    run.font.color.rgb = RGBColor.from_string(color_hex)
    run.font.name = "Calibri"


def _nofit(table) -> None:
    table.autofit = True
