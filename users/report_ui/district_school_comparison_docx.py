"""
Финальное оформление DOCX
«Сравнительный анализ образовательных организаций (муниципалитет)».

Только визуальный слой: тексты, показатели и категории — из payload без изменений.
"""

from __future__ import annotations

from io import BytesIO
from typing import Any

NAVY = "0B1F3A"
BLUE = "1B4F72"
SKY = "5B9BD5"
TEAL = "0E7C7B"
GRAY = "5D6D7E"
LINE = "D5DDE5"
SOFT = "F4F7FA"
WHITE = "FFFFFF"
GOOD = "1E7A4A"
GOOD_BG = "EAF6EF"
LEAD = "1B4F72"
LEAD_BG = "E8F0F8"
STABLE = "A67C00"
STABLE_BG = "FBF5E4"
WATCH = "C27A1E"
WATCH_BG = "F8F0E4"
WARN = "A85A1A"
WARN_BG = "F8F1E6"
RISK = "9B2C2C"
RISK_BG = "F8ECEC"
CRIT = "7A1F1F"
CRIT_BG = "F5E4E4"
INFO_BG = "EAF2F8"
POS_BG = "EAF6EF"
NEG_BG = "F8ECEC"
DOC_VERSION = "1.0"

# Порядок категорий для визуальных блоков (метки из бизнес-логики без изменений)
CATEGORY_VISUAL: list[tuple[str, str, str, str]] = [
    ("устойчивый лидер", "Устойчивый лидер", GOOD, GOOD_BG),
    ("лидер", "Лидер", LEAD, LEAD_BG),
    ("стабильная организация", "Стабильная организация", STABLE, STABLE_BG),
    ("зона наблюдения", "Зона наблюдения", WATCH, WATCH_BG),
    ("зона риска", "Зона риска", WARN, WARN_BG),
    ("критическая зона", "Критическая зона", CRIT, CRIT_BG),
]


def render_district_school_comparison_docx(data: dict[str, Any]) -> BytesIO:
    from docx import Document

    doc = Document()
    _setup_page(doc)

    if not (data or {}).get("has_data"):
        doc.add_paragraph((data or {}).get("message") or "Недостаточно данных.")
        return _save(doc)

    district = str(data.get("district_name") or "Муниципалитет")
    et = (data.get("exam_type") or "ege").lower()
    exam = "ОГЭ" if et == "oge" else "ЕГЭ"
    year = data.get("year") or ""
    formed = str(data.get("generated_at") or "")
    avg_label = "Средняя оценка" if et == "oge" else "Средний балл"

    _setup_header_footer(doc, district=district, exam=exam, year=year, formed=formed)
    _title_page(doc, district=district, exam=exam, year=year, formed=formed)
    _toc(doc)

    # —— Executive Summary ——
    _section_block(doc, "Executive Summary · управленческий срез", "◆", accent=SKY)
    leaders = list(data.get("leaders") or [])
    at_risk = list(data.get("at_risk") or [])
    _kpi_dashboard(
        doc,
        [
            ("⌂", "Школ в срезе", str(data.get("schools_count") or len(data.get("school_rows") or [])), SKY),
            ("👤", "Участники", str(data.get("participants")), SKY),
            ("Σ", avg_label, str(data.get("avg_score")), _tone_avg(data.get("avg_score"), et)),
            ("Q", "Качество", f"{data.get('quality_rate')}%", _tone_pct(data.get("quality_rate"), 40, 25)),
            ("%", "Успеваемость", f"{data.get('pass_rate')}%", _tone_pct(data.get("pass_rate"), 85, 70)),
            ("★", "Лидеры", str(len(leaders)), GOOD),
            ("!", "Школы риска", str(len(at_risk)), RISK if at_risk else GOOD),
            ("↗", "Динамика", _fmt_delta(data.get("avg_delta")), _tone_delta(data.get("avg_delta"))),
        ],
    )

    exec_items = list(data.get("executive_summary") or data.get("summary") or [])
    analytical = _section_items(data, ("итоговое аналитическое резюме", "аналитическое резюме", "резюме"))
    if not analytical:
        analytical = exec_items
    _summary_insight_cards(doc, analytical[:6] if analytical else exec_items[:6])
    _assessment_card(
        doc,
        "Итоговая управленческая оценка",
        (exec_items[0] if exec_items else (analytical[0] if analytical else "—")),
    )
    doc.add_page_break()

    # —— Секции AI / fallback ——
    sections = list(data.get("report_sections") or [])
    if sections:
        for section in sections:
            title = str(section.get("title") or "Раздел")
            items = list(section.get("items") or [])
            kind = _detect_kind(title)
            _section_block(doc, title, _kind_icon(kind), accent=_kind_accent(kind))
            _render_section_by_kind(doc, data, title, items, kind=kind, exam=exam, et=et, avg_label=avg_label)
    else:
        _section_block(doc, "1. Итоговое аналитическое резюме", "◆", accent=SKY)
        _summary_insight_cards(doc, exec_items)
        if data.get("recommendations"):
            _section_block(doc, "Управленческие решения", "▸", accent=BLUE)
            for i, item in enumerate(data.get("recommendations") or [], start=1):
                _plan_card(doc, i, str(item))

    # —— Рейтинговая таблица (данные school_rows без изменения значений) ——
    _section_block(doc, "Рейтинг образовательных организаций · сводная таблица", "◆", accent=BLUE)
    ranking_rows, cat_tones, risk_tones, bar_vals = _build_ranking_rows(data)
    _ranking_table(doc, ranking_rows, cat_tones=cat_tones, risk_tones=risk_tones, bar_vals=bar_vals, et=et)

    return _save(doc)


# ─── section routing ───────────────────────────────────────────────


def _section_items(data: dict, title_parts: tuple[str, ...]) -> list:
    for section in data.get("report_sections") or []:
        t = str(section.get("title") or "").lower()
        if any(p in t for p in title_parts):
            return list(section.get("items") or [])
    return []


def _detect_kind(title: str) -> str:
    t = (title or "").lower()
    if "резюме" in t or "аналитическ" in t and "итог" in t:
        return "summary"
    if "общие статистическ" in t:
        return "stats"
    if "сравнительн" in t:
        return "compare"
    if "рейтинг" in t:
        return "ranking_text"
    if "риск" in t and "школ" in t:
        return "risk_schools"
    if "лидер" in t:
        return "leaders"
    if "предметн" in t:
        return "subjects"
    if "выявленн" in t and "риск" in t:
        return "risks"
    if "прогноз" in t:
        return "forecast"
    if "управленческ" in t or "решен" in t:
        return "management"
    if "первоочередн" in t or "мер" in t and "приоритет" in t:
        return "plan"
    if "заключ" in t:
        return "conclusion"
    return "info"


def _kind_icon(kind: str) -> str:
    return {
        "summary": "◆",
        "stats": "Σ",
        "compare": "↔",
        "ranking_text": "☰",
        "risk_schools": "⚠",
        "leaders": "★",
        "subjects": "◎",
        "risks": "!",
        "forecast": "↗",
        "management": "▸",
        "plan": "▸",
        "conclusion": "◆",
        "info": "◆",
    }.get(kind, "◆")


def _kind_accent(kind: str) -> str:
    return {
        "summary": SKY,
        "stats": BLUE,
        "compare": TEAL,
        "ranking_text": BLUE,
        "risk_schools": RISK,
        "leaders": GOOD,
        "subjects": WARN,
        "risks": RISK,
        "forecast": SKY,
        "management": BLUE,
        "plan": BLUE,
        "conclusion": NAVY,
        "info": SKY,
    }.get(kind, SKY)


def _render_section_by_kind(
    doc,
    data: dict,
    title: str,
    items: list,
    *,
    kind: str,
    exam: str,
    et: str,
    avg_label: str,
) -> None:
    if kind == "summary":
        _summary_insight_cards(doc, items)
        return
    if kind == "stats":
        for item in items:
            _info_card(doc, str(item))
        return
    if kind == "compare":
        _category_blocks(doc, data)
        for item in items:
            _info_card(doc, str(item))
        return
    if kind == "ranking_text":
        for item in items:
            _info_card(doc, str(item))
        return
    if kind == "risk_schools":
        _school_risk_cards(doc, data.get("at_risk") or [], et=et, avg_label=avg_label)
        for i, item in enumerate(items or [], start=1):
            _risk_card(doc, i, str(item), badge="Анализ")
        return
    if kind == "leaders":
        _school_leader_cards(doc, data.get("leaders") or [], et=et, avg_label=avg_label)
        for i, item in enumerate(items or [], start=1):
            _strength_card(doc, i, str(item))
        return
    if kind == "subjects":
        _subject_risk_cards(doc, data)
        for item in items:
            _info_card(doc, str(item))
        return
    if kind == "risks":
        for i, item in enumerate(items or [], start=1):
            _risk_detail_card(doc, i, str(item))
        return
    if kind == "forecast":
        _forecast_scenarios(doc, items)
        return
    if kind == "management":
        _management_audience_cards(doc, items)
        return
    if kind == "plan":
        for i, item in enumerate(items or [], start=1):
            _plan_card(doc, i, str(item))
        return
    if kind == "conclusion":
        _conclusion_page(doc, data, items)
        return
    for item in items or []:
        _info_card(doc, str(item))


# ─── visual blocks ─────────────────────────────────────────────────


def _summary_insight_cards(doc, items: list) -> None:
    """Карточки резюме с цветовой индикацией по ключевым словам (только оформление)."""
    if not items:
        return
    for i, raw in enumerate(items, start=1):
        text = str(raw or "").strip()
        if not text:
            continue
        label, accent, bg = _classify_insight(text)
        _labeled_card(doc, num=i, badge=label, text=text, accent=accent, bg=bg)


def _classify_insight(text: str) -> tuple[str, str, str]:
    t = text.lower()
    if any(k in t for k in ("достижен", "сильн", "успеш", "лидер", "положительн", "рост")):
        return ("Основные достижения", GOOD, GOOD_BG)
    if any(k in t for k in ("проблем", "дефицит", "негатив", "снижен", "ухудш", "риск", "критич")):
        return ("Основные проблемы", RISK, RISK_BG)
    if any(k in t for k in ("динамик", "рост", "улучш", "повыш")) and "отриц" not in t and "негатив" not in t:
        return ("Положительная динамика", GOOD, POS_BG)
    if any(k in t for k in ("отрицательн", "негативн", "снижен", "паден")):
        return ("Негативные тенденции", WARN, WARN_BG)
    if any(k in t for k in ("вывод", "итог", "оценк", "заключ")):
        return ("Главный вывод", BLUE, INFO_BG)
    return ("Ключевой ориентир", SKY, INFO_BG)


def _assessment_card(doc, title: str, text: str) -> None:
    from docx.shared import Cm, Pt, RGBColor

    t = doc.add_table(rows=1, cols=2)
    t.autofit = False
    a, b = t.rows[0].cells
    _set_cell_width(a, Cm(0.22))
    _set_cell_width(b, Cm(16.5))
    _shade_cell(a, NAVY)
    _shade_cell(b, SOFT)
    p = b.paragraphs[0]
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(2)
    br = p.add_run(f"🎯  {title}")
    br.bold = True
    br.font.size = Pt(10)
    br.font.color.rgb = RGBColor(0x0B, 0x1F, 0x3A)
    p2 = b.add_paragraph()
    p2.paragraph_format.space_before = Pt(0)
    p2.paragraph_format.space_after = Pt(6)
    run = p2.add_run(str(text or "—"))
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x1A, 0x23, 0x32)


def _category_blocks(doc, data: dict) -> None:
    from docx.shared import Cm, Pt, RGBColor

    rows = list(data.get("school_rows") or [])
    by_cat: dict[str, list] = {}
    for row in rows:
        cat = str(row.get("category") or row.get("tier") or "стабильная организация").strip().lower()
        by_cat.setdefault(cat, []).append(row)

    for key, label, accent, bg in CATEGORY_VISUAL:
        schools = by_cat.get(key) or []
        if not schools:
            continue
        wrap = doc.add_table(rows=1, cols=2)
        wrap.autofit = False
        a, b = wrap.rows[0].cells
        _set_cell_width(a, Cm(0.22))
        _set_cell_width(b, Cm(16.5))
        _shade_cell(a, accent)
        _shade_cell(b, bg)
        p = b.paragraphs[0]
        p.paragraph_format.space_before = Pt(5)
        p.paragraph_format.space_after = Pt(2)
        badge = p.add_run(f"●  {label}  ·  {len(schools)} ОО")
        badge.bold = True
        badge.font.size = Pt(10)
        badge.font.color.rgb = RGBColor.from_string(accent)
        names = ", ".join(
            str(s.get("student__school__name") or "—")[:70] for s in schools[:12]
        )
        if len(schools) > 12:
            names += f" … (+{len(schools) - 12})"
        p2 = b.add_paragraph()
        p2.paragraph_format.space_before = Pt(0)
        p2.paragraph_format.space_after = Pt(5)
        r = p2.add_run(names)
        r.font.size = Pt(9)
        r.font.color.rgb = RGBColor(0x1A, 0x23, 0x32)


def _school_risk_cards(doc, schools: list, *, et: str, avg_label: str) -> None:
    if not schools:
        _info_card(doc, "Школ в зоне риска и критической зоне в текущем срезе не выявлено.")
        return
    for i, row in enumerate(schools, start=1):
        name = str(row.get("student__school__name") or "—")
        cat = str(row.get("category") or "зона риска")
        risk = str(row.get("risk") or "—")
        avg = f"{float(row.get('avg') or 0):.2f}"
        pr = str(row.get("pass_rate") if row.get("pass_rate") is not None else "—")
        qr = str(row.get("quality_rate") if row.get("quality_rate") is not None else "—")
        delta = row.get("delta_vs_district")
        delta_s = f"{delta:+.2f}" if delta is not None else "—"
        priority = "Высокий" if "критич" in cat.lower() else "Средний"
        body = (
            f"{avg_label}: {avg}  ·  Усп.: {pr}%  ·  Кач.: {qr}%  ·  Δ к МО: {delta_s}\n"
            f"Категория: {cat}  ·  Риск: {risk}  ·  Приоритет вмешательства: {priority}"
        )
        _labeled_card(doc, num=i, badge="Школа риска", text=f"{name}\n{body}", accent=RISK, bg=RISK_BG)


def _school_leader_cards(doc, schools: list, *, et: str, avg_label: str) -> None:
    if not schools:
        _info_card(doc, "Школ-лидеров в текущем срезе недостаточно для выделения группы.")
        return
    for i, row in enumerate(schools, start=1):
        name = str(row.get("student__school__name") or "—")
        cat = str(row.get("category") or "лидер")
        avg = f"{float(row.get('avg') or 0):.2f}"
        pr = str(row.get("pass_rate") if row.get("pass_rate") is not None else "—")
        qr = str(row.get("quality_rate") if row.get("quality_rate") is not None else "—")
        delta = row.get("delta_vs_district")
        delta_s = f"{delta:+.2f}" if delta is not None else "—"
        body = (
            f"Место в группе лидеров: {i}\n"
            f"{avg_label}: {avg}  ·  Усп.: {pr}%  ·  Кач.: {qr}%  ·  Δ к МО: {delta_s}\n"
            f"Категория: {cat}  ·  Потенциал: тиражирование практик подготовки"
        )
        _labeled_card(doc, num=i, badge="Школа-лидер", text=f"{name}\n{body}", accent=GOOD, bg=GOOD_BG)


def _subject_risk_cards(doc, data: dict) -> None:
    weak = list(data.get("weak_subjects") or [])
    rows = weak or [
        r
        for r in (data.get("subject_rows") or [])
        if str(r.get("risk") or "").lower() in {"высокий", "средний", "критический"}
        or float(r.get("pass_rate") or 100) < 75
    ]
    if not rows:
        for r in (data.get("subject_rows") or [])[:6]:
            rows.append(r)
    for i, row in enumerate(rows[:10], start=1):
        name = str(row.get("exam__subject") or "—")
        risk = str(row.get("risk") or "—")
        pr = row.get("pass_rate")
        avg = f"{float(row.get('avg') or 0):.2f}"
        priority = "Высокий" if float(pr or 100) < 60 else ("Средний" if float(pr or 100) < 75 else "Низкий")
        text = f"{name}\nСредний: {avg}  ·  Усп.: {pr}%  ·  Риск: {risk}  ·  Приоритет: {priority}"
        accent = RISK if priority == "Высокий" else (WARN if priority == "Средний" else GOOD)
        bg = RISK_BG if accent == RISK else (WARN_BG if accent == WARN else GOOD_BG)
        _labeled_card(doc, num=i, badge="Предмет", text=text, accent=accent, bg=bg)


def _risk_detail_card(doc, num: int, text: str) -> None:
    t = (text or "").lower()
    level = "Высокий" if any(k in t for k in ("критич", "срочн", "высок")) else ("Средний" if "риск" in t else "Умеренный")
    influence = "Существенное" if level == "Высокий" else "Умеренное"
    probability = "Высокая" if level == "Высокий" else "Средняя"
    body = (
        f"{text}\n"
        f"Уровень риска: {level}  ·  Вероятность: {probability}  ·  Влияние: {influence}"
    )
    _risk_card(doc, num, body, badge="Риск")


def _forecast_scenarios(doc, items: list) -> None:
    """Три сценария — визуальная рамка; тексты из payload без переписывания."""
    scenarios = [
        ("Оптимистичный сценарий", GOOD, GOOD_BG),
        ("Базовый сценарий", BLUE, INFO_BG),
        ("Негативный сценарий", RISK, RISK_BG),
    ]
    cleaned = [str(x).strip() for x in (items or []) if str(x).strip()]
    if not cleaned:
        return
    # Распределяем пункты по сценариям по ключевым словам, иначе по порядку
    buckets: list[list[str]] = [[], [], []]
    for item in cleaned:
        low = item.lower()
        if any(k in low for k in ("оптимист", "улучш", "рост", "повыш")):
            buckets[0].append(item)
        elif any(k in low for k in ("негатив", "ухудш", "снижен", "риск ухудш")):
            buckets[2].append(item)
        elif any(k in low for k in ("базов", "сохранени", "стабил")):
            buckets[1].append(item)
        else:
            # round-robin
            idx = min(range(3), key=lambda i: len(buckets[i]))
            buckets[idx].append(item)
    for i, (title, accent, bg) in enumerate(scenarios):
        texts = buckets[i] or (cleaned[i : i + 1] if i < len(cleaned) else [])
        if not texts:
            continue
        body = "\n".join(texts)
        _labeled_card(doc, num=i + 1, badge=title, text=body, accent=accent, bg=bg)


def _management_audience_cards(doc, items: list) -> None:
    """Группировка по аудитории — только если в тексте уже есть заголовки блоков."""
    audiences = [
        ("муниципалитет", "Для муниципалитета", BLUE, INFO_BG),
        ("муниципального органа", "Для муниципалитета", BLUE, INFO_BG),
        ("руководител", "Для руководителей школ", TEAL, "E6F4F3"),
        ("заместител", "Для заместителей директоров", SKY, INFO_BG),
        ("методическ", "Для руководителей методических объединений", STABLE, STABLE_BG),
        ("педагогическ", "Для педагогов", GOOD, GOOD_BG),
        ("педагогических работников", "Для педагогов", GOOD, GOOD_BG),
    ]
    current_label = "Управленческие решения"
    current_accent, current_bg = BLUE, INFO_BG
    bucket: list[str] = []
    num = 0

    def flush() -> None:
        nonlocal num, bucket
        if not bucket:
            return
        num += 1
        _labeled_card(
            doc,
            num=num,
            badge=current_label,
            text="\n".join(bucket),
            accent=current_accent,
            bg=current_bg,
        )
        bucket = []

    for raw in items or []:
        text = str(raw or "").strip()
        if not text:
            continue
        low = text.lower().rstrip(":")
        matched = False
        for key, label, accent, bg in audiences:
            if key in low and len(text) < 120 and (
                text.endswith(":") or text.lower().startswith("рекомендац") or "для " in low
            ):
                flush()
                current_label, current_accent, current_bg = label, accent, bg
                matched = True
                break
        if matched:
            continue
        bucket.append(text)
    flush()


def _conclusion_page(doc, data: dict, items: list) -> None:
    leaders_n = len(data.get("leaders") or [])
    risk_n = len(data.get("at_risk") or [])
    _assessment_card(
        doc,
        "Итоговая оценка муниципалитета",
        (
            f"Средний показатель: {data.get('avg_score')}; успеваемость: {data.get('pass_rate')}%; "
            f"качество: {data.get('quality_rate')}%; лидеров: {leaders_n}; школ риска: {risk_n}."
        ),
    )
    for i, item in enumerate(items or [], start=1):
        label, accent, bg = _classify_insight(str(item))
        if i == len(items):
            label = "Стратегические задачи следующего года"
            accent, bg = BLUE, INFO_BG
        elif "направлен" in str(item).lower() or "развит" in str(item).lower():
            label = "Ключевые направления развития"
            accent, bg = TEAL, "E6F4F3"
        _labeled_card(doc, num=i, badge=label, text=str(item), accent=accent, bg=bg)


# ─── ranking table ─────────────────────────────────────────────────


def _build_ranking_rows(data: dict) -> tuple[list[list[str]], list[str], list[str], list[float]]:
    rows = sorted(
        list(data.get("school_rows") or []),
        key=lambda x: (-float(x.get("avg") or 0), str(x.get("student__school__name") or "")),
    )
    out: list[list[str]] = []
    cat_tones: list[str] = []
    risk_tones: list[str] = []
    bar_vals: list[float] = []
    for i, row in enumerate(rows, start=1):
        cat = str(row.get("category") or row.get("tier") or "—")
        risk = str(row.get("risk") or "—")
        pr = float(row.get("pass_rate") or 0)
        out.append(
            [
                str(i),
                str(row.get("student__school__name") or "—"),
                f"{float(row.get('avg') or 0):.2f}",
                str(row.get("pass_rate") if row.get("pass_rate") is not None else ""),
                _badge_label(cat),
                risk,
                _level_bar(pr),
            ]
        )
        cat_tones.append(_category_accent(cat))
        risk_tones.append(RISK if "высок" in risk.lower() or "крит" in risk.lower() else (WARN if "сред" in risk.lower() else GOOD))
        bar_vals.append(pr)
    return out, cat_tones, risk_tones, bar_vals


def _badge_label(cat: str) -> str:
    c = (cat or "").lower()
    mapping = {
        "устойчивый лидер": "● Устойчивый лидер",
        "лидер": "● Лидер",
        "стабильная организация": "● Стабильная",
        "зона наблюдения": "● Наблюдение",
        "зона риска": "● Зона риска",
        "критическая зона": "● Критическая зона",
    }
    return mapping.get(c, cat or "—")


def _category_accent(cat: str) -> str:
    c = (cat or "").lower()
    for key, _label, accent, _bg in CATEGORY_VISUAL:
        if key == c:
            return accent
    return SKY


def _level_bar(pass_rate: float) -> str:
    # Визуальный индикатор уровня (символьный, без изменения данных)
    filled = max(0, min(5, int(round(float(pass_rate or 0) / 20))))
    return "■" * filled + "□" * (5 - filled)


def _ranking_table(
    doc,
    rows: list[list[str]],
    *,
    cat_tones: list[str],
    risk_tones: list[str],
    bar_vals: list[float],
    et: str,
) -> None:
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Cm, Pt, RGBColor

    headers = ["№", "Образовательная организация", "Средний", "Усп., %", "Категория", "Риск", "Уровень"]
    widths = [0.8, 7.2, 1.6, 1.5, 3.0, 1.6, 1.5]
    table = doc.add_table(rows=1, cols=len(headers))
    table.autofit = False
    table.allow_autofit = False

    hdr = table.rows[0]
    _set_row_header(hdr)
    for i, h in enumerate(headers):
        cell = hdr.cells[i]
        _set_cell_width(cell, Cm(widths[i]))
        _shade_cell(cell, NAVY)
        p = cell.paragraphs[0]
        p.paragraph_format.space_before = Pt(3)
        p.paragraph_format.space_after = Pt(3)
        if i in {0, 2, 3, 6}:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(h)
        run.bold = True
        run.font.size = Pt(8)
        run.font.color.rgb = RGBColor(0xDC, 0xE8, 0xF3)

    for r_i, row in enumerate(rows):
        cells = table.add_row().cells
        _cant_split_row(table.rows[r_i + 1])
        for i, val in enumerate(row):
            cell = cells[i]
            _set_cell_width(cell, Cm(widths[i]))
            if r_i % 2 == 1:
                _shade_cell(cell, "F7FAFC")
            if i == 4 and r_i < len(cat_tones):
                # лёгкий оттенок категории
                for _k, _l, _a, bg in CATEGORY_VISUAL:
                    if _a == cat_tones[r_i]:
                        _shade_cell(cell, bg)
                        break
            if i == 3 and r_i < len(bar_vals):
                pr = bar_vals[r_i]
                if pr >= 85:
                    _shade_cell(cell, GOOD_BG)
                elif pr < 70:
                    _shade_cell(cell, RISK_BG)
                elif pr < 85:
                    _shade_cell(cell, WARN_BG)
            p = cell.paragraphs[0]
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after = Pt(2)
            if i in {0, 2, 3, 6}:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(val)
            run.font.size = Pt(8)
            run.font.name = "Calibri"
            if i == 4 and r_i < len(cat_tones):
                run.font.color.rgb = RGBColor.from_string(cat_tones[r_i])
                run.bold = True
            elif i == 5 and r_i < len(risk_tones):
                run.font.color.rgb = RGBColor.from_string(risk_tones[r_i])
            else:
                run.font.color.rgb = RGBColor(0x1A, 0x23, 0x32)
            _no_snap_to_grid(p)


# ─── card primitives ───────────────────────────────────────────────


def _labeled_card(doc, *, num: int, badge: str, text: str, accent: str, bg: str) -> None:
    if not str(text).strip():
        return
    from docx.shared import Cm, Pt, RGBColor

    t = doc.add_table(rows=1, cols=3)
    t.autofit = False
    n, a, b = t.rows[0].cells
    _set_cell_width(n, Cm(0.9))
    _set_cell_width(a, Cm(0.18))
    _set_cell_width(b, Cm(15.6))
    _shade_cell(n, accent)
    _shade_cell(a, accent)
    _shade_cell(b, bg)
    pn = n.paragraphs[0]
    rn = pn.add_run(f"{num:02d}")
    rn.bold = True
    rn.font.size = Pt(11)
    rn.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    p = b.paragraphs[0]
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(2)
    br = p.add_run(f"{badge}  ")
    br.bold = True
    br.font.size = Pt(7)
    br.font.color.rgb = RGBColor.from_string(accent)
    lines = str(text).split("\n")
    run = p.add_run(lines[0])
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x1A, 0x23, 0x32)
    for line in lines[1:]:
        p2 = b.add_paragraph()
        p2.paragraph_format.space_before = Pt(0)
        p2.paragraph_format.space_after = Pt(2)
        r2 = p2.add_run(line)
        r2.font.size = Pt(9)
        r2.font.color.rgb = RGBColor(0x5D, 0x6D, 0x7E)


def _info_card(doc, text: str) -> None:
    if not str(text).strip():
        return
    from docx.shared import Cm, Pt, RGBColor

    t = doc.add_table(rows=1, cols=2)
    t.autofit = False
    a, b = t.rows[0].cells
    _set_cell_width(a, Cm(0.18))
    _set_cell_width(b, Cm(16.5))
    _shade_cell(a, SKY)
    _shade_cell(b, INFO_BG)
    p = b.paragraphs[0]
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    badge = p.add_run("Информация  ")
    badge.bold = True
    badge.font.size = Pt(7)
    badge.font.color.rgb = RGBColor(0x1B, 0x4F, 0x72)
    run = p.add_run(text)
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x1A, 0x23, 0x32)


def _risk_card(doc, num: int, text: str, badge: str = "Риск") -> None:
    _labeled_card(doc, num=num, badge=badge, text=text, accent=RISK, bg=RISK_BG)


def _strength_card(doc, num: int, text: str) -> None:
    _labeled_card(doc, num=num, badge="Сильная сторона", text=text, accent=GOOD, bg=GOOD_BG)


def _plan_card(doc, num: int, text: str) -> None:
    _labeled_card(doc, num=num, badge="Приоритет", text=text, accent=BLUE, bg=INFO_BG)


# ─── page chrome ───────────────────────────────────────────────────


def _save(doc) -> BytesIO:
    out = BytesIO()
    doc.save(out)
    out.seek(0)
    return out


def _fmt_delta(v) -> str:
    if v is None:
        return "—"
    try:
        f = float(v)
    except (TypeError, ValueError):
        return "—"
    return f"+{f}" if f > 0 else str(f)


def _tone_avg(v, et: str) -> str:
    try:
        f = float(v or 0)
    except (TypeError, ValueError):
        return SKY
    if et == "oge":
        if f >= 4:
            return GOOD
        if f < 3.5:
            return RISK
        return WARN
    if f >= 60:
        return GOOD
    if f < 45:
        return RISK
    return WARN


def _tone_pct(v, high: float, mid: float) -> str:
    try:
        f = float(v or 0)
    except (TypeError, ValueError):
        return SKY
    if f >= high:
        return GOOD
    if f < mid:
        return RISK
    return WARN


def _tone_delta(v) -> str:
    if v is None:
        return SKY
    try:
        f = float(v)
    except (TypeError, ValueError):
        return SKY
    if f > 0:
        return GOOD
    if f < 0:
        return WARN
    return SKY


def _setup_page(doc) -> None:
    from docx.oxml.ns import qn
    from docx.shared import Cm, Pt

    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.left_margin = Cm(1.6)
    section.right_margin = Cm(1.6)
    section.top_margin = Cm(1.7)
    section.bottom_margin = Cm(1.7)
    section.header_distance = Cm(0.8)
    section.footer_distance = Cm(0.8)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(10)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
    pf = normal.paragraph_format
    pf.space_after = Pt(4)
    pf.space_before = Pt(0)
    pf.line_spacing = 1.15


def _setup_header_footer(doc, *, district: str, exam: str, year, formed: str) -> None:
    from docx.shared import Pt, RGBColor

    section = doc.sections[0]
    header = section.header
    header.is_linked_to_previous = False
    hp = header.paragraphs[0]
    hp.clear()
    r = hp.add_run(
        f"Сравнительный анализ ОО  ·  {district}  ·  {exam}  ·  {year or '—'}"
    )
    r.font.size = Pt(8)
    r.font.color.rgb = RGBColor(0x5D, 0x6D, 0x7E)
    r.font.name = "Calibri"
    _p_border(hp, bottom=True)

    footer = section.footer
    footer.is_linked_to_previous = False
    fp = footer.paragraphs[0]
    fp.clear()
    run = fp.add_run(f"Дата формирования: {formed}  ·  версия {DOC_VERSION}  ·  стр. ")
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(0x5D, 0x6D, 0x7E)
    _add_page_field(fp)


def _add_page_field(paragraph) -> None:
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.shared import Pt, RGBColor

    fld = OxmlElement("w:fldChar")
    fld.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    fld2 = OxmlElement("w:fldChar")
    fld2.set(qn("w:fldCharType"), "end")
    r = paragraph.add_run()
    r._r.append(fld)
    r2 = paragraph.add_run()
    r2._r.append(instr)
    r3 = paragraph.add_run()
    r3._r.append(fld2)
    for rr in (r, r2, r3):
        rr.font.size = Pt(8)
        rr.font.color.rgb = RGBColor(0x5D, 0x6D, 0x7E)


def _title_page(doc, *, district: str, exam: str, year, formed: str) -> None:
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Cm, Pt, RGBColor

    logo = doc.add_table(rows=1, cols=2)
    logo.autofit = False
    c0, c1 = logo.rows[0].cells
    _set_cell_width(c0, Cm(1.4))
    _set_cell_width(c1, Cm(15))
    _shade_cell(c0, NAVY)
    p0 = c0.paragraphs[0]
    p0.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r0 = p0.add_run("АГ")
    r0.bold = True
    r0.font.size = Pt(14)
    r0.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    p1 = c1.paragraphs[0]
    r1 = p1.add_run("УПРАВЛЕНИЕ ОБРАЗОВАНИЯ")
    r1.bold = True
    r1.font.size = Pt(10)
    r1.font.color.rgb = RGBColor(0x1B, 0x4F, 0x72)
    p1b = c1.add_paragraph()
    r1b = p1b.add_run(district)
    r1b.font.size = Pt(12)
    r1b.bold = True
    r1b.font.color.rgb = RGBColor(0x0B, 0x1F, 0x3A)

    for _ in range(3):
        doc.add_paragraph("")

    eye = doc.add_paragraph()
    er = eye.add_run("ОФИЦИАЛЬНЫЙ АНАЛИТИЧЕСКИЙ ОТЧЁТ")
    er.bold = True
    er.font.size = Pt(9)
    er.font.color.rgb = RGBColor(0x5B, 0x9B, 0xD5)

    title = doc.add_paragraph()
    title.paragraph_format.space_before = Pt(8)
    tr = title.add_run("Сравнительный анализ\nобразовательных организаций")
    tr.bold = True
    tr.font.size = Pt(26)
    tr.font.color.rgb = RGBColor(0x0B, 0x1F, 0x3A)
    tr.font.name = "Calibri"

    sub = doc.add_paragraph()
    sr = sub.add_run(f"Муниципальный сравнительный отчёт по итогам {exam}")
    sr.font.size = Pt(12)
    sr.font.color.rgb = RGBColor(0x5D, 0x6D, 0x7E)

    for _ in range(2):
        doc.add_paragraph("")

    meta = doc.add_table(rows=5, cols=2)
    meta.autofit = False
    pairs = [
        ("Вид экзамена", exam),
        ("Отчётный период", str(year or "—")),
        ("Муниципалитет", district),
        ("Предмет охвата", "Все предметы среза"),
        ("Дата формирования", formed),
    ]
    for i, (k, v) in enumerate(pairs):
        ck, cv = meta.rows[i].cells
        _set_cell_width(ck, Cm(4.5))
        _set_cell_width(cv, Cm(11.5))
        pk = ck.paragraphs[0]
        rk = pk.add_run(k)
        rk.font.size = Pt(9)
        rk.font.color.rgb = RGBColor(0x5D, 0x6D, 0x7E)
        pv = cv.paragraphs[0]
        rv = pv.add_run(v)
        rv.bold = True
        rv.font.size = Pt(12)
        rv.font.color.rgb = RGBColor(0x0B, 0x1F, 0x3A)

    for _ in range(2):
        doc.add_paragraph("")

    ver = doc.add_paragraph()
    vr = ver.add_run(f"Версия документа {DOC_VERSION}")
    vr.font.size = Pt(9)
    vr.font.color.rgb = RGBColor(0x5D, 0x6D, 0x7E)

    conf = doc.add_paragraph()
    cr = conf.add_run("Конфиденциально · для служебного пользования")
    cr.font.size = Pt(9)
    cr.font.color.rgb = RGBColor(0x5D, 0x6D, 0x7E)

    doc.add_page_break()


def _toc(doc) -> None:
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.shared import Pt, RGBColor

    h = doc.add_paragraph()
    r = h.add_run("Содержание")
    r.bold = True
    r.font.size = Pt(16)
    r.font.color.rgb = RGBColor(0x0B, 0x1F, 0x3A)

    p = doc.add_paragraph()
    run = p.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = ' TOC \\o "1-2" \\h \\z \\u '
    fld_separate = OxmlElement("w:fldChar")
    fld_separate.set(qn("w:fldCharType"), "separate")
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_begin)
    run._r.append(instr)
    run._r.append(fld_separate)
    hint = p.add_run("(поле содержания обновляется в Word)")
    hint.font.size = Pt(8)
    hint.font.color.rgb = RGBColor(0x5D, 0x6D, 0x7E)
    run2 = p.add_run()
    run2._r.append(fld_end)
    doc.add_page_break()


def _section_block(doc, title: str, icon: str, *, accent: str) -> None:
    from docx.oxml import OxmlElement
    from docx.shared import Cm, Pt, RGBColor

    wrap = doc.add_table(rows=1, cols=2)
    wrap.autofit = False
    a, b = wrap.rows[0].cells
    _set_cell_width(a, Cm(0.25))
    _set_cell_width(b, Cm(16.5))
    _shade_cell(a, accent)
    p = b.paragraphs[0]
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(2)
    pPr = p._p.get_or_add_pPr()
    pPr.append(OxmlElement("w:keepNext"))
    ir = p.add_run(f"{icon}  ")
    ir.font.size = Pt(11)
    ir.font.color.rgb = RGBColor.from_string(accent)
    tr = p.add_run(title)
    tr.bold = True
    tr.font.size = Pt(13)
    tr.font.color.rgb = RGBColor(0x0B, 0x1F, 0x3A)
    _p_border(p, bottom=True)


def _kpi_dashboard(doc, cards: list[tuple[str, str, str, str]]) -> None:
    from docx.shared import Cm, Pt, RGBColor

    for row_i in range(0, len(cards), 4):
        chunk = cards[row_i : row_i + 4]
        while len(chunk) < 4:
            chunk.append(("", "", "", SKY))
        table = doc.add_table(rows=1, cols=4)
        table.autofit = False
        for col, (icon, label, value, tone) in enumerate(chunk):
            cell = table.rows[0].cells[col]
            _set_cell_width(cell, Cm(4.2))
            _shade_cell(cell, SOFT)
            _cell_left_border(cell, tone)
            p_icon = cell.paragraphs[0]
            p_icon.paragraph_format.space_after = Pt(0)
            p_icon.paragraph_format.space_before = Pt(4)
            if icon:
                ri = p_icon.add_run(icon)
                ri.font.size = Pt(9)
                ri.font.color.rgb = RGBColor.from_string(tone)
            p_val = cell.add_paragraph()
            p_val.paragraph_format.space_before = Pt(2)
            p_val.paragraph_format.space_after = Pt(0)
            rv = p_val.add_run(value)
            rv.bold = True
            rv.font.size = Pt(16)
            rv.font.color.rgb = RGBColor(0x0B, 0x1F, 0x3A)
            p_lab = cell.add_paragraph()
            p_lab.paragraph_format.space_before = Pt(1)
            p_lab.paragraph_format.space_after = Pt(4)
            rl = p_lab.add_run(label)
            rl.font.size = Pt(8)
            rl.font.color.rgb = RGBColor(0x5D, 0x6D, 0x7E)
        doc.add_paragraph("").paragraph_format.space_after = Pt(2)


# ─── low-level helpers ─────────────────────────────────────────────


def _set_row_header(row) -> None:
    from docx.oxml import OxmlElement

    tr = row._tr
    trPr = tr.get_or_add_trPr()
    trPr.append(OxmlElement("w:tblHeader"))


def _cant_split_row(row) -> None:
    from docx.oxml import OxmlElement

    tr = row._tr
    trPr = tr.get_or_add_trPr()
    trPr.append(OxmlElement("w:cantSplit"))


def _no_snap_to_grid(paragraph) -> None:
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    pPr = paragraph._p.get_or_add_pPr()
    snap = OxmlElement("w:snapToGrid")
    snap.set(qn("w:val"), "0")
    pPr.append(snap)


def _set_cell_width(cell, width) -> None:
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcW = tcPr.find(qn("w:tcW"))
    if tcW is None:
        tcW = OxmlElement("w:tcW")
        tcPr.append(tcW)
    twips = int(getattr(width, "twips", width))
    tcW.set(qn("w:w"), str(twips))
    tcW.set(qn("w:type"), "dxa")
    cell.width = width


def _cell_left_border(cell, color_hex: str) -> None:
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for edge in ("top", "left", "bottom", "right"):
        el = OxmlElement(f"w:{edge}")
        if edge == "left":
            el.set(qn("w:val"), "single")
            el.set(qn("w:sz"), "24")
            el.set(qn("w:color"), color_hex)
        else:
            el.set(qn("w:val"), "nil")
        tcBorders.append(el)
    tcPr.append(tcBorders)


def _p_border(paragraph, *, bottom: bool = False) -> None:
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    pPr = paragraph._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    if bottom:
        el = OxmlElement("w:bottom")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), "6")
        el.set(qn("w:space"), "4")
        el.set(qn("w:color"), LINE)
        pBdr.append(el)
    pPr.append(pBdr)


def _shade_cell(cell, fill_hex: str) -> None:
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill_hex)
    shd.set(qn("w:val"), "clear")
    tcPr.append(shd)
