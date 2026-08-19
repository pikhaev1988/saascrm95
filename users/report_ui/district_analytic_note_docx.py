"""
DOCX visual shell for municipal GIA analytic note.

Presentation/formatting only — text and section order come from the same UI payload.
"""

from __future__ import annotations

from datetime import date
from io import BytesIO
from typing import Any

# Brand palette (RGB)
NAVY = "0B1F3A"
BLUE = "1B4F72"
SKY = "5B9BD5"
GRAY = "5D6D7E"
LINE = "D5DDE5"
WHITE = "FFFFFF"
FACT_BG = "E8F2FA"
CONCL_BG = "DCEAF5"
HYP_BG = "EEE6F5"
DEC_BG = "E4F4EA"
EFF_BG = "E2F5F2"


def render_district_analytic_note_docx(payload: dict[str, Any], ui: dict[str, Any]) -> BytesIO:
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.shared import Cm, Pt, RGBColor

    doc = Document()
    if not (payload or {}).get("has_data") or not (ui or {}).get("has_data"):
        doc.add_paragraph((payload or ui or {}).get("message") or "Недостаточно данных.")
        out = BytesIO()
        doc.save(out)
        out.seek(0)
        return out

    district = ui.get("district_name") or payload.get("district_name") or "Муниципалитет"
    exam = ui.get("exam_label") or "ГИА"
    year = ui.get("year") or ""
    avg_label = ui.get("avg_label") or "Средний балл"
    formed = date.today().strftime("%d.%m.%Y")

    _setup_page(doc)
    _setup_header_footer(doc, district=district, year=year, formed=formed)
    _add_title_page(doc, district=district, exam=exam, year=year, formed=formed, ui=ui)
    _add_toc(doc)

    passport = ui.get("passport") or {}
    if passport:
        _section_heading(doc, "Управленческий паспорт муниципалитета")
        level = passport.get("system_level") or ui.get("system_level") or {}
        doc.add_paragraph(f"Общая оценка системы: {level.get('icon', '')} {level.get('label', '—')}")
        if passport.get("snapshot"):
            doc.add_paragraph(str(passport["snapshot"]))
        _bullets(doc, "Рейтинг муниципалитета:", passport.get("rating_lines") or [])
        _bullets(doc, "Основные достижения:", passport.get("achievements") or [])
        _bullets(doc, "Основные риски:", passport.get("risks") or [])
        _bullets(
            doc,
            "Ключевые предметы:",
            [
                f"{s.get('name')}: усп. {s.get('pass_rate')}%, вклад {s.get('weight')}%"
                for s in (passport.get("key_subjects") or [])
            ],
        )
        _bullets(
            doc,
            "Школы-лидеры:",
            [
                f"{s.get('name')}: ср. {s.get('avg')}, усп. {s.get('pass_rate')}%"
                for s in (passport.get("leader_schools") or [])
            ],
        )
        _bullets(
            doc,
            "Школы приоритетного сопровождения:",
            [
                f"{s.get('name')}: ср. {s.get('avg')}, усп. {s.get('pass_rate')}%"
                for s in (passport.get("support_schools") or [])
            ],
        )
        _bullets(doc, "Потенциал роста:", passport.get("growth_potential") or [])
        doc.add_paragraph("Три главных управленческих решения:")
        for d in passport.get("top_decisions") or []:
            p = doc.add_paragraph(
                f"{d.get('action')} (проблема: {d.get('problem')}; эффект: {d.get('effect')})",
                style="List Number",
            )
            _style_body(p)

    _section_heading(doc, "1. Управленческое резюме")
    level = ui.get("system_level") or {}
    doc.add_paragraph(f"Уровень муниципальной системы: {level.get('icon', '')} {level.get('label', '—')}")
    exec_ = ui.get("executive") or {}
    _bullets(doc, "Основные достижения:", exec_.get("achievements") or [])
    _bullets(doc, "Основные риски:", exec_.get("risks") or [])
    doc.add_paragraph(f"Главный управленческий вывод: {exec_.get('main_conclusion') or '—'}")
    _add_chain(doc, exec_.get("chain") or [])

    kpi = ui.get("kpi") or {}
    doc.add_paragraph(
        f"Участники: {kpi.get('participants')}; результатов: {kpi.get('total_results')}; "
        f"ОО: {kpi.get('schools_count')}; предметы: {kpi.get('subjects_count')}; "
        f"средний: {kpi.get('avg_score')}; успеваемость: {kpi.get('pass_rate')}%; "
        f"качество: {kpi.get('quality_rate')}%; высокобалльники: {kpi.get('high_count')}; "
        f"школ риска: {kpi.get('risk_count')}."
    )

    regional = ui.get("regional_position") if ui.get("has_regional_position") else None
    if regional:
        _section_heading(doc, "2. Положение муниципалитета")
        doc.add_paragraph(f"Место среди {regional.get('districts_total')} муниципалитетов региона.")
        positions = regional.get("positions") or {}
        labels = {
            "avg_score": avg_label,
            "quality_rate": "Качество знаний, %",
            "pass_rate": "Успеваемость, %",
            "high_count": "Высокобалльники",
            "avg_delta": "Динамика среднего",
        }
        rows = []
        for key, title in labels.items():
            pos = positions.get(key)
            if not pos:
                continue
            rows.append([title, str(pos.get("value") if pos.get("value") is not None else "—"), str(pos.get("place") or "—"), str(pos.get("total") or "—")])
        _add_table(doc, ["Показатель", "Значение", "Место", "из"], rows)

    clusters = ui.get("clusters") or []
    if clusters:
        _section_heading(doc, "3. Кластеры образовательных организаций")
        for c in clusters:
            doc.add_paragraph(
                f"{c.get('title')}: школ {c.get('count')}; "
                f"средний {c.get('avg') if c.get('avg') is not None else '—'}; "
                f"доля участников {c.get('share')}%. {c.get('feature') or ''}"
            )

    contribution = ui.get("contribution") or {}
    if contribution:
        _section_heading(doc, "Анализ вклада")
        _add_chain(doc, contribution.get("chain") or [])
        large = contribution.get("large_schools") or []
        if large:
            doc.add_paragraph("Крупные школы (доля участников):")
            for s in large:
                doc.add_paragraph(
                    f"{s.get('name')}: доля {s.get('weight')}%, средний {s.get('avg')}, "
                    f"вклад в средний {s.get('pull')}",
                    style="List Bullet",
                )
        sig = contribution.get("significant_subjects") or []
        if sig:
            doc.add_paragraph("Значимые предметы (доля результатов):")
            for s in sig:
                doc.add_paragraph(
                    f"{s.get('name')}: доля {s.get('weight')}%, средний {s.get('avg')}, "
                    f"вклад в средний {s.get('pull')}",
                    style="List Bullet",
                )

    system_analysis = ui.get("system_analysis") or {}
    if system_analysis:
        _section_heading(doc, "Системный анализ муниципалитета")
        _bullets(doc, "Локальные проблемы:", system_analysis.get("local_problems") or [])
        _bullets(doc, "Системные проблемы:", system_analysis.get("systemic_problems") or [])
        _bullets(doc, "Решения с наибольшим эффектом:", system_analysis.get("high_effect_decisions") or [])
        _add_chain(doc, system_analysis.get("chain") or [])

    sd = ui.get("school_deviation") or {}
    if sd.get("has_data"):
        _section_heading(doc, "4. Анализ школ муниципалитета")
        _add_chain(doc, sd.get("chain") or [])

    sp = ui.get("system_profile") or {}
    if sp:
        _section_heading(doc, "5. Анализ муниципальной системы")
        traits = sp.get("traits") or []
        if traits:
            doc.add_paragraph("Характеристики: " + ", ".join(str(t) for t in traits))
        _add_chain(doc, sp.get("chain") or [])

    ss = ui.get("subject_structure") or {}
    if ss:
        _section_heading(doc, "6. Структура предметных результатов")
        strong = ss.get("strong") or []
        attention = ss.get("priority_attention") or []
        if strong:
            doc.add_paragraph("Сильные: " + ", ".join(f"{s.get('name')} ({s.get('pass_rate')}%)" for s in strong))
        if attention:
            doc.add_paragraph(
                "Требуют внимания: " + ", ".join(f"{s.get('name')} ({s.get('pass_rate')}%)" for s in attention)
            )
        _add_chain(doc, ss.get("chain") or [])
        subjects = ui.get("subjects") or []
        if subjects:
            rows = []
            for row in subjects:
                name = str(row.get("name") or "")
                if not row.get("statistically_stable"):
                    name += " (малая выборка)"
                rows.append(
                    [
                        name,
                        str(row.get("participants") or ""),
                        str(row.get("avg") or ""),
                        str(row.get("pass_rate") or ""),
                        str(row.get("status") or ""),
                    ]
                )
            _add_table(doc, ["Предмет", "Участ.", "Средний", "Усп., %", "Статус"], rows)

    st = ui.get("statistical_stability") or {}
    if st:
        _section_heading(doc, "7. Статистическая устойчивость")
        _add_chain(doc, st.get("chain") or [])

    mentoring = ui.get("mentoring") or {}
    if mentoring:
        _section_heading(doc, "8. Наставничество между ОО")
        _add_chain(doc, mentoring.get("chain") or [])
        pairs = mentoring.get("pairs") or []
        if pairs:
            rows = [[str(p.get("mentor") or ""), str(p.get("mentee") or ""), str(p.get("basis") or "")] for p in pairs]
            _add_table(doc, ["Наставник", "Сопровождение", "Основание"], rows)

    priorities = ui.get("priorities") or []
    if priorities:
        _section_heading(doc, "9. Приоритетные направления муниципальной работы")
        doc.add_paragraph("Системная приоритизация: тяжесть проблемы × масштаб влияния.")
        rows = []
        for t in priorities:
            rows.append(
                [
                    str(t.get("priority") or ""),
                    str(t.get("problem") or ""),
                    str(t.get("basis") or ""),
                    str(t.get("impact_scale") or ""),
                    str(t.get("expected_effect") or ""),
                    str(t.get("justification") or ""),
                    str(t.get("owner") or ""),
                ]
            )
        _add_table(
            doc,
            ["№", "Проблема", "Основание", "Масштаб", "Эффект", "Обоснование", "Ответственный"],
            rows,
        )

    risk_map = ui.get("risk_map") or []
    if risk_map:
        _section_heading(doc, "10. Муниципальная карта рисков")
        rows = [
            [
                str(r.get("category") or ""),
                str(r.get("risk") or ""),
                str(r.get("probability") or ""),
                str(r.get("impact") or ""),
                str(r.get("priority") or ""),
            ]
            for r in risk_map
        ]
        _add_table(doc, ["Категория", "Риск", "Вероятность", "Влияние", "Приоритет"], rows)

    effects = ui.get("expected_effects") or []
    if effects:
        _section_heading(doc, "11. Ожидаемый эффект реализации мероприятий")
        for e in effects:
            doc.add_paragraph(
                f"{e.get('action')}: улучшение — {e.get('improves')}; "
                f"снижение рисков — {e.get('reduces_risks')}; "
                f"эффект — {e.get('effect_level')}."
            )

    dynamics = ui.get("dynamics") or []
    if dynamics:
        _section_heading(doc, "12. Динамика результатов")
        rows = [
            [
                str(row.get("year") or ""),
                str(row.get("participants") or ""),
                str(row.get("results") or ""),
                str(row.get("avg") or ""),
                str(row.get("pass_rate") or ""),
            ]
            for row in dynamics
        ]
        _add_table(doc, ["Год", "Участников", "Результатов", "Средний", "Усп., %"], rows)

    decision_map = ui.get("decision_map") or []
    if decision_map:
        _section_heading(doc, "13. Карта управленческих решений")
        doc.add_paragraph("Итоговый управленческий документ.")
        rows = [
            [
                str(row.get("priority") or ""),
                str(row.get("problem") or ""),
                str(row.get("action") or ""),
                str(row.get("owner") or ""),
                str(row.get("term") or ""),
                str(row.get("effect") or ""),
                str(row.get("status") or ""),
            ]
            for row in decision_map
        ]
        _add_table(
            doc,
            ["Приоритет", "Проблема", "Действие", "Ответственный", "Срок", "Эффект", "Статус"],
            rows,
        )

    out = BytesIO()
    doc.save(out)
    out.seek(0)
    return out


def _setup_page(doc) -> None:
    from docx.shared import Cm, Pt
    from docx.oxml.ns import qn

    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.left_margin = Cm(2.0)
    section.right_margin = Cm(1.8)
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.header_distance = Cm(1.0)
    section.footer_distance = Cm(1.0)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
    pf = normal.paragraph_format
    pf.space_after = Pt(6)
    pf.line_spacing = 1.15


def _setup_header_footer(doc, *, district: str, year, formed: str) -> None:
    from docx.shared import Pt, RGBColor
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    section = doc.sections[0]
    header = section.header
    header.is_linked_to_previous = False
    hp = header.paragraphs[0]
    hp.text = f"Аналитическая справка по итогам ГИА  ·  {district}  ·  {year or '—'}"
    if hp.runs:
        hp.runs[0].font.size = Pt(8)
        hp.runs[0].font.color.rgb = RGBColor(0x5D, 0x6D, 0x7E)
    _add_bottom_border(hp)

    footer = section.footer
    footer.is_linked_to_previous = False
    fp = footer.paragraphs[0]
    run = fp.add_run(f"Дата формирования: {formed}  ·  Конфиденциально  ·  стр. ")
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(0x5D, 0x6D, 0x7E)
    # PAGE field
    fld = OxmlElement("w:fldChar")
    fld.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    fld2 = OxmlElement("w:fldChar")
    fld2.set(qn("w:fldCharType"), "end")
    r = fp.add_run()
    r._r.append(fld)
    r2 = fp.add_run()
    r2._r.append(instr)
    r3 = fp.add_run()
    r3._r.append(fld2)
    for rr in (r, r2, r3):
        rr.font.size = Pt(8)
        rr.font.color.rgb = RGBColor(0x5D, 0x6D, 0x7E)


def _add_bottom_border(paragraph) -> None:
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    pPr = paragraph._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "4")
    bottom.set(qn("w:color"), LINE)
    pBdr.append(bottom)
    pPr.append(pBdr)


def _add_title_page(doc, *, district: str, exam: str, year, formed: str, ui: dict) -> None:
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Pt, RGBColor

    # navy band simulation
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run("УПРАВЛЕНИЕ ОБРАЗОВАНИЯ")
    run.bold = True
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x1B, 0x4F, 0x72)
    run.font.name = "Calibri"

    p2 = doc.add_paragraph()
    r2 = p2.add_run(district)
    r2.font.size = Pt(12)
    r2.font.color.rgb = RGBColor(0x0B, 0x1F, 0x3A)
    r2.bold = True

    for _ in range(3):
        doc.add_paragraph("")

    eye = doc.add_paragraph()
    er = eye.add_run("ОФИЦИАЛЬНЫЙ АНАЛИТИЧЕСКИЙ ОТЧЁТ")
    er.font.size = Pt(9)
    er.font.color.rgb = RGBColor(0x5B, 0x9B, 0xD5)
    er.bold = True

    title = doc.add_paragraph()
    tr = title.add_run("Аналитическая справка\nпо итогам ГИА")
    tr.bold = True
    tr.font.size = Pt(28)
    tr.font.color.rgb = RGBColor(0x0B, 0x1F, 0x3A)
    tr.font.name = "Calibri"

    for _ in range(2):
        doc.add_paragraph("")

    meta = [
        ("Вид экзамена", str(exam)),
        ("Отчётный период", str(year or "—")),
        ("Муниципалитет", str(district)),
        ("Дата формирования", formed),
    ]
    level = (ui.get("system_level") or {})
    meta.append(("Уровень системы", f"{level.get('icon', '')} {level.get('label', '—')}".strip()))
    for k, v in meta:
        row = doc.add_paragraph()
        rk = row.add_run(f"{k}:  ")
        rk.font.size = Pt(10)
        rk.font.color.rgb = RGBColor(0x5D, 0x6D, 0x7E)
        rv = row.add_run(v)
        rv.bold = True
        rv.font.size = Pt(12)
        rv.font.color.rgb = RGBColor(0x0B, 0x1F, 0x3A)

    for _ in range(4):
        doc.add_paragraph("")
    conf = doc.add_paragraph()
    cr = conf.add_run("Конфиденциально · для служебного пользования")
    cr.font.size = Pt(9)
    cr.font.color.rgb = RGBColor(0x5D, 0x6D, 0x7E)

    doc.add_page_break()


def _add_toc(doc) -> None:
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.shared import Pt, RGBColor

    h = doc.add_paragraph()
    r = h.add_run("Содержание")
    r.bold = True
    r.font.size = Pt(16)
    r.font.color.rgb = RGBColor(0x0B, 0x1F, 0x3A)

    p = doc.add_paragraph()
    run = p.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = ' TOC \\o "1-2" \\h \\z \\u '
    fld_separate = OxmlElement("w:fldChar")
    fld_separate.set(qn("w:fldCharType"), "separate")
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_begin)
    run._r.append(instr)
    run._r.append(fld_separate)
    hint = p.add_run("Обновите поле содержания в Word (ПКМ → Обновить поле).")
    hint.font.size = Pt(9)
    hint.font.color.rgb = RGBColor(0x5D, 0x6D, 0x7E)
    run2 = p.add_run()
    run2._r.append(fld_end)

    doc.add_page_break()


def _section_heading(doc, text: str) -> None:
    from docx.shared import Pt, RGBColor
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(8)
    # keep with next
    pPr = p._p.get_or_add_pPr()
    keep = OxmlElement("w:keepNext")
    pPr.append(keep)
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(0x0B, 0x1F, 0x3A)
    run.font.name = "Calibri"
    _add_bottom_border(p)


def _style_body(p) -> None:
    from docx.shared import Pt

    for run in p.runs:
        run.font.size = Pt(11)
        run.font.name = "Calibri"


def _bullets(doc, title: str, items: list) -> None:
    if title:
        doc.add_paragraph(title)
    for item in items:
        p = doc.add_paragraph(str(item), style="List Bullet")
        _style_body(p)


def _add_chain(doc, chain: list[dict[str, Any]]) -> None:
    from docx.shared import Pt, RGBColor

    colors = {
        "fact": (FACT_BG, "2E75B6"),
        "conclusion": (CONCL_BG, "1B4F72"),
        "hypothesis": (HYP_BG, "6B4F8A"),
        "decision": (DEC_BG, "2F6B4F"),
        "effect": (EFF_BG, "1F6F66"),
    }
    for line in chain or []:
        kind = line.get("kind") or "fact"
        bg, accent = colors.get(kind, (FACT_BG, BLUE))
        label = line.get("label") or ""
        text = line.get("text") or ""
        p = doc.add_paragraph()
        _shade_paragraph(p, bg)
        r1 = p.add_run(f"[{label}] ")
        r1.bold = True
        r1.font.size = Pt(10)
        r1.font.color.rgb = RGBColor.from_string(accent)
        r2 = p.add_run(text)
        r2.font.size = Pt(11)
        r2.font.color.rgb = RGBColor(0x1A, 0x23, 0x32)


def _shade_paragraph(paragraph, fill_hex: str) -> None:
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    pPr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill_hex)
    shd.set(qn("w:val"), "clear")
    pPr.append(shd)


def _add_table(doc, headers: list[str], rows: list[list[str]]) -> None:
    from docx.shared import Pt, RGBColor, Cm
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.autofit = True
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        _shade_cell(hdr[i], NAVY)
        for p in hdr[i].paragraphs:
            for run in p.runs:
                run.bold = True
                run.font.size = Pt(9)
                run.font.color.rgb = RGBColor(0xDC, 0xE8, 0xF3)
                run.font.name = "Calibri"
    for r_i, row in enumerate(rows):
        cells = table.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = val
            if r_i % 2 == 1:
                _shade_cell(cells[i], "F7FAFC")
            for p in cells[i].paragraphs:
                for run in p.runs:
                    run.font.size = Pt(9)
                    run.font.name = "Calibri"
    # spacing after
    doc.add_paragraph("")


def _shade_cell(cell, fill_hex: str) -> None:
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill_hex)
    shd.set(qn("w:val"), "clear")
    tcPr.append(shd)
