"""
Презентационный рендер DOCX для «Свод результатов ГИА (ОО)».

Только внешний вид Word-документа поверх уже готового summary_ui.
Без SQL, пересчётов и изменения исходных показателей.
"""

from __future__ import annotations

from typing import Any


# Палитра BI-отчёта
_C = {
    "ink": "102A43",
    "muted": "627D98",
    "line": "D9E2EC",
    "bg": "F0F4F8",
    "white": "FFFFFF",
    "blue": "2B6CB0",
    "blue_bg": "EBF8FF",
    "good": "2F855A",
    "good_bg": "F0FFF4",
    "mid": "B7791F",
    "mid_bg": "FFFBEB",
    "warn": "C05621",
    "warn_bg": "FFFAF0",
    "low": "C53030",
    "low_bg": "FFF5F5",
    "title_bg": "1A365D",
}


def render_gia_summary_docx(doc, ui: dict[str, Any], *, exam_label: str, year, avg_label: str, class_rows: list | None = None) -> None:
    """Собирает BI-layout внутри уже созданного Document."""
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.shared import Cm, Pt, RGBColor

    _setup_page(doc)
    executive = ui.get("executive") or {}
    grade = ui.get("school_grade") or {}
    tones = executive.get("kpi_tones") or {}

    # ===== Title banner =====
    year_bit = f" · {year}" if year else ""
    _banner(
        doc,
        f"СВОД РЕЗУЛЬТАТОВ ГИА (ОО)  ·  {exam_label}{year_bit}",
        "Аналитическая панель образовательной организации",
    )

    # ===== Module: Панель руководителя =====
    _module_title(doc, "01", "Панель руководителя")

    # Grade hero
    grade_tone = grade.get("tone") or "mid"
    grade_table = doc.add_table(rows=1, cols=1)
    gcell = grade_table.rows[0].cells[0]
    _shade(gcell, _tone_bg(grade_tone))
    _set_cell_margins(gcell, 140, 140, 140, 140)
    p = gcell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("ОБЩАЯ ОЦЕНКА РЕЗУЛЬТАТОВ")
    _font(r, size=9, bold=True, color=_C["muted"])
    p2 = gcell.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = p2.add_run(grade.get("stars_display") or "★★★☆☆")
    _font(r2, size=22, bold=True, color="D69E2E")
    p3 = gcell.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r3 = p3.add_run(grade.get("label") or executive.get("overall") or "")
    _font(r3, size=16, bold=True, color=_tone_fg(grade_tone))
    _nofit(grade_table)
    doc.add_paragraph()

    # KPI row
    kpi_items = [
        ("👥", str(executive.get("participants") or 0), "Участников", tones.get("participants", "neutral")),
        ("📈", str(executive.get("avg_score") if executive.get("avg_score") is not None else "—"), avg_label, tones.get("avg", "mid")),
        ("🎯", f"{executive.get('quality_rate')}%", "Качество знаний", tones.get("quality", "mid")),
        ("✅", f"{executive.get('pass_rate')}%", "Успеваемость", tones.get("pass", "mid")),
    ]
    if executive.get("total_results") and executive.get("total_results") != executive.get("participants"):
        # вторая строка: результаты как отдельная мини-карточка рядом не нужна в 4 KPI — покажем в chips
        pass

    kpi = doc.add_table(rows=1, cols=4)
    for i, (icon, value, label, tone) in enumerate(kpi_items):
        cell = kpi.rows[0].cells[i]
        _fill_kpi_card(cell, icon=icon, value=value, label=label, tone=tone)
    _nofit(kpi)
    doc.add_paragraph()

    # Side chips: best / problem / risk / results
    chips = []
    if executive.get("total_results"):
        chips.append(("📄 Результатов", str(executive.get("total_results")), "neutral"))
    if executive.get("best_subject"):
        chips.append(("🟢 Лучший предмет", str(executive["best_subject"]), "high"))
    if executive.get("worst_subject"):
        chips.append(("🔴 Проблемный предмет", str(executive["worst_subject"]), "low"))
    if executive.get("main_risk"):
        chips.append(("⚠ Основной риск", str(executive["main_risk"]), "warn"))
    if chips:
        n = min(len(chips), 4)
        chip_t = doc.add_table(rows=1, cols=n)
        for i in range(n):
            title, value, tone = chips[i]
            cell = chip_t.rows[0].cells[i]
            _shade(cell, _tone_bg(tone))
            _set_cell_margins(cell, 80, 80, 80, 80)
            p = cell.paragraphs[0]
            r = p.add_run(title)
            _font(r, size=8, bold=True, color=_C["muted"])
            p2 = cell.add_paragraph()
            r2 = p2.add_run(value)
            _font(r2, size=11, bold=True, color=_C["ink"])
        _nofit(chip_t)

    if executive.get("text"):
        doc.add_paragraph()
        note = doc.add_table(rows=1, cols=1)
        ncell = note.rows[0].cells[0]
        _shade(ncell, _C["blue_bg"])
        _set_cell_margins(ncell, 100, 100, 100, 100)
        p = ncell.paragraphs[0]
        r = p.add_run(str(executive["text"]))
        _font(r, size=10, color=_C["ink"])
        _nofit(note)

    # Progress indicators for key KPIs
    doc.add_paragraph()
    _add_progress_block(
        doc,
        [
            ("Качество знаний", float(executive.get("quality_rate") or 0), tones.get("quality", "mid")),
            ("Успеваемость", float(executive.get("pass_rate") or 0), tones.get("pass", "mid")),
            (
                avg_label,
                _avg_to_pct(float(executive.get("avg_score") or 0), avg_label),
                tones.get("avg", "mid"),
            ),
        ],
    )

    # ===== Distribution =====
    dist = ui.get("distribution") or []
    if dist:
        _module_title(doc, "02", "Распределение результатов")
        total_dist = ui.get("distribution_total") or sum(int(x.get("value") or 0) for x in dist)
        hint = doc.add_paragraph()
        rh = hint.add_run(f"Всего результатов в распределении: {total_dist}")
        _font(rh, size=9, color=_C["muted"])
        for idx, row in enumerate(dist, start=1):
            _hbar_row(
                doc,
                label=str(row.get("label") or ""),
                meta=f"{row.get('value')} · {row.get('percent')}%",
                pct=float(row.get("bar") or row.get("percent") or 0),
                color=_dist_color(idx),
            )
        for line in ui.get("distribution_analysis") or []:
            p = doc.add_paragraph()
            r = p.add_run(f"• {line}")
            _font(r, size=10, color=_C["ink"])

    # ===== Charts: avg + pass by subject =====
    cards = ui.get("subject_cards") or []
    chart_cards = [c for c in cards if not c.get("thin_sample")] or cards
    if chart_cards:
        _module_title(doc, "03", f"{avg_label} по предметам")
        max_avg = max((float(c.get("avg") or 0) for c in chart_cards), default=1) or 1
        for c in chart_cards:
            avg = float(c.get("avg") or 0)
            pct = min(100.0, round(100.0 * avg / max_avg, 1)) if max_avg else 0
            _hbar_row(
                doc,
                label=str(c.get("name") or ""),
                meta=str(c.get("avg") or ""),
                pct=pct,
                color=_tone_fg(c.get("tone") or "mid"),
            )

        _module_title(doc, "04", "Успеваемость по предметам")
        for c in chart_cards:
            _hbar_row(
                doc,
                label=str(c.get("name") or ""),
                meta=f"{c.get('pass_rate')}%",
                pct=float(c.get("pass_bar") or c.get("pass_rate") or 0),
                color=_tone_fg(c.get("tone") or "mid"),
            )

    # ===== Ranking =====
    _module_title(doc, "05", "Рейтинг предметов")
    rank = doc.add_table(rows=1, cols=2)
    left, right = rank.rows[0].cells
    _shade(left, _C["good_bg"])
    _shade(right, _C["low_bg"])
    _set_cell_margins(left, 100, 100, 100, 100)
    _set_cell_margins(right, 100, 100, 100, 100)
    _rank_panel(left, "ТОП лучших", ui.get("best_subjects") or [], avg_label, good=True)
    _rank_panel(right, "Проблемные", ui.get("worst_subjects") or [], avg_label, good=False)
    _nofit(rank)

    thin = ui.get("thin_subjects") or []
    if thin:
        p = doc.add_paragraph()
        r = p.add_run("Недостаточно статистики для объективной оценки")
        _font(r, size=10, bold=True, color=_C["muted"])
        for row in thin:
            p = doc.add_paragraph()
            r = p.add_run(f"• {row.get('name')}: {row.get('participants')} рез.")
            _font(r, size=9, color=_C["muted"])

    # ===== Subject cards grid =====
    if cards:
        _module_title(doc, "06", "Предметный анализ")
        # 2 columns
        for i in range(0, len(cards), 2):
            pair = cards[i : i + 2]
            t = doc.add_table(rows=1, cols=2 if len(pair) == 2 else 1)
            for j, card in enumerate(pair):
                _subject_card(t.rows[0].cells[j], card, avg_label=avg_label, count_label=ui.get("subject_count_label") or "Результатов")
            _nofit(t)
            doc.add_paragraph()

    # ===== Classes (compact, only if data) =====
    if class_rows:
        _module_title(doc, "07", "Рейтинг классов")
        for row in class_rows:
            name = row.get("name") or "Класс"
            avg = row.get("avg")
            pr = float(row.get("pass_rate") or 0)
            n = row.get("participants") or 0
            _hbar_row(
                doc,
                label=f"{name}  ·  n={n}  ·  {avg}",
                meta=f"{pr}%",
                pct=pr,
                color=_C["blue"],
            )

    # ===== Recommendations =====
    recs = ui.get("recommendations") or []
    if recs:
        _module_title(doc, "08", "Рекомендации")
        for line in recs:
            box = doc.add_table(rows=1, cols=1)
            cell = box.rows[0].cells[0]
            _shade(cell, _C["blue_bg"])
            _set_cell_margins(cell, 90, 90, 90, 90)
            p = cell.paragraphs[0]
            r = p.add_run(str(line))
            _font(r, size=10, color=_C["ink"])
            _nofit(box)
            doc.add_paragraph()

    # ===== Final panels =====
    panels = ui.get("panels") or ui.get("final_conclusion") or {}
    if panels:
        _module_title(doc, "09", "Общий вывод")
        blocks = [
            ("Сильные стороны", panels.get("strengths") or [], "high", "🟢"),
            ("Слабые стороны", panels.get("weaknesses") or [], "low", "🔴"),
            ("Основные риски", panels.get("risks") or [], "warn", "🟠"),
            ("Следующие шаги", panels.get("next_steps") or panels.get("priorities") or [], "neutral", "🔵"),
        ]
        for i in range(0, len(blocks), 2):
            pair = blocks[i : i + 2]
            t = doc.add_table(rows=1, cols=len(pair))
            for j, (title, items, tone, icon) in enumerate(pair):
                cell = t.rows[0].cells[j]
                _shade(cell, _tone_bg(tone))
                _set_cell_margins(cell, 100, 100, 100, 100)
                p = cell.paragraphs[0]
                r = p.add_run(f"{icon}  {title}")
                _font(r, size=11, bold=True, color=_tone_fg(tone))
                for item in items:
                    p2 = cell.add_paragraph()
                    r2 = p2.add_run(f"• {item}")
                    _font(r2, size=9, color=_C["ink"])
            _nofit(t)
            doc.add_paragraph()

        assessment = panels.get("assessment") or []
        if assessment:
            box = doc.add_table(rows=1, cols=1)
            cell = box.rows[0].cells[0]
            _shade(cell, _C["bg"])
            _set_cell_margins(cell, 100, 100, 100, 100)
            p = cell.paragraphs[0]
            r = p.add_run("Итоговая оценка")
            _font(r, size=11, bold=True, color=_C["blue"])
            for line in assessment:
                p2 = cell.add_paragraph()
                r2 = p2.add_run(str(line))
                _font(r2, size=10, color=_C["ink"])
            _nofit(box)

    foot = doc.add_paragraph()
    fr = foot.add_run(
        "Методика: "
        + ("ОГЭ — оценки 2–5." if exam_label == "ОГЭ" else "ЕГЭ — тестовые баллы 0–100.")
        + " Показатели сформированы по фактическим результатам школы."
    )
    _font(fr, size=8, color=_C["muted"])


# ---------- helpers ----------

def _setup_page(doc) -> None:
    from docx.shared import Cm

    section = doc.sections[0]
    section.top_margin = Cm(1.4)
    section.bottom_margin = Cm(1.4)
    section.left_margin = Cm(1.5)
    section.right_margin = Cm(1.5)


def _banner(doc, title: str, subtitle: str) -> None:
    t = doc.add_table(rows=1, cols=1)
    cell = t.rows[0].cells[0]
    _shade(cell, _C["title_bg"])
    _set_cell_margins(cell, 140, 140, 140, 140)
    p = cell.paragraphs[0]
    r = p.add_run(title)
    _font(r, size=14, bold=True, color="FFFFFF")
    p2 = cell.add_paragraph()
    r2 = p2.add_run(subtitle)
    _font(r2, size=9, color="BEE3F8")
    _nofit(t)
    doc.add_paragraph()


def _module_title(doc, num: str, title: str) -> None:
    doc.add_paragraph()
    t = doc.add_table(rows=1, cols=1)
    cell = t.rows[0].cells[0]
    _shade(cell, _C["bg"])
    _set_cell_margins(cell, 70, 70, 70, 70)
    p = cell.paragraphs[0]
    r = p.add_run(f"{num}   {title.upper()}")
    _font(r, size=11, bold=True, color=_C["title_bg"])
    _nofit(t)
    doc.add_paragraph()


def _fill_kpi_card(cell, *, icon: str, value: str, label: str, tone: str) -> None:
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    _shade(cell, _tone_bg(tone))
    _set_cell_margins(cell, 100, 90, 100, 90)
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(icon)
    _font(r, size=12, color=_tone_fg(tone))
    p2 = cell.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = p2.add_run(value)
    _font(r2, size=18, bold=True, color=_C["ink"])
    p3 = cell.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r3 = p3.add_run(label)
    _font(r3, size=8, color=_C["muted"])


def _add_progress_block(doc, items: list[tuple[str, float, str]]) -> None:
    for label, pct, tone in items:
        p = doc.add_paragraph()
        r = p.add_run(f"{label}: {pct:.1f}%   {_bar(pct)}")
        _font(r, size=9, color=_tone_fg(tone))


def _hbar_row(doc, *, label: str, meta: str, pct: float, color: str) -> None:
    t = doc.add_table(rows=1, cols=3)
    t.autofit = True
    c0, c1, c2 = t.rows[0].cells
    _set_cell_margins(c0, 40, 40, 40, 40)
    _set_cell_margins(c1, 40, 40, 40, 40)
    _set_cell_margins(c2, 40, 40, 40, 40)
    p0 = c0.paragraphs[0]
    r0 = p0.add_run(label)
    _font(r0, size=9, bold=True, color=_C["ink"])
    # visual bar in middle cell
    _shade(c1, "E2E8F0")
    filled = max(0, min(100, int(round(pct))))
    # approximate fill by nested text bar
    p1 = c1.paragraphs[0]
    r1 = p1.add_run(_bar(filled))
    _font(r1, size=9, bold=True, color=color)
    p2 = c2.paragraphs[0]
    r2 = p2.add_run(meta)
    _font(r2, size=9, color=_C["muted"])
    # widths hint
    try:
        t.columns[0].width = int(4.2 * 360000)  # EMU-ish via cm later
    except Exception:
        pass
    _nofit(t)


def _rank_panel(cell, title: str, rows: list, avg_label: str, *, good: bool) -> None:
    p = cell.paragraphs[0]
    r = p.add_run(title)
    _font(r, size=11, bold=True, color=_C["good"] if good else _C["low"])
    if not rows:
        p2 = cell.add_paragraph()
        r2 = p2.add_run("Недостаточно статистики для объективной оценки.")
        _font(r2, size=9, color=_C["muted"])
        return
    for idx, row in enumerate(rows, start=1):
        p2 = cell.add_paragraph()
        name = row.get("name") or ""
        avg = row.get("avg")
        pr = row.get("pass_rate")
        r2 = p2.add_run(f"{idx}. {name}")
        _font(r2, size=10, bold=True, color=_C["ink"])
        p3 = cell.add_paragraph()
        r3 = p3.add_run(f"{avg_label}: {avg}   ·   успеваемость {pr}%   {_bar(float(pr or 0))}")
        _font(r3, size=8, color=_C["muted"])


def _subject_card(cell, card: dict, *, avg_label: str, count_label: str) -> None:
    tone = card.get("tone") or "mid"
    _shade(cell, _tone_bg(tone))
    _set_cell_margins(cell, 110, 100, 110, 100)
    status = card.get("status") or "Средний"
    icon = {"high": "🟢", "mid": "🟡", "warn": "🟠", "low": "🔴"}.get(tone, "🟡")
    if card.get("thin_sample"):
        icon = "⚪"
        status = "Недостаточно данных"

    p = cell.paragraphs[0]
    r = p.add_run(str(card.get("name") or "Предмет"))
    _font(r, size=12, bold=True, color=_C["ink"])
    p2 = cell.add_paragraph()
    r2 = p2.add_run(f"{icon}  {status}")
    _font(r2, size=9, bold=True, color=_tone_fg(tone))

    p3 = cell.add_paragraph()
    r3 = p3.add_run(
        f"{count_label}: {card.get('participants')}    ·    {avg_label}: {card.get('avg')}    ·    "
        f"Успеваемость: {card.get('pass_rate')}%"
    )
    _font(r3, size=9, color=_C["muted"])

    p4 = cell.add_paragraph()
    r4 = p4.add_run(_bar(float(card.get("pass_bar") or card.get("pass_rate") or 0)))
    _font(r4, size=9, color=_tone_fg(tone))

    if card.get("conclusion"):
        p5 = cell.add_paragraph()
        r5 = p5.add_run(str(card["conclusion"]))
        _font(r5, size=9, color=_C["ink"])
    if card.get("recommendation") and not card.get("thin_sample"):
        p6 = cell.add_paragraph()
        r6 = p6.add_run(str(card["recommendation"]))
        _font(r6, size=9, color=_C["blue"])


def _bar(pct: float, width: int = 12) -> str:
    pct = max(0.0, min(100.0, float(pct or 0)))
    filled = int(round(width * pct / 100.0))
    return "█" * filled + "░" * (width - filled)


def _avg_to_pct(avg: float, avg_label: str) -> float:
    if "оценка" in (avg_label or "").lower() or avg <= 5:
        return min(100.0, round((avg / 5.0) * 100.0, 1))
    return min(100.0, float(avg))


def _tone_bg(tone: str) -> str:
    return {
        "high": _C["good_bg"],
        "success": _C["good_bg"],
        "mid": _C["mid_bg"],
        "warn": _C["warn_bg"],
        "low": _C["low_bg"],
        "danger": _C["low_bg"],
        "neutral": _C["blue_bg"],
    }.get(tone or "neutral", _C["blue_bg"])


def _tone_fg(tone: str) -> str:
    return {
        "high": _C["good"],
        "success": _C["good"],
        "mid": _C["mid"],
        "warn": _C["warn"],
        "low": _C["low"],
        "danger": _C["low"],
        "neutral": _C["blue"],
    }.get(tone or "neutral", _C["blue"])


def _dist_color(idx: int) -> str:
    return {1: _C["low"], 2: _C["warn"], 3: _C["blue"], 4: _C["good"]}.get(idx, _C["blue"])


def _shade(cell, fill_hex: str) -> None:
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill_hex)
    shd.set(qn("w:val"), "clear")
    tc_pr.append(shd)


def _set_cell_margins(cell, top: int, bottom: int, left: int, right: int) -> None:
    """Margins in twips-ish via tcMar (DXA)."""
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = OxmlElement("w:tcMar")
    for name, val in (("top", top), ("bottom", bottom), ("left", left), ("right", right)):
        node = OxmlElement(f"w:{name}")
        node.set(qn("w:w"), str(val))
        node.set(qn("w:type"), "dxa")
        tc_mar.append(node)
    tc_pr.append(tc_mar)


def _font(run, *, size: int = 10, bold: bool = False, color: str = "102A43") -> None:
    from docx.shared import Pt, RGBColor

    run.bold = bold
    run.font.size = Pt(size)
    run.font.name = "Calibri"
    try:
        run._element.rPr.rFonts.set(
            "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}ascii",
            "Calibri",
        )
    except Exception:
        pass
    color = (color or "102A43").lstrip("#")
    run.font.color.rgb = RGBColor(int(color[0:2], 16), int(color[2:4], 16), int(color[4:6], 16))


def _nofit(table) -> None:
    try:
        table.autofit = True
    except Exception:
        pass
    # remove heavy grid look where possible
    for style_name in ("Table Grid", "Normal Table"):
        try:
            table.style = style_name
            break
        except (KeyError, ValueError):
            continue
