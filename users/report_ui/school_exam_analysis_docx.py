"""
Оформление DOCX предметной аналитической справки по экзамену.

Только Presentation Layer: графики, диаграммы, визуальный каркас.
Тексты и расчёты берутся из готового payload / ExamData без изменений.
"""

from __future__ import annotations

import struct
import zlib
from io import BytesIO
from math import atan2, pi
from typing import Iterable


_C = {
    "ink": "102A43",
    "muted": "627D98",
    "bg": "F0F4F8",
    "blue": "2B6CB0",
    "blue_bg": "EBF8FF",
    "navy": "0B1F3A",
    "good": "2F855A",
    "good_bg": "F0FFF4",
    "mid": "B7791F",
    "mid_bg": "FFFBEB",
    "warn": "C05621",
    "warn_bg": "FFFAF0",
    "low": "C53030",
    "low_bg": "FFF5F5",
    "title": "1A365D",
    "white": "FFFFFF",
    "line": "D9E2EC",
}


def render_exam_analysis_docx(data) -> BytesIO:
    from docx import Document
    from docx.shared import Cm

    from users.export_reports import _build_analysis_payload, _exam_type_label_ru

    doc = Document()
    section = doc.sections[0]
    section.top_margin = Cm(1.3)
    section.bottom_margin = Cm(1.3)
    section.left_margin = Cm(1.4)
    section.right_margin = Cm(1.4)

    exam_label = _exam_type_label_ru(data.exam_type)
    payload = _build_analysis_payload(data)
    sections = payload.get("sections") or {}

    _banner(
        doc,
        f"АНАЛИТИЧЕСКАЯ СПРАВКА  ·  {exam_label}",
        f"{data.subject}  ·  дата экзамена {data.date}",
    )

    kpi_items = [
        ("👥", str(data.students_count), "Участников", "neutral"),
        ("📈", str(data.avg_score), "Средний балл", _tone_avg(data.avg_score, data)),
        ("↓", str(data.min_score), "Минимум", "neutral"),
        ("↑", str(data.max_score), "Максимум", "neutral"),
        ("✅", f"{data.pass_rate}%", "Доля сдавших", _tone_pct(data.pass_rate)),
        ("★", str(len(data.strong_tasks or [])), "Сильные задания", "high"),
        ("⚠", str(len(data.weak_tasks or [])), "Задания <50%", "warn" if data.weak_tasks else "neutral"),
        ("▣", str(len(data.tasks or [])), "Заданий КИМ", "neutral"),
    ]
    kpi_t = doc.add_table(rows=2, cols=4)
    for idx, (icon, value, label, tone) in enumerate(kpi_items):
        r, c = divmod(idx, 4)
        _kpi_cell(kpi_t.rows[r].cells[c], icon, value, label, tone)
    _nofit(kpi_t)
    _set_borders(kpi_t, "FFFFFF")
    _gap(doc)

    _insert_charts(doc, data)

    for title, lines in sections.items():
        _module(doc, str(title))
        for line in lines or []:
            _body(doc, str(line))

    _module(doc, "Сводные показатели")
    table = doc.add_table(rows=2, cols=5)
    headers = ["Участников", "Средний", "Мин", "Макс", "Доля сдавших"]
    vals = [
        str(data.students_count),
        str(data.avg_score),
        str(data.min_score),
        str(data.max_score),
        f"{data.pass_rate}%",
    ]
    for i, h in enumerate(headers):
        _cell_text(table.rows[0].cells[i], h, bold=True, size=8, color=_C["muted"], fill=_C["bg"])
    for i, v in enumerate(vals):
        _cell_text(table.rows[1].cells[i], v, bold=True, size=12, color=_C["ink"])
    _nofit(table)
    _set_borders(table, _C["line"])
    _gap(doc)

    if data.tasks:
        _module(doc, "Успешность по заданиям")
        ttasks = doc.add_table(rows=1, cols=5)
        for i, h in enumerate(["Задание", "Успешность, %", "Верно (+)", "Ошибок (−)", "Всего ответов"]):
            _cell_text(ttasks.rows[0].cells[i], h, bold=True, size=8, color=_C["white"], fill=_C["navy"])
        for task in data.tasks:
            cells = ttasks.add_row().cells
            rate = float(task.get("success_rate") or 0)
            _cell_text(cells[0], f"№{task['id']}", bold=True, size=9, color=_C["ink"])
            _cell_text(
                cells[1],
                str(task.get("success_rate", "")),
                bold=True,
                size=9,
                color=_rate_fg(rate),
                fill=_rate_bg(rate),
            )
            _cell_text(cells[2], str(task.get("correct", "")), size=9, color=_C["ink"])
            _cell_text(cells[3], str(task.get("wrong", "")), size=9, color=_C["ink"])
            _cell_text(cells[4], str(task.get("total", "")), size=9, color=_C["ink"])
        _nofit(ttasks)
        _set_borders(ttasks, _C["line"])
        _gap(doc)

    recs = payload.get("recommendations") or {}
    _module(doc, "11. Рекомендации")
    for section_name, lines in recs.items():
        _sub(doc, str(section_name))
        for line in lines or []:
            _bullet(doc, str(line))

    _module(doc, "12. План контроля заданий ниже 50%")
    control_rows = payload.get("control_plan") or []
    if control_rows:
        ctable = doc.add_table(rows=1, cols=5)
        for i, h in enumerate(["Задание", "Тематический блок", "Уровень дефицита", "Классы", "Действия"]):
            _cell_text(ctable.rows[0].cells[i], h, bold=True, size=8, color=_C["white"], fill=_C["navy"])
        for item in control_rows:
            cells = ctable.add_row().cells
            _cell_text(cells[0], str(item.get("task") or ""), bold=True, size=9, color=_C["ink"])
            _cell_text(cells[1], str(item.get("block") or ""), size=8, color=_C["ink"])
            _cell_text(
                cells[2],
                str(item.get("severity") or "значимый"),
                size=8,
                color=_C["warn"],
                fill=_C["warn_bg"],
            )
            _cell_text(cells[3], str(item.get("classes") or ""), size=8, color=_C["ink"])
            _cell_text(cells[4], str(item.get("action") or ""), size=8, color=_C["ink"])
        _nofit(ctable)
        _set_borders(ctable, _C["line"])
    else:
        _body(doc, "Заданий с успешностью ниже 50% не выявлено.")

    out = BytesIO()
    doc.save(out)
    out.seek(0)
    return out


def _insert_charts(doc, data) -> None:
    from users.export_reports import _score_groups, _subject_key

    _module(doc, "Диаграммы")

    tasks = list(data.tasks or [])
    if tasks:
        _sub(doc, "Успешность выполнения заданий КИМ, %")
        labels = [str(t.get("id")) for t in tasks]
        values = [float(t.get("success_rate") or 0) for t in tasks]
        colors = [_rate_hex(v) for v in values]
        png = _bar_chart_png(labels, values, y_max=100, colors=colors, y_label="%")
        _picture(doc, png, width_cm=16.4)
        _legend(
            doc,
            [
                ("< 30%", _C["low"]),
                ("30–49%", _C["warn"]),
                ("50–79%", _C["mid"]),
                ("≥ 80%", _C["good"]),
            ],
        )

    subject_key = _subject_key(data.subject, data.exam_type)
    groups = _score_groups(list(data.score_values or []), subject_key, data.exam_type)
    dist_items = _distribution_items(groups)
    if dist_items and any(item[1] for item in dist_items):
        _sub(doc, "Распределение обучающихся по уровням")
        png = _pie_chart_png([(label, count, color) for label, count, color in dist_items])
        _picture(doc, png, width_cm=16.4)
        _legend(doc, [(f"{label} · {count}", color) for label, count, color in dist_items])
        _stacked_bar(doc, dist_items)
        total = max(1, int(data.students_count or 0))
        for label, count, color in dist_items:
            pct = round(100.0 * count / total, 1) if total else 0.0
            _width_bar(doc, label=label, meta=f"{count} чел.  ·  {pct}%", pct=pct, color=color)

    dynamics = list(data.dynamics or [])
    if dynamics:
        _sub(doc, "Динамика среднего балла по годам")
        years = [str(row.get("year") or "") for row in dynamics]
        avgs = [float(row.get("avg_score") or 0) for row in dynamics]
        y_max = 5 if avgs and max(avgs) <= 5 else 100
        png = _line_chart_png(years, avgs, y_max=y_max, y_label="avg")
        _picture(doc, png, width_cm=16.4)
        for year, avg in zip(years, avgs):
            pct = (avg / y_max) * 100 if y_max else 0
            _width_bar(doc, label=str(year), meta=str(avg), pct=pct, color=_C["blue"])


def _distribution_items(groups: dict[str, int]) -> list[tuple[str, int, str]]:
    if "grade_2" in groups:
        return [
            ("Отметка 2", int(groups.get("grade_2") or 0), _C["low"]),
            ("Отметка 3", int(groups.get("grade_3") or 0), _C["warn"]),
            ("Отметка 4", int(groups.get("grade_4") or 0), _C["mid"]),
            ("Отметка 5", int(groups.get("grade_5") or 0), _C["good"]),
        ]
    return [
        ("Слабый", int(groups.get("weak") or 0), _C["low"]),
        ("Базовый", int(groups.get("basic") or 0), _C["warn"]),
        ("Хороший", int(groups.get("good") or 0), _C["mid"]),
        ("Высокий", int(groups.get("high") or 0), _C["good"]),
    ]


def _stacked_bar(doc, items: list[tuple[str, int, str]]) -> None:
    total = sum(max(0, int(x[1])) for x in items) or 1
    cols = [x for x in items if int(x[1]) > 0] or items
    t = doc.add_table(rows=1, cols=len(cols))
    widths = [16.4 * (max(0, int(count)) / total) for _, count, _ in cols]
    _set_tbl_widths(t, widths)
    for i, (label, count, color) in enumerate(cols):
        cell = t.rows[0].cells[i]
        pct = 100.0 * count / total
        _cell_text(
            cell,
            f"{label}\n{count} ({pct:.0f}%)",
            bold=True,
            size=8,
            color=_C["white"],
            fill=color,
            center=True,
        )
    _nofit(t)
    _set_borders(t, color)
    _gap(doc, 80)


def _legend(doc, items: list[tuple[str, str]]) -> None:
    t = doc.add_table(rows=1, cols=len(items))
    for i, (label, color) in enumerate(items):
        _cell_text(t.rows[0].cells[i], f"■  {label}", bold=True, size=8, color=color, fill=_C["bg"], center=True)
    _nofit(t)
    _set_borders(t, _C["white"])
    _gap(doc, 60)


def _picture(doc, buf: BytesIO | None, *, width_cm: float) -> None:
    from docx.shared import Cm, Pt

    if buf is None:
        return
    buf.seek(0)
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run()
    run.add_picture(buf, width=Cm(width_cm))


def _bar_chart_png(labels: list[str], values: list[float], *, y_max: float, colors: list[str], y_label: str) -> BytesIO:
    w, h = 1100, 430
    c = _Canvas(w, h)
    pad_l, pad_r, pad_t, pad_b = 62, 18, 22, 58
    x0, y0, x1, y1 = pad_l, pad_t, w - pad_r, h - pad_b
    plot_w, plot_h = x1 - x0, y1 - y0
    y_max = max(float(y_max or 0), 1.0)
    c.fill_rect(x0, y0, x1, y1, (255, 255, 255))
    c.rect(x0, y0, x1, y1, (16, 42, 67))
    c.text(8, 8, y_label, (16, 42, 67), scale=2)
    _draw_grid(c, x0, y0, x1, y1, y_max)
    ref = y1 - int(50 / y_max * plot_h) if y_max >= 50 else None
    if ref is not None:
        c.hline(x0, ref, x1, (197, 48, 48))
    n = max(1, len(values))
    gap = 6 if n <= 24 else 3
    bar_w = max(8, int((plot_w - gap * (n + 1)) / n))
    for i, (lab, v) in enumerate(zip(labels, values)):
        bx = x0 + gap + i * (bar_w + gap)
        bh = int(max(0.0, min(float(v), y_max)) / y_max * plot_h)
        by = y1 - bh
        c.fill_rect(bx, by, bx + bar_w, y1, _hex_rgb(colors[i % len(colors)] if colors else _C["blue"]))
        if n <= 28 or i % 2 == 0:
            c.text(bx + max(0, (bar_w - 6) // 2), y1 + 8, str(lab)[:4], (16, 42, 67), scale=1)
    return c.png()


def _line_chart_png(labels: list[str], values: list[float], *, y_max: float, y_label: str) -> BytesIO:
    w, h = 1100, 400
    c = _Canvas(w, h)
    pad_l, pad_r, pad_t, pad_b = 62, 24, 22, 52
    x0, y0, x1, y1 = pad_l, pad_t, w - pad_r, h - pad_b
    plot_w, plot_h = x1 - x0, y1 - y0
    y_max = max(float(y_max or 0), 1.0)
    c.rect(x0, y0, x1, y1, (16, 42, 67))
    c.text(8, 8, y_label, (16, 42, 67), scale=2)
    _draw_grid(c, x0, y0, x1, y1, y_max)
    n = max(1, len(values))
    pts = []
    for i, v in enumerate(values):
        x = x0 + int((i + 0.5) * plot_w / n)
        y = y1 - int(max(0.0, min(float(v), y_max)) / y_max * plot_h)
        pts.append((x, y))
    for a, b in zip(pts, pts[1:]):
        c.line(a[0], a[1], b[0], b[1], (43, 108, 176), width=3)
    for (x, y), lab, val in zip(pts, labels, values):
        c.fill_circle(x, y, 6, (43, 108, 176))
        c.text(x - 12, y1 + 10, str(lab), (16, 42, 67), scale=2)
        c.text(x - 10, y - 18, _fmt_axis(val), (16, 42, 67), scale=2)
    return c.png()


def _pie_chart_png(items: list[tuple[str, int, str]]) -> BytesIO:
    w, h = 980, 400
    c = _Canvas(w, h)
    total = sum(max(0, int(v)) for _, v, _ in items) or 1
    cx, cy, r, inner = 490, 200, 155, 82
    start = -pi / 2
    nonzero = [(label, count, color) for label, count, color in items if int(count) > 0]
    if len(nonzero) == 1:
        c.fill_circle(cx, cy, r, _hex_rgb(nonzero[0][2]))
    else:
        for _, count, color in items:
            share = max(0, int(count)) / total
            if share <= 0:
                continue
            sweep = share * 2 * pi
            c.pie_slice(cx, cy, r, start, start + sweep, _hex_rgb(color))
            start += sweep
    c.fill_circle(cx, cy, inner, (255, 255, 255))
    total_s = str(total)
    c.text(cx - 6 * len(total_s), cy - 10, total_s, (16, 42, 67), scale=3)
    return c.png()


class _Canvas:
    def __init__(self, w: int, h: int, bg: tuple[int, int, int] = (255, 255, 255)):
        self.w = w
        self.h = h
        self.px = bytearray(bytes(bg) * (w * h))

    def _put(self, x: int, y: int, rgb: tuple[int, int, int]) -> None:
        if 0 <= x < self.w and 0 <= y < self.h:
            i = (y * self.w + x) * 3
            self.px[i] = rgb[0]
            self.px[i + 1] = rgb[1]
            self.px[i + 2] = rgb[2]

    def fill_rect(self, x0: int, y0: int, x1: int, y1: int, rgb: tuple[int, int, int]) -> None:
        x0, x1 = max(0, int(x0)), min(self.w, int(x1))
        y0, y1 = max(0, int(y0)), min(self.h, int(y1))
        if x1 <= x0 or y1 <= y0:
            return
        chunk = bytes(rgb) * (x1 - x0)
        for y in range(y0, y1):
            start = (y * self.w + x0) * 3
            self.px[start : start + len(chunk)] = chunk

    def rect(self, x0: int, y0: int, x1: int, y1: int, rgb: tuple[int, int, int]) -> None:
        self.hline(x0, y0, x1, rgb)
        self.hline(x0, y1, x1, rgb)
        self.vline(x0, y0, y1, rgb)
        self.vline(x1, y0, y1, rgb)

    def hline(self, x0: int, y: int, x1: int, rgb: tuple[int, int, int]) -> None:
        self.fill_rect(x0, y, x1, y + 1, rgb)

    def vline(self, x: int, y0: int, y1: int, rgb: tuple[int, int, int]) -> None:
        self.fill_rect(x, y0, x + 1, y1, rgb)

    def line(self, x0: int, y0: int, x1: int, y1: int, rgb: tuple[int, int, int], width: int = 1) -> None:
        dx = abs(x1 - x0)
        dy = abs(y1 - y0)
        sx = 1 if x0 < x1 else -1
        sy = 1 if y0 < y1 else -1
        err = dx - dy
        x, y = x0, y0
        while True:
            if width <= 1:
                self._put(x, y, rgb)
            else:
                r = max(1, width // 2)
                self.fill_circle(x, y, r, rgb)
            if x == x1 and y == y1:
                break
            e2 = 2 * err
            if e2 > -dy:
                err -= dy
                x += sx
            if e2 < dx:
                err += dx
                y += sy

    def fill_circle(self, cx: int, cy: int, r: int, rgb: tuple[int, int, int]) -> None:
        rr = r * r
        for y in range(cy - r, cy + r + 1):
            for x in range(cx - r, cx + r + 1):
                if (x - cx) * (x - cx) + (y - cy) * (y - cy) <= rr:
                    self._put(x, y, rgb)

    def pie_slice(self, cx: int, cy: int, r: int, a0: float, a1: float, rgb: tuple[int, int, int]) -> None:
        if a1 - a0 <= 0:
            return
        two_pi = 2 * pi

        def _norm(a: float) -> float:
            while a < 0:
                a += two_pi
            while a >= two_pi:
                a -= two_pi
            return a

        a0n, a1n = _norm(a0), _norm(a1)
        wrap = a1n < a0n
        rr = r * r
        for y in range(cy - r, cy + r + 1):
            for x in range(cx - r, cx + r + 1):
                dx, dy = x - cx, y - cy
                if dx * dx + dy * dy > rr:
                    continue
                ang = _norm(atan2(dy, dx))
                inside = (ang >= a0n or ang <= a1n) if wrap else (a0n <= ang <= a1n)
                if inside:
                    self._put(x, y, rgb)

    def text(self, x: int, y: int, s: str, rgb: tuple[int, int, int], scale: int = 2) -> None:
        cx = x
        for ch in str(s):
            glyph = _GLYPHS.get(ch.upper() if ch.isalpha() else ch)
            if glyph is None:
                cx += 4 * scale
                continue
            for gy, row in enumerate(glyph):
                for gx, bit in enumerate(row):
                    if bit == "1":
                        for sy in range(scale):
                            for sx in range(scale):
                                self._put(cx + gx * scale + sx, y + gy * scale + sy, rgb)
            cx += (len(glyph[0]) + 1) * scale

    def png(self) -> BytesIO:
        raw = b"".join(
            b"\x00" + bytes(self.px[y * self.w * 3 : (y + 1) * self.w * 3]) for y in range(self.h)
        )
        ihdr = struct.pack(">IIBBBBB", self.w, self.h, 8, 2, 0, 0, 0)

        def chunk(tag: bytes, data: bytes) -> bytes:
            return struct.pack(">I", len(data)) + tag + data + struct.pack(">I", zlib.crc32(tag + data) & 0xFFFFFFFF)

        buf = BytesIO()
        buf.write(b"\x89PNG\r\n\x1a\n")
        buf.write(chunk(b"IHDR", ihdr))
        buf.write(chunk(b"IDAT", zlib.compress(raw, 9)))
        buf.write(chunk(b"IEND", b""))
        buf.seek(0)
        return buf


_GLYPHS = {
    "0": ["01110", "10001", "10001", "10001", "10001", "10001", "01110"],
    "1": ["00100", "01100", "00100", "00100", "00100", "00100", "01110"],
    "2": ["01110", "10001", "00001", "00010", "00100", "01000", "11111"],
    "3": ["11110", "00001", "00001", "01110", "00001", "00001", "11110"],
    "4": ["00010", "00110", "01010", "10010", "11111", "00010", "00010"],
    "5": ["11111", "10000", "11110", "00001", "00001", "10001", "01110"],
    "6": ["01110", "10000", "11110", "10001", "10001", "10001", "01110"],
    "7": ["11111", "00001", "00010", "00100", "01000", "01000", "01000"],
    "8": ["01110", "10001", "10001", "01110", "10001", "10001", "01110"],
    "9": ["01110", "10001", "10001", "01111", "00001", "00001", "01110"],
    "%": ["11001", "11010", "00100", "01000", "10110", "10011", "00000"],
    ".": ["00000", "00000", "00000", "00000", "00000", "01100", "01100"],
    "-": ["00000", "00000", "00000", "11111", "00000", "00000", "00000"],
    " ": ["000", "000", "000", "000", "000", "000", "000"],
    "(": ["0010", "0100", "1000", "1000", "1000", "0100", "0010"],
    ")": ["0100", "0010", "0001", "0001", "0001", "0010", "0100"],
    ":": ["000", "011", "011", "000", "011", "011", "000"],
    "A": ["01110", "10001", "10001", "11111", "10001", "10001", "10001"],
    "B": ["11110", "10001", "10001", "11110", "10001", "10001", "11110"],
    "C": ["01110", "10001", "10000", "10000", "10000", "10001", "01110"],
    "E": ["11111", "10000", "10000", "11110", "10000", "10000", "11111"],
    "G": ["01110", "10001", "10000", "10111", "10001", "10001", "01110"],
    "H": ["10001", "10001", "10001", "11111", "10001", "10001", "10001"],
    "I": ["111", "010", "010", "010", "010", "010", "111"],
    "K": ["10001", "10010", "10100", "11000", "10100", "10010", "10001"],
    "L": ["10000", "10000", "10000", "10000", "10000", "10000", "11111"],
    "N": ["10001", "11001", "10101", "10011", "10001", "10001", "10001"],
    "O": ["01110", "10001", "10001", "10001", "10001", "10001", "01110"],
    "R": ["11110", "10001", "10001", "11110", "10100", "10010", "10001"],
    "S": ["01111", "10000", "10000", "01110", "00001", "00001", "11110"],
    "V": ["10001", "10001", "10001", "10001", "10001", "01010", "00100"],
    "Y": ["10001", "10001", "01010", "00100", "00100", "00100", "00100"],
    "Z": ["11111", "00001", "00010", "00100", "01000", "10000", "11111"],
}


def _banner(doc, title: str, subtitle: str) -> None:
    t = doc.add_table(rows=1, cols=1)
    cell = t.rows[0].cells[0]
    _shade(cell, _C["blue_bg"])
    _pad(cell, 160)
    p = cell.paragraphs[0]
    r = p.add_run(title)
    _font_run(r, 15, True, _C["title"])
    p2 = cell.add_paragraph()
    r2 = p2.add_run(subtitle)
    _font_run(r2, 10, False, _C["muted"])
    _nofit(t)
    _set_borders(t, "BEE3F8")
    _gap(doc)


def _module(doc, title: str) -> None:
    from docx.shared import Pt

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run(title)
    _font_run(r, 13, True, _C["title"])


def _sub(doc, text: str) -> None:
    from docx.shared import Pt

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(text)
    _font_run(r, 10, True, _C["blue"])


def _body(doc, text: str) -> None:
    from docx.shared import Pt

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run(text)
    _font_run(r, 10, False, _C["ink"])


def _bullet(doc, text: str) -> None:
    from docx.shared import Pt, Twips

    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Twips(180)
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(f"•  {text}")
    _font_run(r, 10, False, _C["ink"])


def _kpi_cell(cell, icon: str, value: str, label: str, tone: str) -> None:
    _shade(cell, _tone_bg(tone))
    _pad(cell)
    p = cell.paragraphs[0]
    r = p.add_run(f"{icon}  {label}")
    _font_run(r, 8, True, _C["muted"])
    p2 = cell.add_paragraph()
    r2 = p2.add_run(value)
    _font_run(r2, 16, True, _C["ink"])


def _width_bar(doc, *, label: str, meta: str, pct: float, color: str) -> None:
    fill = max(0.0, min(100.0, float(pct or 0)))
    empty = 100.0 - fill
    t = doc.add_table(rows=1, cols=3)
    _set_tbl_widths(t, [3.6, 12.8 * fill / 100.0 if fill else 0.15, 12.8 * empty / 100.0 if empty else 0.15])
    left, bar, rest = t.rows[0].cells
    _shade(left, _C["bg"])
    _pad(left, 60)
    p = left.paragraphs[0]
    r = p.add_run(label)
    _font_run(r, 8, True, _C["ink"])
    p2 = left.add_paragraph()
    r2 = p2.add_run(meta)
    _font_run(r2, 8, False, _C["muted"])
    _shade(bar, color)
    _shade(rest, _C["bg"])
    _nofit(t)
    _set_borders(t, _C["white"])


def _cell_text(
    cell,
    text: str,
    *,
    bold: bool = False,
    size: int = 9,
    color: str = "102A43",
    fill: str | None = None,
    center: bool = False,
) -> None:
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    if fill:
        _shade(cell, fill)
    p = cell.paragraphs[0]
    if center:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.clear()
    r = p.add_run(text)
    _font_run(r, size, bold, color)


def _gap(doc, twips: int = 80) -> None:
    from docx.shared import Pt

    p = doc.add_paragraph("")
    p.paragraph_format.space_after = Pt(max(2, twips // 20))


def _tone_avg(avg: float, data) -> str:
    scores = list(data.score_values or [])
    if data.exam_type == "oge" or (scores and max(scores) <= 5):
        if avg >= 4.2:
            return "high"
        if avg >= 3.5:
            return "mid"
        return "low"
    if avg >= 60:
        return "high"
    if avg >= 45:
        return "mid"
    return "low"


def _tone_pct(value: float) -> str:
    if value >= 85:
        return "high"
    if value >= 70:
        return "mid"
    return "low"


def _tone_bg(tone: str) -> str:
    return {
        "high": _C["good_bg"],
        "mid": _C["mid_bg"],
        "warn": _C["warn_bg"],
        "low": _C["low_bg"],
        "neutral": _C["blue_bg"],
    }.get(tone, _C["bg"])


def _rate_bg(pct: float) -> str:
    if pct < 30:
        return _C["low_bg"]
    if pct < 50:
        return _C["warn_bg"]
    if pct < 80:
        return _C["mid_bg"]
    return _C["good_bg"]


def _rate_fg(pct: float) -> str:
    if pct < 30:
        return _C["low"]
    if pct < 50:
        return _C["warn"]
    if pct < 80:
        return _C["mid"]
    return _C["good"]


def _rate_hex(pct: float) -> str:
    if pct < 30:
        return _C["low"]
    if pct < 50:
        return _C["warn"]
    if pct < 80:
        return _C["mid"]
    return _C["good"]


def _draw_grid(c: _Canvas, x0: int, y0: int, x1: int, y1: int, y_max: float) -> None:
    plot_h = y1 - y0
    if abs(y_max - 5) < 1e-6:
        ticks = [0, 1, 2, 3, 4, 5]
    elif abs(y_max - 100) < 1e-6:
        ticks = [0, 25, 50, 75, 100]
    else:
        ticks = [y_max * i / 4 for i in range(5)]
    for val in ticks:
        y = y1 - int(float(val) / y_max * plot_h)
        c.hline(x0, y, x1, (230, 236, 242))
        c.text(8, y - 6, _fmt_axis(val), (93, 109, 126), scale=2)


def _fmt_axis(v: float) -> str:
    if abs(v - round(v)) < 1e-6:
        return str(int(round(v)))
    return f"{v:.1f}"


def _hex_rgb(hex_color: str) -> tuple[int, int, int]:
    h = (hex_color or "2B6CB0").lstrip("#")
    return int(h[0:2], 16), int(h[2:4], 16), int(h[4:6], 16)


def _font_run(run, size_pt: int, bold: bool, color_hex: str) -> None:
    from docx.shared import Pt, RGBColor

    run.bold = bold
    run.font.size = Pt(size_pt)
    run.font.color.rgb = RGBColor.from_string(color_hex)
    run.font.name = "Calibri"


def _shade(cell, hex_color: str) -> None:
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), hex_color)
    shd.set(qn("w:val"), "clear")
    tcPr.append(shd)


def _pad(cell, twips: int = 100) -> None:
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement("w:tcMar")
    for edge in ("top", "left", "bottom", "right"):
        node = OxmlElement(f"w:{edge}")
        node.set(qn("w:w"), str(twips))
        node.set(qn("w:type"), "dxa")
        tcMar.append(node)
    tcPr.append(tcMar)


def _set_tbl_widths(table, widths_cm: Iterable[float]) -> None:
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.shared import Cm

    widths = [Cm(max(0.15, float(w))) for w in widths_cm]
    tbl = table._tbl
    tblGrid = tbl.find(qn("w:tblGrid"))
    if tblGrid is not None:
        for child in list(tblGrid):
            tblGrid.remove(child)
    else:
        tblGrid = OxmlElement("w:tblGrid")
        tbl.insert(0, tblGrid)
    for width in widths:
        grid_col = OxmlElement("w:gridCol")
        grid_col.set(qn("w:w"), str(width.twips))
        tblGrid.append(grid_col)
    for row in table.rows:
        for cell, width in zip(row.cells, widths):
            tcPr = cell._tc.get_or_add_tcPr()
            tcW = tcPr.find(qn("w:tcW"))
            if tcW is None:
                tcW = OxmlElement("w:tcW")
                tcPr.append(tcW)
            tcW.set(qn("w:w"), str(width.twips))
            tcW.set(qn("w:type"), "dxa")
            cell.width = width


def _set_borders(table, color_hex: str) -> None:
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    tbl = table._tbl
    tblPr = tbl.tblPr
    if tblPr is None:
        tblPr = OxmlElement("w:tblPr")
        tbl.insert(0, tblPr)
    borders = tblPr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tblPr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = borders.find(qn(f"w:{edge}"))
        if el is None:
            el = OxmlElement(f"w:{edge}")
            borders.append(el)
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), "4")
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), color_hex)


def _nofit(table) -> None:
    table.autofit = True
