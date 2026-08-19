"""
Финальное оформление DOCX «Свод результатов ГИА по району».

Только визуальный слой: тексты, показатели и данные таблиц — из payload без изменений.
"""

from __future__ import annotations

from io import BytesIO
from typing import Any

NAVY = "0B1F3A"
BLUE = "1B4F72"
SKY = "5B9BD5"
GRAY = "5D6D7E"
LINE = "D5DDE5"
SOFT = "F4F7FA"
WHITE = "FFFFFF"
GOOD = "1E7A4A"
GOOD_BG = "EAF6EF"
WARN = "A85A1A"
WARN_BG = "F8F1E6"
RISK = "9B2C2C"
RISK_BG = "F8ECEC"
INFO_BG = "EAF2F8"
DOC_VERSION = "1.0"


def render_district_gia_summary_docx(data: dict[str, Any]) -> BytesIO:
    from docx import Document
    from docx.shared import Cm, Pt

    doc = Document()
    _setup_page(doc)

    if not (data or {}).get("has_data"):
        doc.add_paragraph((data or {}).get("message") or "Недостаточно данных.")
        return _save(doc)

    district = data.get("district_name") or "Муниципалитет"
    et = (data.get("exam_type") or "ege").lower()
    exam = "ОГЭ" if et == "oge" else "ЕГЭ"
    year = data.get("year") or ""
    formed = data.get("generated_at") or ""
    avg_label = "Средняя оценка" if et == "oge" else "Средний балл"

    _setup_header_footer(doc, district=str(district), year=year, formed=str(formed))
    _title_page(doc, district=str(district), exam=exam, year=year, formed=str(formed))
    _toc(doc)

    # Dashboard KPI (visual only — values from payload)
    _section_block(doc, "Общие показатели", "◆", accent=SKY)
    _kpi_dashboard(
        doc,
        [
            ("👤", "Участники", str(data.get("participants")), _tone_neutral()),
            ("Σ", avg_label, str(data.get("avg_score")), _tone_avg(data.get("avg_score"), et)),
            ("Q", "Качество знаний", f"{data.get('quality_rate')}%", _tone_pct(data.get("quality_rate"), 40, 25)),
            ("%", "Успеваемость", f"{data.get('pass_rate')}%", _tone_pct(data.get("pass_rate"), 85, 70)),
            ("★", "Высокобалльники", str(data.get("high_count")), _tone_neutral()),
            ("!", "Неудовлетворительные", str(data.get("failed_count")), _tone_fail(data.get("failed_count"))),
            ("↗", "Динамика", _fmt_delta(data.get("avg_delta")), _tone_delta(data.get("avg_delta"))),
            ("⌂", "Количество школ", str(data.get("schools_count")), _tone_neutral()),
        ],
    )
    doc.add_paragraph("")

    sections = data.get("report_sections") or []
    if sections:
        for section in sections:
            title = str(section.get("title") or "Раздел")
            items = list(section.get("items") or [])
            # Skip duplicate KPI text block title if identical metrics already on dashboard
            # — still render items; only avoid renaming. Always render section.
            kind = _detect_kind(title)
            icon = _kind_icon(kind)
            _section_block(doc, title, icon, accent=_kind_accent(kind))
            _render_section_items(doc, items, kind=kind)
    else:
        _section_block(doc, "Краткое резюме", "◆", accent=SKY)
        _render_section_items(doc, data.get("executive_summary") or data.get("summary") or [], kind="info")
        if data.get("ai_insights"):
            _section_block(doc, "Ключевые выводы", "◆", accent=BLUE)
            _render_section_items(doc, data.get("ai_insights") or [], kind="info")
        if data.get("conclusions"):
            _section_block(doc, "Заключение", "◆", accent=BLUE)
            _render_section_items(doc, data.get("conclusions") or [], kind="info")
        if data.get("recommendations"):
            _section_block(doc, "Управленческие рекомендации", "◆", accent=GOOD)
            _render_section_items(doc, data.get("recommendations") or [], kind="plan")

    # Subject / school tables (данные без изменения; заголовок только для навигации по таблице)
    from docx.shared import Pt, RGBColor

    cap = doc.add_paragraph()
    cap.paragraph_format.space_before = Pt(8)
    cr = cap.add_run("Таблица · предметные результаты")
    cr.bold = True
    cr.font.size = Pt(10)
    cr.font.color.rgb = RGBColor(0x1B, 0x4F, 0x72)
    subj_rows = []
    for row in data.get("subject_rows") or []:
        subj_rows.append(
            [
                str(row.get("exam__subject") or "—"),
                str(row.get("participants") or ""),
                f"{float(row.get('avg') or 0):.2f}",
                str(row.get("pass_rate") if row.get("pass_rate") is not None else ""),
                str(row.get("quality_rate") if row.get("quality_rate") is not None else ""),
                str(row.get("risk") if row.get("risk") not in (None, "") else "—"),
            ]
        )
    _data_table(
        doc,
        ["Предмет", "Участ.", "Средний", "Усп., %", "Кач., %", "Риск"],
        subj_rows,
        widths_cm=[5.2, 1.6, 1.8, 1.7, 1.7, 3.2],
        num_cols={1, 2, 3, 4},
    )

    cap2 = doc.add_paragraph()
    cap2.paragraph_format.space_before = Pt(8)
    cr2 = cap2.add_run("Таблица · результаты образовательных организаций")
    cr2.bold = True
    cr2.font.size = Pt(10)
    cr2.font.color.rgb = RGBColor(0x1B, 0x4F, 0x72)
    school_rows = []
    pass_tones = []
    for row in data.get("school_rows") or []:
        pr = float(row.get("pass_rate") or 0)
        school_rows.append(
            [
                str(row.get("student__school__name") or "—"),
                str(row.get("student__school__code") or "—"),
                str(row.get("participants") or ""),
                f"{float(row.get('avg') or 0):.2f}",
                str(row.get("pass_rate") if row.get("pass_rate") is not None else ""),
            ]
        )
        pass_tones.append(pr)
    _data_table(
        doc,
        ["ОО", "Код", "Участ.", "Средний", "Усп., %"],
        school_rows,
        widths_cm=[8.0, 1.8, 1.6, 1.8, 1.8],
        num_cols={2, 3, 4},
        pass_rate_col=4,
        pass_tones=pass_tones,
    )

    return _save(doc)


def _save(doc) -> BytesIO:
    out = BytesIO()
    doc.save(out)
    out.seek(0)
    return out


def _fmt_delta(v) -> str:
    if v is None:
        return "—"
    try:
        f = float(v)
    except (TypeError, ValueError):
        return "—"
    return f"+{f}" if f > 0 else str(f)


def _tone_neutral() -> str:
    return SKY


def _tone_avg(v, et: str) -> str:
    try:
        f = float(v or 0)
    except (TypeError, ValueError):
        return SKY
    if et == "oge":
        if f >= 4:
            return GOOD
        if f < 3.5:
            return RISK
        return WARN
    if f >= 60:
        return GOOD
    if f < 45:
        return RISK
    return WARN


def _tone_pct(v, high: float, mid: float) -> str:
    try:
        f = float(v or 0)
    except (TypeError, ValueError):
        return SKY
    if f >= high:
        return GOOD
    if f < mid:
        return RISK
    return WARN


def _tone_fail(v) -> str:
    try:
        return RISK if int(v or 0) > 0 else GOOD
    except (TypeError, ValueError):
        return SKY


def _tone_delta(v) -> str:
    if v is None:
        return SKY
    try:
        f = float(v)
    except (TypeError, ValueError):
        return SKY
    if f > 0:
        return GOOD
    if f < 0:
        return WARN
    return SKY


def _detect_kind(title: str) -> str:
    t = (title or "").lower()
    if "пять главных проблем" in t or (t.startswith("пять") and "проблем" in t):
        return "problems"
    if "сильн" in t:
        return "strengths"
    if "прогноз риск" in t:
        return "risk_forecast"
    if "приоритетный план" in t or "план действий" in t:
        return "plan"
    if "рекоменд" in t:
        return "plan"
    if "риск" in t:
        return "risk_forecast"
    if "вывод" in t or "заключ" in t or "резюме" in t:
        return "info"
    return "info"


def _kind_icon(kind: str) -> str:
    return {
        "problems": "⚠",
        "strengths": "★",
        "risk_forecast": "◆",
        "plan": "▸",
        "info": "◆",
    }.get(kind, "◆")


def _kind_accent(kind: str) -> str:
    return {
        "problems": RISK,
        "strengths": GOOD,
        "risk_forecast": WARN,
        "plan": BLUE,
        "info": SKY,
    }.get(kind, SKY)


def _setup_page(doc) -> None:
    from docx.oxml.ns import qn
    from docx.shared import Cm, Pt

    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.left_margin = Cm(1.6)
    section.right_margin = Cm(1.6)
    section.top_margin = Cm(1.7)
    section.bottom_margin = Cm(1.7)
    section.header_distance = Cm(0.8)
    section.footer_distance = Cm(0.8)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(10)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
    pf = normal.paragraph_format
    pf.space_after = Pt(4)
    pf.space_before = Pt(0)
    pf.line_spacing = 1.12


def _setup_header_footer(doc, *, district: str, year, formed: str) -> None:
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.shared import Pt, RGBColor

    section = doc.sections[0]
    header = section.header
    header.is_linked_to_previous = False
    hp = header.paragraphs[0]
    hp.clear()
    r = hp.add_run(f"Свод результатов ГИА по району  ·  {district}  ·  {year or '—'}")
    r.font.size = Pt(8)
    r.font.color.rgb = RGBColor(0x5D, 0x6D, 0x7E)
    r.font.name = "Calibri"
    _p_border(hp, bottom=True)

    footer = section.footer
    footer.is_linked_to_previous = False
    fp = footer.paragraphs[0]
    fp.clear()
    run = fp.add_run(f"Дата формирования: {formed}  ·  версия {DOC_VERSION}  ·  стр. ")
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(0x5D, 0x6D, 0x7E)
    _add_page_field(fp)


def _add_page_field(paragraph) -> None:
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.shared import Pt, RGBColor

    fld = OxmlElement("w:fldChar")
    fld.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    fld2 = OxmlElement("w:fldChar")
    fld2.set(qn("w:fldCharType"), "end")
    r = paragraph.add_run()
    r._r.append(fld)
    r2 = paragraph.add_run()
    r2._r.append(instr)
    r3 = paragraph.add_run()
    r3._r.append(fld2)
    for rr in (r, r2, r3):
        rr.font.size = Pt(8)
        rr.font.color.rgb = RGBColor(0x5D, 0x6D, 0x7E)


def _title_page(doc, *, district: str, exam: str, year, formed: str) -> None:
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Pt, RGBColor, Cm

    # Logo mark
    logo = doc.add_table(rows=1, cols=2)
    logo.autofit = False
    c0, c1 = logo.rows[0].cells
    _set_cell_width(c0, Cm(1.4))
    _set_cell_width(c1, Cm(15))
    _shade_cell(c0, NAVY)
    p0 = c0.paragraphs[0]
    p0.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r0 = p0.add_run("АГ")
    r0.bold = True
    r0.font.size = Pt(14)
    r0.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    p1 = c1.paragraphs[0]
    r1 = p1.add_run("УПРАВЛЕНИЕ ОБРАЗОВАНИЯ")
    r1.bold = True
    r1.font.size = Pt(10)
    r1.font.color.rgb = RGBColor(0x1B, 0x4F, 0x72)
    p1b = c1.add_paragraph()
    r1b = p1b.add_run(district)
    r1b.font.size = Pt(12)
    r1b.bold = True
    r1b.font.color.rgb = RGBColor(0x0B, 0x1F, 0x3A)

    for _ in range(4):
        doc.add_paragraph("")

    eye = doc.add_paragraph()
    er = eye.add_run("ОФИЦИАЛЬНЫЙ АНАЛИТИЧЕСКИЙ ОТЧЁТ")
    er.bold = True
    er.font.size = Pt(9)
    er.font.color.rgb = RGBColor(0x5B, 0x9B, 0xD5)

    title = doc.add_paragraph()
    title.paragraph_format.space_before = Pt(8)
    tr = title.add_run("Свод результатов ГИА\nпо району")
    tr.bold = True
    tr.font.size = Pt(30)
    tr.font.color.rgb = RGBColor(0x0B, 0x1F, 0x3A)
    tr.font.name = "Calibri"

    sub = doc.add_paragraph()
    sr = sub.add_run(f"Муниципальный свод итогов {exam}")
    sr.font.size = Pt(12)
    sr.font.color.rgb = RGBColor(0x5D, 0x6D, 0x7E)

    for _ in range(2):
        doc.add_paragraph("")

    meta = doc.add_table(rows=4, cols=2)
    meta.autofit = False
    pairs = [
        ("Вид экзамена", exam),
        ("Отчётный период", str(year or "—")),
        ("Муниципалитет", district),
        ("Дата формирования", formed),
    ]
    for i, (k, v) in enumerate(pairs):
        ck, cv = meta.rows[i].cells
        _set_cell_width(ck, Cm(4.5))
        _set_cell_width(cv, Cm(11.5))
        pk = ck.paragraphs[0]
        rk = pk.add_run(k)
        rk.font.size = Pt(9)
        rk.font.color.rgb = RGBColor(0x5D, 0x6D, 0x7E)
        pv = cv.paragraphs[0]
        rv = pv.add_run(v)
        rv.bold = True
        rv.font.size = Pt(12)
        rv.font.color.rgb = RGBColor(0x0B, 0x1F, 0x3A)

    for _ in range(3):
        doc.add_paragraph("")

    ver = doc.add_paragraph()
    vr = ver.add_run(f"Версия документа {DOC_VERSION}")
    vr.font.size = Pt(9)
    vr.font.color.rgb = RGBColor(0x5D, 0x6D, 0x7E)

    conf = doc.add_paragraph()
    cr = conf.add_run("Конфиденциально · для служебного пользования")
    cr.font.size = Pt(9)
    cr.font.color.rgb = RGBColor(0x5D, 0x6D, 0x7E)

    doc.add_page_break()


def _toc(doc) -> None:
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.shared import Pt, RGBColor

    h = doc.add_paragraph()
    r = h.add_run("Содержание")
    r.bold = True
    r.font.size = Pt(16)
    r.font.color.rgb = RGBColor(0x0B, 0x1F, 0x3A)

    p = doc.add_paragraph()
    run = p.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = ' TOC \\o "1-2" \\h \\z \\u '
    fld_separate = OxmlElement("w:fldChar")
    fld_separate.set(qn("w:fldCharType"), "separate")
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_begin)
    run._r.append(instr)
    run._r.append(fld_separate)
    hint = p.add_run("(поле содержания обновляется в Word)")
    hint.font.size = Pt(8)
    hint.font.color.rgb = RGBColor(0x5D, 0x6D, 0x7E)
    run2 = p.add_run()
    run2._r.append(fld_end)
    doc.add_page_break()


def _section_block(doc, title: str, icon: str, *, accent: str) -> None:
    from docx.oxml import OxmlElement
    from docx.shared import Pt, RGBColor, Cm

    wrap = doc.add_table(rows=1, cols=2)
    wrap.autofit = False
    a, b = wrap.rows[0].cells
    _set_cell_width(a, Cm(0.25))
    _set_cell_width(b, Cm(16.5))
    _shade_cell(a, accent)
    p = b.paragraphs[0]
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(2)
    pPr = p._p.get_or_add_pPr()
    pPr.append(OxmlElement("w:keepNext"))
    ir = p.add_run(f"{icon}  ")
    ir.font.size = Pt(11)
    ir.font.color.rgb = RGBColor.from_string(accent)
    tr = p.add_run(title)
    tr.bold = True
    tr.font.size = Pt(13)
    tr.font.color.rgb = RGBColor(0x0B, 0x1F, 0x3A)
    _p_border(p, bottom=True)


def _kpi_dashboard(doc, cards: list[tuple[str, str, str, str]]) -> None:
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Cm, Pt, RGBColor

    # 2 rows × 4 cards
    for row_i in range(0, len(cards), 4):
        chunk = cards[row_i : row_i + 4]
        while len(chunk) < 4:
            chunk.append(("", "", "", SKY))
        table = doc.add_table(rows=1, cols=4)
        table.autofit = False
        for col, (icon, label, value, tone) in enumerate(chunk):
            cell = table.rows[0].cells[col]
            _set_cell_width(cell, Cm(4.2))
            _shade_cell(cell, SOFT)
            _cell_left_border(cell, tone)
            # clear default para
            p_icon = cell.paragraphs[0]
            p_icon.paragraph_format.space_after = Pt(0)
            p_icon.paragraph_format.space_before = Pt(4)
            if icon:
                ri = p_icon.add_run(icon)
                ri.font.size = Pt(9)
                ri.font.color.rgb = RGBColor.from_string(tone)
            p_val = cell.add_paragraph()
            p_val.paragraph_format.space_before = Pt(2)
            p_val.paragraph_format.space_after = Pt(0)
            rv = p_val.add_run(value)
            rv.bold = True
            rv.font.size = Pt(16)
            rv.font.color.rgb = RGBColor(0x0B, 0x1F, 0x3A)
            p_lab = cell.add_paragraph()
            p_lab.paragraph_format.space_before = Pt(1)
            p_lab.paragraph_format.space_after = Pt(4)
            rl = p_lab.add_run(label)
            rl.font.size = Pt(8)
            rl.font.color.rgb = RGBColor(0x5D, 0x6D, 0x7E)
        # spacer
        doc.add_paragraph("").paragraph_format.space_after = Pt(2)


def _render_section_items(doc, items: list, *, kind: str) -> None:
    if kind == "problems":
        for i, item in enumerate(items or [], start=1):
            _risk_card(doc, i, str(item))
    elif kind == "strengths":
        for i, item in enumerate(items or [], start=1):
            _strength_card(doc, i, str(item))
    elif kind == "risk_forecast":
        for i, item in enumerate(items or [], start=1):
            _risk_card(doc, i, str(item), badge="Риск")
    elif kind == "plan":
        for i, item in enumerate(items or [], start=1):
            _plan_card(doc, i, str(item))
    else:
        for item in items or []:
            _info_card(doc, str(item))


def _info_card(doc, text: str) -> None:
    if not text.strip():
        return
    from docx.shared import Pt, RGBColor, Cm

    t = doc.add_table(rows=1, cols=2)
    t.autofit = False
    a, b = t.rows[0].cells
    _set_cell_width(a, Cm(0.18))
    _set_cell_width(b, Cm(16.5))
    _shade_cell(a, SKY)
    _shade_cell(b, INFO_BG)
    p = b.paragraphs[0]
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    badge = p.add_run("Информация  ")
    badge.bold = True
    badge.font.size = Pt(7)
    badge.font.color.rgb = RGBColor(0x1B, 0x4F, 0x72)
    run = p.add_run(text)
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x1A, 0x23, 0x32)


def _risk_card(doc, num: int, text: str, badge: str = "Риск") -> None:
    if not text.strip():
        return
    from docx.shared import Pt, RGBColor, Cm

    t = doc.add_table(rows=1, cols=3)
    t.autofit = False
    n, a, b = t.rows[0].cells
    _set_cell_width(n, Cm(0.9))
    _set_cell_width(a, Cm(0.18))
    _set_cell_width(b, Cm(15.6))
    _shade_cell(n, RISK)
    _shade_cell(a, RISK)
    _shade_cell(b, RISK_BG)
    pn = n.paragraphs[0]
    rn = pn.add_run(f"{num:02d}")
    rn.bold = True
    rn.font.size = Pt(11)
    rn.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    p = b.paragraphs[0]
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(3)
    br = p.add_run(f"{badge}  ")
    br.bold = True
    br.font.size = Pt(7)
    br.font.color.rgb = RGBColor(0x9B, 0x2C, 0x2C)
    run = p.add_run(text)
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x1A, 0x23, 0x32)


def _strength_card(doc, num: int, text: str) -> None:
    if not text.strip():
        return
    from docx.shared import Pt, RGBColor, Cm

    t = doc.add_table(rows=1, cols=3)
    t.autofit = False
    n, a, b = t.rows[0].cells
    _set_cell_width(n, Cm(0.9))
    _set_cell_width(a, Cm(0.18))
    _set_cell_width(b, Cm(15.6))
    _shade_cell(n, GOOD)
    _shade_cell(a, GOOD)
    _shade_cell(b, GOOD_BG)
    pn = n.paragraphs[0]
    rn = pn.add_run(f"{num:02d}")
    rn.bold = True
    rn.font.size = Pt(11)
    rn.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    p = b.paragraphs[0]
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(3)
    br = p.add_run("Рост  ")
    br.bold = True
    br.font.size = Pt(7)
    br.font.color.rgb = RGBColor(0x1E, 0x7A, 0x4A)
    run = p.add_run(text)
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x1A, 0x23, 0x32)


def _plan_card(doc, num: int, text: str) -> None:
    if not text.strip():
        return
    from docx.shared import Pt, RGBColor, Cm

    t = doc.add_table(rows=1, cols=3)
    t.autofit = False
    n, a, b = t.rows[0].cells
    _set_cell_width(n, Cm(0.9))
    _set_cell_width(a, Cm(0.18))
    _set_cell_width(b, Cm(15.6))
    _shade_cell(n, BLUE)
    _shade_cell(a, SKY)
    _shade_cell(b, INFO_BG)
    pn = n.paragraphs[0]
    rn = pn.add_run(f"{num:02d}")
    rn.bold = True
    rn.font.size = Pt(11)
    rn.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    p = b.paragraphs[0]
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(3)
    br = p.add_run("Приоритет  ")
    br.bold = True
    br.font.size = Pt(7)
    br.font.color.rgb = RGBColor(0x1B, 0x4F, 0x72)
    run = p.add_run(text)
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x1A, 0x23, 0x32)


def _data_table(
    doc,
    headers: list[str],
    rows: list[list[str]],
    *,
    widths_cm: list[float],
    num_cols: set[int] | None = None,
    pass_rate_col: int | None = None,
    pass_tones: list[float] | None = None,
) -> None:
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.shared import Cm, Pt, RGBColor

    num_cols = num_cols or set()
    table = doc.add_table(rows=1, cols=len(headers))
    table.autofit = False
    table.allow_autofit = False

    # header
    hdr = table.rows[0]
    _set_row_header(hdr)
    for i, h in enumerate(headers):
        cell = hdr.cells[i]
        _set_cell_width(cell, Cm(widths_cm[i]))
        _shade_cell(cell, NAVY)
        p = cell.paragraphs[0]
        p.paragraph_format.space_before = Pt(3)
        p.paragraph_format.space_after = Pt(3)
        if i in num_cols:
            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        run = p.add_run(h)
        run.bold = True
        run.font.size = Pt(8)
        run.font.color.rgb = RGBColor(0xDC, 0xE8, 0xF3)
        run.font.name = "Calibri"

    for r_i, row in enumerate(rows):
        cells = table.add_row().cells
        _cant_split_row(table.rows[r_i + 1])
        for i, val in enumerate(row):
            cell = cells[i]
            _set_cell_width(cell, Cm(widths_cm[i]))
            if r_i % 2 == 1:
                _shade_cell(cell, "F7FAFC")
            # subtle pass-rate highlight (visual only)
            if pass_rate_col is not None and i == pass_rate_col and pass_tones and r_i < len(pass_tones):
                pr = pass_tones[r_i]
                if pr >= 85:
                    _shade_cell(cell, GOOD_BG)
                elif pr < 70:
                    _shade_cell(cell, RISK_BG)
                elif pr < 85:
                    _shade_cell(cell, WARN_BG)
            p = cell.paragraphs[0]
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after = Pt(2)
            if i in num_cols:
                p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            run = p.add_run(val)
            run.font.size = Pt(8)
            run.font.name = "Calibri"
            run.font.color.rgb = RGBColor(0x1A, 0x23, 0x32)
            # avoid letter-by-letter wrapping artifacts
            _no_snap_to_grid(p)

    doc.add_paragraph("").paragraph_format.space_after = Pt(6)


def _set_row_header(row) -> None:
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    tr = row._tr
    trPr = tr.get_or_add_trPr()
    tblHeader = OxmlElement("w:tblHeader")
    trPr.append(tblHeader)


def _cant_split_row(row) -> None:
    from docx.oxml import OxmlElement

    tr = row._tr
    trPr = tr.get_or_add_trPr()
    cant = OxmlElement("w:cantSplit")
    trPr.append(cant)


def _no_snap_to_grid(paragraph) -> None:
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    pPr = paragraph._p.get_or_add_pPr()
    # keep words together better
    snap = OxmlElement("w:snapToGrid")
    snap.set(qn("w:val"), "0")
    pPr.append(snap)


def _set_cell_width(cell, width) -> None:
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcW = tcPr.find(qn("w:tcW"))
    if tcW is None:
        tcW = OxmlElement("w:tcW")
        tcPr.append(tcW)
    twips = int(getattr(width, "twips", width))
    tcW.set(qn("w:w"), str(twips))
    tcW.set(qn("w:type"), "dxa")
    cell.width = width


def _cell_left_border(cell, color_hex: str) -> None:
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for edge in ("top", "left", "bottom", "right"):
        el = OxmlElement(f"w:{edge}")
        if edge == "left":
            el.set(qn("w:val"), "single")
            el.set(qn("w:sz"), "24")
            el.set(qn("w:color"), color_hex)
        else:
            el.set(qn("w:val"), "nil")
        tcBorders.append(el)
    tcPr.append(tcBorders)


def _p_border(paragraph, *, bottom: bool = False) -> None:
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    pPr = paragraph._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    if bottom:
        el = OxmlElement("w:bottom")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), "6")
        el.set(qn("w:space"), "4")
        el.set(qn("w:color"), LINE)
        pBdr.append(el)
    pPr.append(pBdr)


def _shade_cell(cell, fill_hex: str) -> None:
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill_hex)
    shd.set(qn("w:val"), "clear")
    tcPr.append(shd)
