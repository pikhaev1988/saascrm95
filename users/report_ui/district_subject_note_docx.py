"""
Финальное визуальное оформление DOCX «Предметная аналитическая справка».

Только Presentation/Layout. Тексты, метрики и структура UI — без изменений.
"""

from __future__ import annotations

from io import BytesIO
from typing import Any

from users.report_ui.district_subject_note import DOC_VERSION, build_district_subject_note_presentation

# Design system
NAVY = "0B1F3A"
BLUE = "1B4F72"
SKY = "5B9BD5"
GRAY = "5D6D7E"
LINE = "D6DEE8"
SOFT = "F5F8FB"
WHITE = "FFFFFF"
INK = "1A2332"

# Performance scale
CRIT = "9B2C2C"
CRIT_BG = "F8E8E8"
LOW = "C45C16"
LOW_BG = "F8EEE4"
MID = "A88B1A"
MID_BG = "F7F3DE"
GOOD = "2F8F55"
GOOD_BG = "E8F6EF"
EXCEL = "1B6B45"
EXCEL_BG = "DFF0E8"

INFO_BG = "EAF2F8"
DEC_BG = "E8F1FA"


def render_district_subject_note_docx(data: dict[str, Any]) -> BytesIO:
    from docx import Document

    doc = Document()
    _setup_page(doc)
    ui = build_district_subject_note_presentation(data)
    if not ui.get("has_data"):
        doc.add_paragraph(ui.get("message") or "Недостаточно данных.")
        return _save(doc)

    district = ui["district"]
    subject = ui["subject"]
    exam = ui["exam_label"]
    year = ui.get("year") or ""
    formed = ui.get("generated_at") or ""
    avg_label = ui.get("avg_label") or "Средний балл"
    level = ui.get("level") or {}
    resume = ui.get("resume") or {}
    kpi = ui.get("kpi") or {}

    _setup_header_footer(doc, district=district, subject=subject, year=year, formed=formed)
    _title_page(doc, district=district, subject=subject, exam=exam, year=year, formed=formed, level=level)
    _toc(doc)

    # ——— Executive Summary (passport + resume on one analytical page) ———
    _module(
        doc,
        num="01",
        title="Executive Summary · Паспорт предмета",
        lead="Dashboard ключевых показателей и краткая управленческая оценка.",
        accent=SKY,
        icon="◆",
    )
    _badge_row(
        doc,
        [
            (_level_badge(level), _tone_fill(level.get("tone"))),
            ("Паспорт предмета", INFO_BG),
            (f"Версия {DOC_VERSION}", SOFT),
        ],
    )
    _kpi_grid(doc, ui.get("passport") or [])
    _spacer(doc, 4)

    # Mini KPI strip
    _mini_kpi_strip(
        doc,
        [
            ("Заданий КИМ", str(kpi.get("tasks_count", "—")), EXCEL),
            ("Сильные", str(kpi.get("strong_count", "—")), GOOD),
            ("Проблемные", str(kpi.get("weak_count", "—")), LOW),
            ("Критические", str(kpi.get("critical_count", "—")), CRIT),
        ],
    )
    _spacer(doc, 6)

    _two_col_cards(
        doc,
        "Основные достижения",
        resume.get("achievements") or [],
        GOOD,
        GOOD_BG,
        "★",
        "Главные проблемы",
        resume.get("problems") or [],
        CRIT,
        CRIT_BG,
        "!",
    )
    for line in resume.get("level") or []:
        _insight_card(doc, line, accent=BLUE, badge="Оценка", bg=INFO_BG)
    for line in resume.get("assessment") or []:
        _insight_card(doc, line, accent=SKY, badge="Вывод", bg=SOFT)

    # ——— Statistics ———
    _module(
        doc,
        num="02",
        title="Предметная статистика",
        lead="Сводные показатели среза и распределение заданий по успешности.",
        accent=BLUE,
        icon="Σ",
    )
    _data_table(
        doc,
        ["Показатель", "Значение"],
        [
            ["Участники", str(kpi.get("participants") or kpi.get("total", ""))],
            ["Результаты (записи)", str(kpi.get("total", ""))],
            ["Образовательные организации", str(kpi.get("schools_count", ""))],
            [avg_label, _s(kpi.get("avg_score"))],
            ["Успеваемость, %", _s(kpi.get("pass_rate"))],
            ["Качество знаний, %", _s(kpi.get("quality_rate"))],
            ["Высокобалльники", str(kpi.get("high_count") if kpi.get("high_count") is not None else "—")],
        ],
        widths_cm=[11.2, 5.6],
        num_cols={1},
        compact=True,
    )
    _subhead(doc, "Распределение заданий по успешности")
    for row in ui.get("distribution") or []:
        _progress_row(
            doc,
            label=str(row.get("label") or ""),
            meta=f"{row.get('value')} зад. · {_s(row.get('percent'))}%",
            pct=float(row.get("percent") or 0),
            color=_band_color(str(row.get("label") or "")),
        )

    # ——— KIM ———
    _module(
        doc,
        num="03",
        title="Анализ выполнения заданий КИМ",
        lead="Процент выполнения, сложность и тематическая привязка каждого задания.",
        accent=BLUE,
        icon="☰",
    )
    kim_rows = []
    rate_vals = []
    for t in ui.get("tasks") or []:
        rate = float(t.get("success_rate") or 0)
        kim_rows.append(
            [
                str(t["number"]),
                f"{_s(rate)}%",
                str(t.get("max_score") if t.get("max_score") is not None else "—"),
                str(t.get("difficulty") or "—"),
                _short(str(t.get("topic") or "—"), 42),
            ]
        )
        rate_vals.append(rate)
    _data_table(
        doc,
        ["№", "Выполнение", "Макс.", "Сложность", "Тема / линия"],
        kim_rows,
        widths_cm=[1.1, 2.2, 1.6, 2.6, 9.3],
        num_cols={0, 1, 2},
        rate_col=1,
        rate_vals=rate_vals,
        compact=True,
        topic_cols={4},
    )

    _module(
        doc,
        num="03+",
        title="Лучшие задания",
        lead="Задания с устойчиво высоким процентом выполнения.",
        accent=GOOD,
        icon="★",
    )
    _task_cards(doc, ui.get("best_tasks") or [], mode="best")

    _module(
        doc,
        num="03−",
        title="Проблемные задания",
        lead="Задания с низкой успешностью — зона методического внимания.",
        accent=CRIT,
        icon="⚠",
    )
    _task_cards(doc, ui.get("problem_tasks") or [], mode="problem")

    # ——— Topics ———
    _module(
        doc,
        num="04",
        title="Анализ содержательных разделов",
        lead="Группировка заданий по темам и линиям содержания.",
        accent=SKY,
        icon="▣",
    )
    topic_rows = []
    topic_rates = []
    for t in ui.get("topics") or []:
        topic_rows.append(
            [
                _short(str(t.get("topic") or "—"), 48),
                str(t.get("tasks_count") or ""),
                ", ".join((t.get("tasks") or [])[:6]),
                f"{_s(t.get('avg_success'))}%",
                str(t.get("status") or "—"),
            ]
        )
        topic_rates.append(float(t.get("avg_success") or 0))
    _data_table(
        doc,
        ["Тема / раздел", "N", "Номера", "Ср. %", "Статус"],
        topic_rows,
        widths_cm=[6.0, 1.2, 4.0, 2.0, 3.6],
        num_cols={1, 3},
        rate_col=3,
        rate_vals=topic_rates,
        compact=True,
        topic_cols={0},
    )

    # ——— Deficits ———
    _module(
        doc,
        num="05",
        title="Карта предметных дефицитов",
        lead="Управленческая панель ТОП дефицитов с приоритетом и риском.",
        accent=CRIT,
        icon="!",
    )
    for d in ui.get("deficits") or []:
        _deficit_card(doc, d)

    # ——— Strengths ———
    _module(
        doc,
        num="06",
        title="Карта сильных сторон",
        lead="Темы и задания с устойчиво высоким результатом.",
        accent=GOOD,
        icon="★",
    )
    for s in ui.get("strengths") or []:
        _strength_card(doc, s.get("title") or "", s.get("detail") or "")

    # ——— Comparison ———
    _module(
        doc,
        num="07",
        title="Сравнение: муниципалитет · республика",
        lead="Отклонения по ключевым показателям предмета (только загруженные протоколы).",
        accent=BLUE,
        icon="↔",
    )
    cmp_rows = []
    cmp_tones = []
    for row in ui.get("comparison") or []:
        cmp_rows.append(
            [row["indicator"], row["municipality"], row["republic"], row["delta"]]
        )
        cmp_tones.append(row.get("tone") or "neutral")
    _data_table(
        doc,
        ["Показатель", "Муниципалитет", "Республика", "Δ"],
        cmp_rows,
        widths_cm=[5.4, 3.6, 3.6, 4.2],
        num_cols={1, 2, 3},
        tone_col=3,
        tones=cmp_tones,
        compact=True,
    )
    _insight_card(
        doc,
        "Сравнение с республикой выполняется по тому же предмету и году "
        "на основе загруженных протоколов муниципалитетов.",
        accent=SKY,
        badge="Справка",
        bg=INFO_BG,
    )

    # ——— Schools ———
    _module(
        doc,
        num="08",
        title="Анализ образовательных организаций",
        lead="Компактный рейтинг ОО по среднему результату, качеству и успеваемости.",
        accent=NAVY,
        icon="⌂",
    )
    school_rows = []
    pass_vals = []
    for i, row in enumerate(ui.get("schools") or [], start=1):
        qr = row.get("quality_rate")
        school_rows.append(
            [
                str(i),
                _short(str(row.get("name") or "—"), 55),
                str(row.get("participants") or ""),
                _s(row.get("avg")),
                f"{_s(qr)}%" if qr is not None else "—",
                f"{_s(row.get('pass_rate'))}%",
            ]
        )
        pass_vals.append(float(row.get("pass_rate") or 0))
    _data_table(
        doc,
        ["№", "ОО", "Участ.", avg_label, "Кач., %", "Усп., %"],
        school_rows,
        widths_cm=[0.9, 8.0, 1.5, 2.2, 1.9, 2.0],
        num_cols={0, 2, 3, 4, 5},
        rate_col=5,
        rate_vals=pass_vals,
        compact=True,
        topic_cols={1},
    )

    # ——— Recommendations ———
    _module(
        doc,
        num="09",
        title="Методические рекомендации",
        lead="Адресные меры для учителя, МО, администрации школы и муниципалитета.",
        accent=BLUE,
        icon="▸",
    )
    for group in ui.get("recommendations") or []:
        _audience_banner(doc, group.get("audience") or "Рекомендации")
        for item in group.get("items") or []:
            _reco_card(doc, item)

    # ——— Practice ———
    _module(
        doc,
        num="10",
        title="Практические рекомендации по подготовке",
        lead="Что изменить в учебном процессе, диагностике, повторении и работе с заданиями.",
        accent=SKY,
        icon="⚙",
    )
    practice = ui.get("practice") or {}
    _practice_block(doc, "Учебный процесс", practice.get("process") or [], BLUE)
    _practice_block(doc, "Диагностика", practice.get("diagnostics") or [], SKY)
    _practice_block(doc, "Повторение", practice.get("revision") or [], GOOD)
    _practice_block(doc, "Работа с заданиями", practice.get("tasks") or [], LOW)

    # ——— Conclusion ———
    _module(
        doc,
        num="11",
        title="Заключение",
        lead="Итоговая оценка, ключевые выводы и приоритеты развития.",
        accent=NAVY,
        icon="◆",
    )
    _badge_row(
        doc,
        [
            (_level_badge(level), _tone_fill(level.get("tone"))),
            ("Итоговая аналитическая страница", INFO_BG),
        ],
    )
    for i, line in enumerate(ui.get("conclusion") or [], start=1):
        _conclusion_card(doc, i, line)

    return _save(doc)


# ── page / chrome ─────────────────────────────────────────────


def _save(doc) -> BytesIO:
    out = BytesIO()
    doc.save(out)
    out.seek(0)
    return out


def _setup_page(doc) -> None:
    from docx.oxml.ns import qn
    from docx.shared import Cm, Pt

    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.left_margin = Cm(1.5)
    section.right_margin = Cm(1.5)
    section.top_margin = Cm(1.5)
    section.bottom_margin = Cm(1.5)
    section.header_distance = Cm(0.7)
    section.footer_distance = Cm(0.7)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(10)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
    pf = normal.paragraph_format
    pf.space_after = Pt(2)
    pf.space_before = Pt(0)
    pf.line_spacing = 1.08


def _setup_header_footer(doc, *, district: str, subject: str, year, formed: str) -> None:
    from docx.shared import Pt, RGBColor

    section = doc.sections[0]
    header = section.header
    header.is_linked_to_previous = False
    hp = header.paragraphs[0]
    hp.clear()
    r = hp.add_run(
        f"Предметная аналитическая справка  ·  {subject}  ·  {district}  ·  {year or '—'}"
    )
    r.font.size = Pt(8)
    r.font.color.rgb = RGBColor(0x5D, 0x6D, 0x7E)
    r.font.name = "Calibri"
    _p_border(hp, bottom=True)

    footer = section.footer
    footer.is_linked_to_previous = False
    fp = footer.paragraphs[0]
    fp.clear()
    run = fp.add_run(f"Дата формирования: {formed}  ·  версия {DOC_VERSION}  ·  стр. ")
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(0x5D, 0x6D, 0x7E)
    _add_page_field(fp)


def _add_page_field(paragraph) -> None:
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.shared import Pt, RGBColor

    fld = OxmlElement("w:fldChar")
    fld.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    fld2 = OxmlElement("w:fldChar")
    fld2.set(qn("w:fldCharType"), "end")
    r = paragraph.add_run()
    r._r.append(fld)
    r2 = paragraph.add_run()
    r2._r.append(instr)
    r3 = paragraph.add_run()
    r3._r.append(fld2)
    for rr in (r, r2, r3):
        rr.font.size = Pt(8)
        rr.font.color.rgb = RGBColor(0x5D, 0x6D, 0x7E)


def _title_page(doc, *, district: str, subject: str, exam: str, year, formed: str, level: dict) -> None:
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Cm, Pt, RGBColor

    logo = doc.add_table(rows=1, cols=2)
    logo.autofit = False
    c0, c1 = logo.rows[0].cells
    _set_cell_width(c0, Cm(1.35))
    _set_cell_width(c1, Cm(15.3))
    _shade_cell(c0, NAVY)
    p0 = c0.paragraphs[0]
    p0.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r0 = p0.add_run("ЦО")
    r0.bold = True
    r0.font.size = Pt(12)
    r0.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    p1 = c1.paragraphs[0]
    r1 = p1.add_run("ЦЕНТР ОЦЕНКИ КАЧЕСТВА ОБРАЗОВАНИЯ")
    r1.bold = True
    r1.font.size = Pt(9)
    r1.font.color.rgb = RGBColor(0x1B, 0x4F, 0x72)
    p1b = c1.add_paragraph()
    r1b = p1b.add_run(f"Управление образования · {district}")
    r1b.bold = True
    r1b.font.size = Pt(11)
    r1b.font.color.rgb = RGBColor(0x0B, 0x1F, 0x3A)

    for _ in range(2):
        doc.add_paragraph("")

    eye = doc.add_paragraph()
    er = eye.add_run("ОФИЦИАЛЬНЫЙ МЕТОДИЧЕСКИЙ ОТЧЁТ")
    er.bold = True
    er.font.size = Pt(8)
    er.font.color.rgb = RGBColor(0x5B, 0x9B, 0xD5)

    title = doc.add_paragraph()
    title.paragraph_format.space_before = Pt(6)
    tr = title.add_run("Предметная аналитическая\nсправка")
    tr.bold = True
    tr.font.size = Pt(26)
    tr.font.color.rgb = RGBColor(0x0B, 0x1F, 0x3A)

    subj = doc.add_paragraph()
    sr = subj.add_run(subject)
    sr.bold = True
    sr.font.size = Pt(16)
    sr.font.color.rgb = RGBColor(0x1B, 0x4F, 0x72)

    meta = doc.add_paragraph()
    mr = meta.add_run(f"{exam}  ·  {year or '—'}  ·  {district}")
    mr.font.size = Pt(11)
    mr.font.color.rgb = RGBColor(0x5D, 0x6D, 0x7E)

    for _ in range(2):
        doc.add_paragraph("")

    _badge_row(
        doc,
        [
            (_level_badge(level), _tone_fill(level.get("tone"))),
            (f"Дата формирования: {formed}", SOFT),
            (f"Версия {DOC_VERSION}", INFO_BG),
        ],
    )

    for _ in range(2):
        doc.add_paragraph("")
    foot = doc.add_paragraph()
    fr = foot.add_run("Для методических служб · РМО · педагогов · администрации ОО")
    fr.font.size = Pt(9)
    fr.font.color.rgb = RGBColor(0x5D, 0x6D, 0x7E)
    doc.add_page_break()


def _toc(doc) -> None:
    from docx.shared import Pt, RGBColor

    h = doc.add_paragraph()
    r = h.add_run("Содержание")
    r.bold = True
    r.font.size = Pt(14)
    r.font.color.rgb = RGBColor(0x0B, 0x1F, 0x3A)
    items = [
        "01  Executive Summary · Паспорт предмета",
        "02  Предметная статистика",
        "03  Анализ выполнения заданий КИМ",
        "04  Анализ содержательных разделов",
        "05  Карта предметных дефицитов",
        "06  Карта сильных сторон",
        "07  Сравнение с республикой",
        "08  Анализ образовательных организаций",
        "09  Методические рекомендации",
        "10  Практические рекомендации по подготовке",
        "11  Заключение",
    ]
    for item in items:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(3)
        run = p.add_run(item)
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(0x1A, 0x23, 0x32)
    doc.add_page_break()


# ── modules / cards ───────────────────────────────────────────


def _module(doc, *, num: str, title: str, lead: str, accent: str, icon: str) -> None:
    from docx.oxml import OxmlElement
    from docx.shared import Cm, Pt, RGBColor

    wrap = doc.add_table(rows=1, cols=3)
    wrap.autofit = False
    a, n, b = wrap.rows[0].cells
    _set_cell_width(a, Cm(0.22))
    _set_cell_width(n, Cm(1.15))
    _set_cell_width(b, Cm(15.3))
    _shade_cell(a, accent)
    _shade_cell(n, accent)
    pn = n.paragraphs[0]
    rn = pn.add_run(num)
    rn.bold = True
    rn.font.size = Pt(9)
    rn.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

    p = b.paragraphs[0]
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(0)
    pPr = p._p.get_or_add_pPr()
    pPr.append(OxmlElement("w:keepNext"))
    ir = p.add_run(f"{icon}  ")
    ir.font.size = Pt(10)
    ir.font.color.rgb = RGBColor.from_string(accent)
    tr = p.add_run(title)
    tr.bold = True
    tr.font.size = Pt(12)
    tr.font.color.rgb = RGBColor(0x0B, 0x1F, 0x3A)

    p2 = b.add_paragraph()
    p2.paragraph_format.space_before = Pt(1)
    p2.paragraph_format.space_after = Pt(4)
    lr = p2.add_run(lead)
    lr.font.size = Pt(8)
    lr.font.color.rgb = RGBColor(0x5D, 0x6D, 0x7E)
    _p_border(p2, bottom=True)


def _kpi_grid(doc, cards: list[dict]) -> None:
    from docx.shared import Cm, Pt, RGBColor

    for row_i in range(0, len(cards), 3):
        chunk = cards[row_i : row_i + 3]
        while len(chunk) < 3:
            chunk.append({"icon": "", "label": "", "value": "", "tone": "sky"})
        table = doc.add_table(rows=1, cols=3)
        table.autofit = False
        for col, card in enumerate(chunk):
            cell = table.rows[0].cells[col]
            _set_cell_width(cell, Cm(5.7))
            _shade_cell(cell, WHITE)
            tone = _tone_hex(card.get("tone"))
            _cell_box(cell, left=tone)
            p_icon = cell.paragraphs[0]
            p_icon.paragraph_format.space_before = Pt(5)
            p_icon.paragraph_format.space_after = Pt(0)
            if card.get("icon"):
                ri = p_icon.add_run(str(card["icon"]))
                ri.font.size = Pt(9)
                ri.font.color.rgb = RGBColor.from_string(tone)
            p_val = cell.add_paragraph()
            p_val.paragraph_format.space_before = Pt(1)
            p_val.paragraph_format.space_after = Pt(0)
            rv = p_val.add_run(str(card.get("value") or ""))
            rv.bold = True
            rv.font.size = Pt(15)
            rv.font.color.rgb = RGBColor(0x0B, 0x1F, 0x3A)
            p_lab = cell.add_paragraph()
            p_lab.paragraph_format.space_before = Pt(1)
            p_lab.paragraph_format.space_after = Pt(5)
            rl = p_lab.add_run(str(card.get("label") or ""))
            rl.font.size = Pt(7.5)
            rl.font.color.rgb = RGBColor(0x5D, 0x6D, 0x7E)
        _spacer(doc, 2)


def _mini_kpi_strip(doc, items: list[tuple[str, str, str]]) -> None:
    from docx.shared import Cm, Pt, RGBColor

    table = doc.add_table(rows=1, cols=len(items))
    table.autofit = False
    for i, (label, value, color) in enumerate(items):
        cell = table.rows[0].cells[i]
        _set_cell_width(cell, Cm(4.25))
        _shade_cell(cell, SOFT)
        _cell_box(cell, left=color)
        p = cell.paragraphs[0]
        p.paragraph_format.space_before = Pt(3)
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(value)
        r.bold = True
        r.font.size = Pt(13)
        r.font.color.rgb = RGBColor(0x0B, 0x1F, 0x3A)
        p2 = cell.add_paragraph()
        p2.paragraph_format.space_after = Pt(3)
        r2 = p2.add_run(label)
        r2.font.size = Pt(7)
        r2.font.color.rgb = RGBColor(0x5D, 0x6D, 0x7E)


def _insight_card(doc, text: str, *, accent: str, badge: str, bg: str) -> None:
    if not str(text).strip():
        return
    from docx.shared import Cm, Pt, RGBColor

    t = doc.add_table(rows=1, cols=2)
    t.autofit = False
    a, b = t.rows[0].cells
    _set_cell_width(a, Cm(0.18))
    _set_cell_width(b, Cm(16.5))
    _shade_cell(a, accent)
    _shade_cell(b, bg)
    p = b.paragraphs[0]
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(3)
    br = p.add_run(f"{badge}  ")
    br.bold = True
    br.font.size = Pt(7)
    br.font.color.rgb = RGBColor.from_string(accent)
    run = p.add_run(text)
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x1A, 0x23, 0x32)
    _spacer(doc, 2)


def _two_col_cards(doc, t1, items1, c1, bg1, i1, t2, items2, c2, bg2, i2) -> None:
    from docx.shared import Cm, Pt, RGBColor

    table = doc.add_table(rows=1, cols=2)
    table.autofit = False
    left, right = table.rows[0].cells
    _set_cell_width(left, Cm(8.4))
    _set_cell_width(right, Cm(8.4))
    for cell, title, items, color, bg, icon in (
        (left, t1, items1, c1, bg1, i1),
        (right, t2, items2, c2, bg2, i2),
    ):
        _shade_cell(cell, bg)
        _cell_box(cell, left=color)
        p = cell.paragraphs[0]
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after = Pt(2)
        r = p.add_run(f"{icon}  {title}")
        r.bold = True
        r.font.size = Pt(9)
        r.font.color.rgb = RGBColor.from_string(color)
        for item in items or ["—"]:
            pi = cell.add_paragraph()
            pi.paragraph_format.space_after = Pt(1)
            ri = pi.add_run(f"•  {item}")
            ri.font.size = Pt(8)
            ri.font.color.rgb = RGBColor(0x1A, 0x23, 0x32)
    _spacer(doc, 4)


def _task_cards(doc, rows: list[dict], *, mode: str) -> None:
    accent = GOOD if mode == "best" else CRIT
    bg = GOOD_BG if mode == "best" else CRIT_BG
    badge = "Лучший показатель" if mode == "best" else "Требует внимания"
    for t in rows:
        rate = float(t.get("success_rate") or 0)
        _panel_card(
            doc,
            title=f"Задание №{t.get('number')}  ·  {_s(rate)}%",
            body=f"{t.get('topic') or '—'}  ·  сложность: {t.get('difficulty') or '—'}",
            accent=accent,
            bg=bg,
            badge=badge,
            pct=rate,
        )


def _deficit_card(doc, d: dict) -> None:
    risk = str(d.get("risk") or "")
    if "крит" in risk.lower():
        accent, bg, badge = CRIT, CRIT_BG, "Критический дефицит"
    elif "высок" in risk.lower():
        accent, bg, badge = LOW, LOW_BG, "Высокий риск"
    else:
        accent, bg, badge = MID, MID_BG, "Требует внимания"
    title = f"{d.get('rank', '')}. {d.get('title') or ''}"
    body = (
        f"{d.get('description') or ''}\n"
        f"Задания: {d.get('tasks') or '—'}  ·  Риск: {d.get('risk') or '—'}  ·  Приоритет: {d.get('priority') or '—'}"
    )
    _panel_card(doc, title=title, body=body, accent=accent, bg=bg, badge=badge, pct=float(d.get("success_rate") or 0))


def _strength_card(doc, title: str, detail: str) -> None:
    _panel_card(
        doc,
        title=title,
        body=detail,
        accent=GOOD,
        bg=GOOD_BG,
        badge="Сильная сторона",
        pct=None,
    )


def _panel_card(doc, *, title: str, body: str, accent: str, bg: str, badge: str, pct) -> None:
    from docx.shared import Cm, Pt, RGBColor

    t = doc.add_table(rows=1, cols=2)
    t.autofit = False
    a, b = t.rows[0].cells
    _set_cell_width(a, Cm(0.18))
    _set_cell_width(b, Cm(16.5))
    _shade_cell(a, accent)
    _shade_cell(b, bg)
    p = b.paragraphs[0]
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(0)
    br = p.add_run(f"{badge}  ")
    br.bold = True
    br.font.size = Pt(7)
    br.font.color.rgb = RGBColor.from_string(accent)
    tr = p.add_run(title)
    tr.bold = True
    tr.font.size = Pt(9)
    tr.font.color.rgb = RGBColor(0x0B, 0x1F, 0x3A)
    p2 = b.add_paragraph()
    p2.paragraph_format.space_before = Pt(1)
    p2.paragraph_format.space_after = Pt(2)
    for i, line in enumerate(str(body).split("\n")):
        if i:
            p2 = b.add_paragraph()
            p2.paragraph_format.space_after = Pt(1)
        r2 = p2.add_run(line)
        r2.font.size = Pt(8)
        r2.font.color.rgb = RGBColor(0x1A, 0x23, 0x32)
    if pct is not None:
        _progress_in_cell(b, float(pct), _rate_color(float(pct)))
    _spacer(doc, 2)


def _reco_card(doc, item: dict) -> None:
    from docx.shared import Cm, Pt, RGBColor

    priority = str(item.get("priority") or "—")
    accent = CRIT if "высок" in priority.lower() else (MID if "сред" in priority.lower() else BLUE)
    bg = CRIT_BG if accent == CRIT else (MID_BG if accent == MID else INFO_BG)
    t = doc.add_table(rows=1, cols=2)
    t.autofit = False
    a, b = t.rows[0].cells
    _set_cell_width(a, Cm(0.18))
    _set_cell_width(b, Cm(16.5))
    _shade_cell(a, accent)
    _shade_cell(b, bg)
    p = b.paragraphs[0]
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(0)
    br = p.add_run("Приоритет  ")
    br.bold = True
    br.font.size = Pt(7)
    br.font.color.rgb = RGBColor.from_string(accent)
    pr = p.add_run(priority)
    pr.bold = True
    pr.font.size = Pt(8)
    pr.font.color.rgb = RGBColor(0x0B, 0x1F, 0x3A)
    p2 = b.add_paragraph()
    p2.paragraph_format.space_after = Pt(1)
    r2 = p2.add_run(str(item.get("text") or ""))
    r2.font.size = Pt(9)
    r2.font.color.rgb = RGBColor(0x1A, 0x23, 0x32)
    p3 = b.add_paragraph()
    p3.paragraph_format.space_after = Pt(3)
    r3 = p3.add_run(
        f"Эффект: {item.get('effect') or '—'}  ·  Связь с дефицитом: {item.get('deficit') or '—'}"
    )
    r3.font.size = Pt(7.5)
    r3.font.color.rgb = RGBColor(0x5D, 0x6D, 0x7E)
    _spacer(doc, 2)


def _practice_block(doc, title: str, items: list[str], accent: str) -> None:
    _subhead(doc, title)
    for item in items or []:
        _insight_card(doc, item, accent=accent, badge="Практика", bg=SOFT)


def _conclusion_card(doc, num: int, text: str) -> None:
    if not str(text).strip():
        return
    from docx.shared import Cm, Pt, RGBColor

    t = doc.add_table(rows=1, cols=2)
    t.autofit = False
    n, b = t.rows[0].cells
    _set_cell_width(n, Cm(0.95))
    _set_cell_width(b, Cm(15.75))
    _shade_cell(n, NAVY)
    _shade_cell(b, SOFT)
    pn = n.paragraphs[0]
    rn = pn.add_run(f"{num:02d}")
    rn.bold = True
    rn.font.size = Pt(11)
    rn.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    p = b.paragraphs[0]
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x1A, 0x23, 0x32)
    _spacer(doc, 2)


def _audience_banner(doc, text: str) -> None:
    from docx.shared import Cm, Pt, RGBColor

    t = doc.add_table(rows=1, cols=1)
    t.autofit = False
    cell = t.rows[0].cells[0]
    _set_cell_width(cell, Cm(16.8))
    _shade_cell(cell, DEC_BG)
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(10)
    r.font.color.rgb = RGBColor(0x1B, 0x4F, 0x72)
    _spacer(doc, 3)


def _badge_row(doc, items: list[tuple[str, str]]) -> None:
    from docx.shared import Cm, Pt, RGBColor

    cols = max(1, len(items))
    t = doc.add_table(rows=1, cols=cols)
    t.autofit = False
    w = 16.8 / cols
    for i, (text, fill) in enumerate(items):
        cell = t.rows[0].cells[i]
        _set_cell_width(cell, Cm(w))
        _shade_cell(cell, fill)
        p = cell.paragraphs[0]
        p.paragraph_format.space_before = Pt(3)
        p.paragraph_format.space_after = Pt(3)
        r = p.add_run(text)
        r.bold = True
        r.font.size = Pt(8)
        r.font.color.rgb = RGBColor(0x0B, 0x1F, 0x3A)
    _spacer(doc, 4)


def _subhead(doc, text: str) -> None:
    from docx.shared import Pt, RGBColor

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(9)
    r.font.color.rgb = RGBColor(0x1B, 0x4F, 0x72)


def _spacer(doc, pt: int = 4) -> None:
    from docx.shared import Pt

    p = doc.add_paragraph("")
    p.paragraph_format.space_after = Pt(pt)


# ── progress / tables ─────────────────────────────────────────


def _progress_row(doc, *, label: str, meta: str, pct: float, color: str) -> None:
    from docx.shared import Cm, Pt, RGBColor

    wrap = doc.add_table(rows=1, cols=2)
    wrap.autofit = False
    left, right = wrap.rows[0].cells
    _set_cell_width(left, Cm(4.2))
    _set_cell_width(right, Cm(12.6))
    _shade_cell(left, SOFT)
    p = left.paragraphs[0]
    p.paragraph_format.space_before = Pt(3)
    r = p.add_run(label)
    r.bold = True
    r.font.size = Pt(8)
    r.font.color.rgb = RGBColor(0x0B, 0x1F, 0x3A)
    p2 = left.add_paragraph()
    p2.paragraph_format.space_after = Pt(3)
    r2 = p2.add_run(meta)
    r2.font.size = Pt(7)
    r2.font.color.rgb = RGBColor(0x5D, 0x6D, 0x7E)
    _progress_in_cell(right, pct, color)
    _spacer(doc, 2)


def _progress_in_cell(cell, pct: float, color: str) -> None:
    from docx.shared import Pt, RGBColor

    pct = max(0.0, min(100.0, float(pct or 0)))
    fill_n = max(1, int(round(pct / 5))) if pct > 0 else 0
    empty_n = 20 - fill_n
    p = cell.paragraphs[0] if cell.paragraphs else cell.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(2)
    filled = "█" * fill_n
    empty = "░" * empty_n
    r = p.add_run(filled)
    r.font.size = Pt(8)
    r.font.color.rgb = RGBColor.from_string(color)
    r2 = p.add_run(empty)
    r2.font.size = Pt(8)
    r2.font.color.rgb = RGBColor(0xD5, 0xDD, 0xE5)
    r3 = p.add_run(f"  {_s(pct)}%")
    r3.bold = True
    r3.font.size = Pt(8)
    r3.font.color.rgb = RGBColor(0x0B, 0x1F, 0x3A)


def _data_table(
    doc,
    headers: list[str],
    rows: list[list[str]],
    *,
    widths_cm: list[float],
    num_cols: set[int] | None = None,
    rate_col: int | None = None,
    rate_vals: list[float] | None = None,
    tone_col: int | None = None,
    tones: list[str] | None = None,
    compact: bool = False,
    topic_cols: set[int] | None = None,
) -> None:
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_ROW_HEIGHT_RULE
    from docx.shared import Cm, Pt, RGBColor

    num_cols = num_cols or set()
    topic_cols = topic_cols or set()
    table = doc.add_table(rows=1, cols=len(headers))
    table.autofit = False
    table.allow_autofit = False

    hdr = table.rows[0]
    _set_row_header(hdr)
    for i, h in enumerate(headers):
        cell = hdr.cells[i]
        _set_cell_width(cell, Cm(widths_cm[i]))
        _shade_cell(cell, NAVY)
        p = cell.paragraphs[0]
        p.paragraph_format.space_before = Pt(2 if compact else 3)
        p.paragraph_format.space_after = Pt(2 if compact else 3)
        if i in num_cols:
            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        run = p.add_run(h)
        run.bold = True
        run.font.size = Pt(7.5 if compact else 8)
        run.font.color.rgb = RGBColor(0xDC, 0xE8, 0xF3)
        run.font.name = "Calibri"

    for r_i, row in enumerate(rows):
        cells = table.add_row().cells
        row_obj = table.rows[r_i + 1]
        _cant_split_row(row_obj)
        if compact:
            try:
                row_obj.height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST
                row_obj.height = Cm(0.55)
            except Exception:
                pass
        for i, val in enumerate(row):
            cell = cells[i]
            _set_cell_width(cell, Cm(widths_cm[i]))
            fill = "F7FAFC" if r_i % 2 else WHITE
            if rate_col is not None and i == rate_col and rate_vals and r_i < len(rate_vals):
                fill = _rate_bg(rate_vals[r_i])
            if tone_col is not None and i == tone_col and tones and r_i < len(tones):
                tone = tones[r_i]
                if tone == "good":
                    fill = GOOD_BG
                elif tone == "risk":
                    fill = CRIT_BG
                elif tone == "warn":
                    fill = LOW_BG
            _shade_cell(cell, fill)
            p = cell.paragraphs[0]
            p.paragraph_format.space_before = Pt(1 if compact else 2)
            p.paragraph_format.space_after = Pt(1 if compact else 2)
            if i in num_cols:
                p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            run = p.add_run(str(val))
            run.font.size = Pt(7 if (compact and i in topic_cols) else (8 if not compact else 7.5))
            run.font.name = "Calibri"
            run.font.color.rgb = RGBColor(0x1A, 0x23, 0x32)
            _no_snap_to_grid(p)
    _spacer(doc, 4)


# ── helpers ───────────────────────────────────────────────────


def _s(v) -> str:
    if v is None:
        return "—"
    try:
        f = float(v)
    except (TypeError, ValueError):
        return str(v)
    if abs(f - round(f)) < 1e-9:
        return str(int(round(f)))
    return f"{f:.2f}".rstrip("0").rstrip(".")


def _short(text: str, n: int) -> str:
    t = " ".join(str(text or "").split())
    return t if len(t) <= n else t[: n - 1].rstrip() + "…"


def _level_badge(level: dict) -> str:
    return f"Общая оценка: {level.get('icon', '◆')} {level.get('label', '—')}"


def _tone_hex(tone) -> str:
    return {
        "sky": SKY,
        "good": GOOD,
        "warn": LOW,
        "risk": CRIT,
        "critical": CRIT,
        "neutral": SKY,
    }.get(tone or "sky", SKY)


def _tone_fill(tone) -> str:
    return {
        "good": GOOD_BG,
        "warn": LOW_BG,
        "risk": CRIT_BG,
        "sky": INFO_BG,
        "neutral": SOFT,
    }.get(tone or "neutral", SOFT)


def _rate_color(pct: float) -> str:
    if pct <= 30:
        return CRIT
    if pct <= 50:
        return LOW
    if pct <= 70:
        return MID
    if pct <= 85:
        return GOOD
    return EXCEL


def _rate_bg(pct: float) -> str:
    if pct <= 30:
        return CRIT_BG
    if pct <= 50:
        return LOW_BG
    if pct <= 70:
        return MID_BG
    if pct <= 85:
        return GOOD_BG
    return EXCEL_BG


def _band_color(label: str) -> str:
    l = label.lower()
    if "0" in l and "29" in l:
        return CRIT
    if "30" in l:
        return LOW
    if "50" in l:
        return MID
    if "70" in l:
        return GOOD
    return EXCEL


def _set_row_header(row) -> None:
    from docx.oxml import OxmlElement

    tr = row._tr
    trPr = tr.get_or_add_trPr()
    trPr.append(OxmlElement("w:tblHeader"))


def _cant_split_row(row) -> None:
    from docx.oxml import OxmlElement

    tr = row._tr
    trPr = tr.get_or_add_trPr()
    trPr.append(OxmlElement("w:cantSplit"))


def _no_snap_to_grid(paragraph) -> None:
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    pPr = paragraph._p.get_or_add_pPr()
    snap = OxmlElement("w:snapToGrid")
    snap.set(qn("w:val"), "0")
    pPr.append(snap)


def _set_cell_width(cell, width) -> None:
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcW = tcPr.find(qn("w:tcW"))
    if tcW is None:
        tcW = OxmlElement("w:tcW")
        tcPr.append(tcW)
    twips = int(getattr(width, "twips", width))
    tcW.set(qn("w:w"), str(twips))
    tcW.set(qn("w:type"), "dxa")
    cell.width = width


def _shade_cell(cell, fill_hex: str) -> None:
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill_hex)
    shd.set(qn("w:val"), "clear")
    tcPr.append(shd)


def _cell_box(cell, *, left: str) -> None:
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    borders = OxmlElement("w:tcBorders")
    for edge in ("top", "left", "bottom", "right"):
        el = OxmlElement(f"w:{edge}")
        if edge == "left":
            el.set(qn("w:val"), "single")
            el.set(qn("w:sz"), "24")
            el.set(qn("w:color"), left)
        else:
            el.set(qn("w:val"), "single")
            el.set(qn("w:sz"), "4")
            el.set(qn("w:color"), LINE)
        borders.append(el)
    tcPr.append(borders)


def _p_border(paragraph, *, bottom: bool = False) -> None:
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    pPr = paragraph._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    if bottom:
        el = OxmlElement("w:bottom")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), "6")
        el.set(qn("w:space"), "3")
        el.set(qn("w:color"), LINE)
        pBdr.append(el)
    pPr.append(pBdr)
