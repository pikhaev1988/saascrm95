"""
BI-оформление DOCX информационно-статистического отчёта школы.

Только Presentation Layer поверх готового payload / info_stat_ui.
"""

from __future__ import annotations

from typing import Any


_C = {
    "ink": "102A43",
    "muted": "627D98",
    "bg": "F0F4F8",
    "blue": "2B6CB0",
    "blue_bg": "EBF8FF",
    "good": "2F855A",
    "good_bg": "F0FFF4",
    "mid": "B7791F",
    "mid_bg": "FFFBEB",
    "warn": "C05621",
    "warn_bg": "FFFAF0",
    "low": "C53030",
    "low_bg": "FFF5F5",
    "title": "1A365D",
}


def render_info_stat_docx(doc, ui: dict[str, Any]) -> None:
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Cm

    section = doc.sections[0]
    section.top_margin = Cm(1.4)
    section.bottom_margin = Cm(1.4)
    section.left_margin = Cm(1.5)
    section.right_margin = Cm(1.5)

    year = ui.get("year")
    year_bit = f" · {year}" if year else ""
    _banner(
        doc,
        f"ИНФОРМАЦИОННО-СТАТИСТИЧЕСКИЙ ОТЧЁТ ГИА  ·  {ui.get('exam_label')}{year_bit}",
        f"Дата формирования: {ui.get('generated_at') or '—'}",
    )

    avg_label = ui.get("avg_label") or "Средний балл"
    kpi = ui.get("kpi") or {}
    tones = kpi.get("tones") or {}

    # 01 Общие сведения
    _module(doc, "01", "Общие сведения")
    kpi_t = doc.add_table(rows=1, cols=4)
    items = [
        ("👥", str(kpi.get("participants") or 0), "Количество участников", tones.get("participants", "neutral")),
        ("📈", str(kpi.get("avg_score") if kpi.get("avg_score") is not None else "—"), avg_label, tones.get("avg", "mid")),
        ("🎯", f"{kpi.get('quality_rate')}%", "Качество знаний", tones.get("quality", "mid")),
        ("✅", f"{kpi.get('pass_rate')}%", "Успеваемость", tones.get("pass", "mid")),
    ]
    for i, (icon, value, label, tone) in enumerate(items):
        _kpi_cell(kpi_t.rows[0].cells[i], icon, value, label, tone)
    _nofit(kpi_t)
    doc.add_paragraph()

    extras = ui.get("extra") or []
    if extras:
        cols = 3
        rows_n = (len(extras) + cols - 1) // cols
        et = doc.add_table(rows=rows_n, cols=cols)
        for idx, card in enumerate(extras):
            r, c = divmod(idx, cols)
            cell = et.rows[r].cells[c]
            _shade(cell, _tone_bg(card.get("tone") or "neutral"))
            _pad(cell)
            p = cell.paragraphs[0]
            rr = p.add_run(f"{card.get('icon', '')}  {card.get('label', '')}")
            _font(rr, 8, True, _C["muted"])
            p2 = cell.add_paragraph()
            rr2 = p2.add_run(str(card.get("value") if card.get("value") is not None else "—"))
            _font(rr2, 14, True, _C["ink"])
        _nofit(et)

    # 02 Динамика и сравнение
    _module(doc, "02", "Динамика и сравнение")
    cmp_rows = ui.get("comparison") or []
    if cmp_rows:
        ct = doc.add_table(rows=1, cols=min(4, len(cmp_rows)))
        for i, card in enumerate(cmp_rows[:4]):
            cell = ct.rows[0].cells[i]
            _shade(cell, _tone_bg(card.get("tone") or "neutral"))
            _pad(cell)
            p = cell.paragraphs[0]
            r = p.add_run(f"{card.get('arrow', '')}  {card.get('label', '')}")
            _font(r, 9, True, _tone_fg(card.get("tone") or "neutral"))
            p2 = cell.add_paragraph()
            r2 = p2.add_run(str(card.get("value") or "—"))
            _font(r2, 14, True, _C["ink"])
            p3 = cell.add_paragraph()
            r3 = p3.add_run(str(card.get("text") or ""))
            _font(r3, 8, False, _C["muted"])
        _nofit(ct)

    # 03 Ключевые выводы
    insights = ui.get("insights") or []
    if insights:
        _module(doc, "03", "Ключевые выводы")
        for note in insights:
            tone = note.get("tone") or "mid"
            icon = {"high": "🟢", "warn": "🟡", "low": "🔴"}.get(tone, "🟡")
            box = doc.add_table(rows=1, cols=1)
            cell = box.rows[0].cells[0]
            _shade(cell, _tone_bg(tone))
            _pad(cell)
            p = cell.paragraphs[0]
            r = p.add_run(f"{icon}  {note.get('title', '')}")
            _font(r, 10, True, _tone_fg(tone))
            p2 = cell.add_paragraph()
            r2 = p2.add_run(str(note.get("text") or ""))
            _font(r2, 10, False, _C["ink"])
            _nofit(box)
            doc.add_paragraph()

    # 04 Распределение
    dist = ui.get("distribution") or []
    if dist:
        _module(doc, "04", "Распределение по уровням")
        hint = doc.add_paragraph()
        hr = hint.add_run(f"Всего результатов: {ui.get('distribution_total') or 0}")
        _font(hr, 9, False, _C["muted"])
        for idx, row in enumerate(dist, start=1):
            _hbar(
                doc,
                label=str(row.get("label") or ""),
                meta=f"{row.get('value')} · {row.get('percent')}%",
                pct=float(row.get("bar") or row.get("percent") or 0),
                color=_dist_color(idx),
            )

    # 05 Динамика прошлых лет
    dynamics = ui.get("dynamics") or []
    if dynamics:
        _module(doc, "05", "Динамика прошлых лет")
        for row in dynamics:
            _hbar(
                doc,
                label=f"{row.get('year')}  ·  n={row.get('results') or row.get('participants')}  ·  ср. {row.get('avg')}",
                meta=f"{row.get('pass_rate')}%",
                pct=float(row.get("pass_bar") or row.get("pass_rate") or 0),
                color=_C["blue"],
            )

    # 06 Предметы
    subjects = ui.get("subjects") or []
    if subjects:
        _module(doc, "06", "Результаты по предметам")
        for i in range(0, len(subjects), 2):
            pair = subjects[i : i + 2]
            t = doc.add_table(rows=1, cols=len(pair))
            for j, card in enumerate(pair):
                _subject_card(t.rows[0].cells[j], card, avg_label=avg_label)
            _nofit(t)
            doc.add_paragraph()

    # 07 Классы
    classes = ui.get("classes") or []
    if classes:
        _module(doc, "07", "Поклассная статистика")
        for row in classes:
            _hbar(
                doc,
                label=f"{row.get('name')}  ·  n={row.get('participants')}  ·  {row.get('avg')}",
                meta=f"{row.get('pass_rate')}%",
                pct=float(row.get("pass_bar") or 0),
                color=_C["blue"],
            )

    # 08 Проблемные зоны
    zones = ui.get("weak_zones") or []
    if zones:
        _module(doc, "08", "Проблемные зоны")
        for i in range(0, len(zones), 2):
            pair = zones[i : i + 2]
            t = doc.add_table(rows=1, cols=len(pair))
            for j, card in enumerate(pair):
                cell = t.rows[0].cells[j]
                tone = card.get("tone") or "warn"
                _shade(cell, _tone_bg(tone))
                _pad(cell)
                p = cell.paragraphs[0]
                r = p.add_run(str(card.get("name") or ""))
                _font(r, 11, True, _C["ink"])
                p2 = cell.add_paragraph()
                r2 = p2.add_run(
                    f"{avg_label}: {card.get('avg')}   ·   успеваемость {card.get('pass_rate')}%   ·   риск: {card.get('risk')}"
                )
                _font(r2, 9, False, _tone_fg(tone))
                p3 = cell.add_paragraph()
                r3 = p3.add_run(_bar(float(card.get("pass_bar") or 0)))
                _font(r3, 9, True, _tone_fg(tone))
                p4 = cell.add_paragraph()
                r4 = p4.add_run(str(card.get("description") or ""))
                _font(r4, 9, False, _C["ink"])
            _nofit(t)
            doc.add_paragraph()

    # 09 Рекомендации
    groups = ui.get("reco_groups") or []
    if groups:
        _module(doc, "09", "Рекомендации")
        for group in groups:
            box = doc.add_table(rows=1, cols=1)
            cell = box.rows[0].cells[0]
            _shade(cell, _C["blue_bg"])
            _pad(cell)
            p = cell.paragraphs[0]
            r = p.add_run(f"{group.get('icon', '')}  {group.get('title', '')}")
            _font(r, 11, True, _C["blue"])
            for item in group.get("items") or []:
                p2 = cell.add_paragraph()
                r2 = p2.add_run(f"• {item}")
                _font(r2, 9, False, _C["ink"])
            _nofit(box)
            doc.add_paragraph()

    foot = doc.add_paragraph()
    fr = foot.add_run("Показатели сформированы по фактическим результатам школы. Оформление — аналитическая панель отчёта.")
    _font(fr, 8, False, _C["muted"])


def _banner(doc, title: str, subtitle: str) -> None:
    t = doc.add_table(rows=1, cols=1)
    cell = t.rows[0].cells[0]
    _shade(cell, _C["title"])
    _pad(cell, 140)
    p = cell.paragraphs[0]
    r = p.add_run(title)
    _font(r, 13, True, "FFFFFF")
    p2 = cell.add_paragraph()
    r2 = p2.add_run(subtitle)
    _font(r2, 9, False, "BEE3F8")
    _nofit(t)
    doc.add_paragraph()


def _module(doc, num: str, title: str) -> None:
    doc.add_paragraph()
    t = doc.add_table(rows=1, cols=1)
    cell = t.rows[0].cells[0]
    _shade(cell, _C["bg"])
    _pad(cell, 70)
    p = cell.paragraphs[0]
    r = p.add_run(f"{num}   {title.upper()}")
    _font(r, 11, True, _C["title"])
    _nofit(t)
    doc.add_paragraph()


def _kpi_cell(cell, icon: str, value: str, label: str, tone: str) -> None:
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    _shade(cell, _tone_bg(tone))
    _pad(cell)
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(icon)
    _font(r, 12, False, _tone_fg(tone))
    p2 = cell.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = p2.add_run(value)
    _font(r2, 16, True, _C["ink"])
    p3 = cell.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r3 = p3.add_run(label)
    _font(r3, 8, False, _C["muted"])


def _subject_card(cell, card: dict, *, avg_label: str) -> None:
    tone = card.get("tone") or "mid"
    _shade(cell, _tone_bg(tone))
    _pad(cell)
    p = cell.paragraphs[0]
    r = p.add_run(str(card.get("name") or "Предмет"))
    _font(r, 11, True, _C["ink"])
    p2 = cell.add_paragraph()
    r2 = p2.add_run(str(card.get("status") or ""))
    _font(r2, 9, True, _tone_fg(tone))
    p3 = cell.add_paragraph()
    r3 = p3.add_run(
        f"Результатов: {card.get('participants')}  ·  {avg_label}: {card.get('avg')}  ·  "
        f"успеваемость {card.get('pass_rate')}%  ·  качество {card.get('quality_rate')}%"
    )
    _font(r3, 8, False, _C["muted"])
    p4 = cell.add_paragraph()
    r4 = p4.add_run(_bar(float(card.get("pass_bar") or 0)))
    _font(r4, 9, True, _tone_fg(tone))


def _hbar(doc, *, label: str, meta: str, pct: float, color: str) -> None:
    t = doc.add_table(rows=1, cols=3)
    c0, c1, c2 = t.rows[0].cells
    _pad(c0, 40)
    _pad(c1, 40)
    _pad(c2, 40)
    p0 = c0.paragraphs[0]
    r0 = p0.add_run(label)
    _font(r0, 9, True, _C["ink"])
    _shade(c1, "E2E8F0")
    p1 = c1.paragraphs[0]
    r1 = p1.add_run(_bar(pct))
    _font(r1, 9, True, color)
    p2 = c2.paragraphs[0]
    r2 = p2.add_run(meta)
    _font(r2, 9, False, _C["muted"])
    _nofit(t)


def _bar(pct: float, width: int = 12) -> str:
    pct = max(0.0, min(100.0, float(pct or 0)))
    filled = int(round(width * pct / 100.0))
    return "█" * filled + "░" * (width - filled)


def _dist_color(idx: int) -> str:
    return {1: "805AD5", 2: _C["good"], 3: _C["blue"], 4: _C["mid"], 5: _C["warn"], 6: _C["low"]}.get(idx, _C["blue"])


def _tone_bg(tone: str) -> str:
    return {
        "high": _C["good_bg"],
        "mid": _C["mid_bg"],
        "warn": _C["warn_bg"],
        "low": _C["low_bg"],
        "neutral": _C["blue_bg"],
    }.get(tone or "neutral", _C["blue_bg"])


def _tone_fg(tone: str) -> str:
    return {
        "high": _C["good"],
        "mid": _C["mid"],
        "warn": _C["warn"],
        "low": _C["low"],
        "neutral": _C["blue"],
    }.get(tone or "neutral", _C["blue"])


def _shade(cell, fill_hex: str) -> None:
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill_hex)
    shd.set(qn("w:val"), "clear")
    tc_pr.append(shd)


def _pad(cell, v: int = 90) -> None:
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = OxmlElement("w:tcMar")
    for name in ("top", "bottom", "left", "right"):
        node = OxmlElement(f"w:{name}")
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")
        tc_mar.append(node)
    tc_pr.append(tc_mar)


def _font(run, size: int = 10, bold: bool = False, color: str = "102A43") -> None:
    from docx.shared import Pt, RGBColor

    run.bold = bold
    run.font.size = Pt(size)
    run.font.name = "Calibri"
    color = (color or "102A43").lstrip("#")
    run.font.color.rgb = RGBColor(int(color[0:2], 16), int(color[2:4], 16), int(color[4:6], 16))


def _nofit(table) -> None:
    try:
        table.autofit = True
    except Exception:
        pass
    for style_name in ("Table Grid", "Normal Table"):
        try:
            table.style = style_name
            break
        except (KeyError, ValueError):
            continue
