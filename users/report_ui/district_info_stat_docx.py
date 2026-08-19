"""
DOCX информационно-статистического отчёта по муниципалитету.

Официальный статистический бюллетень: максимум таблиц и KPI, минимум текста.
Только Presentation Layer поверх готового payload.
"""

from __future__ import annotations

from io import BytesIO
from typing import Any

from users.report_ui.district_info_stat import DOC_VERSION, build_district_info_stat_presentation

NAVY = "0B1F3A"
BLUE = "1B4F72"
SKY = "5B9BD5"
GRAY = "5D6D7E"
LINE = "D5DDE5"
SOFT = "F4F7FA"
WHITE = "FFFFFF"
GOOD = "1E7A4A"
GOOD_BG = "EAF6EF"
WARN = "A85A1A"
WARN_BG = "F8F1E6"
RISK = "9B2C2C"
RISK_BG = "F8ECEC"
INFO_BG = "EAF2F8"
TOP_BG = "E8F5EE"
LOW_BG = "F9EDED"

TONE_HEX = {
    "sky": SKY,
    "good": GOOD,
    "warn": WARN,
    "risk": RISK,
    "neutral": SKY,
}


def render_district_info_stat_docx(data: dict[str, Any]) -> BytesIO:
    from docx import Document

    doc = Document()
    _setup_page(doc)

    ui = build_district_info_stat_presentation(data)
    if not ui.get("has_data"):
        doc.add_paragraph(ui.get("message") or "Недостаточно данных.")
        return _save(doc)

    district = ui["district"]
    exam = ui["exam_label"]
    year = ui.get("year") or ""
    formed = ui.get("generated_at") or ""
    avg_label = ui.get("avg_label") or "Средний балл"
    notes = ui.get("notes") or {}

    _setup_header_footer(doc, district=district, year=year, formed=formed)
    _title_page(doc, district=district, exam=exam, year=year, formed=formed)
    _toc(doc)

    # 1. Паспорт
    _section(doc, "1. Паспорт муниципалитета")
    _note(doc, notes.get("passport", ""))
    _kpi_dashboard(doc, ui.get("passport") or [])

    # 2. Общая статистика
    _section(doc, "2. Общая статистика")
    _note(doc, notes.get("general", ""))
    kpi = ui.get("kpi_summary") or {}
    _data_table(
        doc,
        ["Показатель", "Значение"],
        [
            ["Участники", str(kpi.get("participants", ""))],
            ["Результаты (записи)", str(kpi.get("total", ""))],
            ["Образовательные организации", str(kpi.get("schools_count", ""))],
            ["Предметы", str(kpi.get("subjects_count", ""))],
            [avg_label, _s(kpi.get("avg_score"))],
            ["Качество знаний, %", _s(kpi.get("quality_rate"))],
            ["Успеваемость, %", _s(kpi.get("pass_rate"))],
            ["Высокобалльники", str(kpi.get("high_count", ""))],
            ["Неудовлетворительные", str(kpi.get("failed_count", ""))],
        ],
        widths_cm=[10.0, 6.8],
        num_cols={1},
    )

    # 3. Предметная статистика
    _section(doc, "3. Предметная статистика")
    _note(doc, notes.get("subjects", ""))
    subj_rows = []
    for row in ui.get("subjects") or []:
        subj_rows.append(
            [
                row["name"],
                str(row["participants"]),
                _s(row["avg"]),
                _s(row["quality_rate"]),
                _s(row["pass_rate"]),
                str(row["high"]),
                _fmt_delta(row.get("mun_delta")),
                _fmt_delta(row.get("rep_delta")),
            ]
        )
    _data_table(
        doc,
        [
            "Предмет",
            "Участ.",
            "Средний",
            "Кач., %",
            "Усп., %",
            "Высок.",
            "Δ мун.",
            "Δ респ.",
        ],
        subj_rows,
        widths_cm=[4.0, 1.4, 1.6, 1.5, 1.5, 1.4, 1.6, 1.6],
        num_cols={1, 2, 3, 4, 5, 6, 7},
        highlight_numeric_extremes={2},
    )

    # 4. Рейтинг предметов
    _section(doc, "4. Рейтинг предметов")
    _caption(doc, "ТОП предметов по среднему результату")
    _rank_table(
        doc,
        ui.get("rank_avg") or [],
        value_key="avg",
        value_header=avg_label,
        widths_cm=[1.2, 10.4, 5.2],
    )
    _caption(doc, "ТОП предметов по качеству знаний")
    _rank_table(
        doc,
        ui.get("rank_quality") or [],
        value_key="quality_rate",
        value_header="Качество, %",
        widths_cm=[1.2, 10.4, 5.2],
        suffix="%",
    )
    _caption(doc, "ТОП предметов по массовости")
    _rank_table(
        doc,
        ui.get("rank_mass") or [],
        value_key="participants",
        value_header="Участники",
        widths_cm=[1.2, 10.4, 5.2],
    )

    # 5. Статистика ОО
    _section(doc, "5. Статистика образовательных организаций")
    _note(doc, notes.get("schools", ""))
    school_rows = []
    pass_tones = []
    for i, row in enumerate(sorted(ui.get("schools") or [], key=lambda x: (-float(x["avg"]), x["name"])), start=1):
        school_rows.append(
            [
                str(i),
                row["name"],
                row["code"],
                str(row["participants"]),
                _s(row["avg"]),
                _s(row["pass_rate"]),
            ]
        )
        pass_tones.append(float(row["pass_rate"]))
    _data_table(
        doc,
        ["№", "ОО", "Код", "Участ.", "Средний", "Усп., %"],
        school_rows,
        widths_cm=[0.9, 8.2, 1.8, 1.5, 1.8, 1.8],
        num_cols={0, 3, 4, 5},
        pass_rate_col=5,
        pass_tones=pass_tones,
        highlight_numeric_extremes={4},
    )

    # 6. ТОП-10
    _section(doc, "6. ТОП-10 школ по среднему результату")
    _top_bottom_table(doc, ui.get("top10") or [], avg_label=avg_label, mode="top")

    # 7. Bottom-10
    _section(doc, "7. 10 школ с наименьшими результатами")
    _top_bottom_table(doc, ui.get("bottom10") or [], avg_label=avg_label, mode="low")

    # 8. Структура результатов
    _section(doc, "8. Структура результатов")
    _note(doc, notes.get("distribution", ""))
    _distribution_block(doc, ui.get("distribution") or [], total=ui.get("distribution_total") or 0)

    # 9. Сравнение
    _section(doc, "9. Сравнение: муниципалитет и республика")
    _note(doc, notes.get("comparison", ""))
    cmp_rows = []
    cmp_tones = []
    for row in ui.get("comparison") or []:
        cmp_rows.append(
            [
                row["indicator"],
                row["municipality"],
                row["republic"],
                row["delta"],
            ]
        )
        cmp_tones.append(row.get("tone") or "neutral")
    _data_table(
        doc,
        ["Показатель", "Муниципалитет", "Республика", "Разница"],
        cmp_rows,
        widths_cm=[5.5, 3.8, 3.8, 3.7],
        num_cols={1, 2, 3},
        delta_col=3,
        delta_tones=cmp_tones,
    )

    # 10. Динамика
    _section(doc, "10. Динамика")
    _note(doc, notes.get("dynamics", ""))
    dyn_rows = [
        [r["indicator"], r["current"], r["previous"], r["delta"], r["rate"]]
        for r in (ui.get("dynamics") or [])
    ]
    _data_table(
        doc,
        ["Показатель", "Текущий год", "Прошлый год", "Разница", "Темп, %"],
        dyn_rows,
        widths_cm=[4.6, 3.0, 3.0, 3.0, 3.2],
        num_cols={1, 2, 3, 4},
    )

    # 11. Сводные таблицы
    _section(doc, "11. Сводные статистические таблицы")
    _caption(doc, "Сводка ключевых показателей")
    _data_table(
        doc,
        ["Блок", "Показатель", "Значение"],
        [
            ["Контингент", "Участники", str(kpi.get("participants", ""))],
            ["Контингент", "ОО", str(kpi.get("schools_count", ""))],
            ["Контингент", "Предметы", str(kpi.get("subjects_count", ""))],
            ["Результативность", avg_label, _s(kpi.get("avg_score"))],
            ["Результативность", "Качество знаний, %", _s(kpi.get("quality_rate"))],
            ["Результативность", "Успеваемость, %", _s(kpi.get("pass_rate"))],
            ["Экстремумы", "Высокобалльники", str(kpi.get("high_count", ""))],
            ["Экстремумы", "Неудовлетворительные", str(kpi.get("failed_count", ""))],
        ],
        widths_cm=[4.5, 7.5, 4.8],
        num_cols={2},
    )

    # 12. Основные выводы
    _section(doc, "12. Основные выводы")
    for i, line in enumerate(ui.get("conclusions") or [], start=1):
        _conclusion_item(doc, i, line)

    return _save(doc)


# ── page / chrome ─────────────────────────────────────────────


def _save(doc) -> BytesIO:
    out = BytesIO()
    doc.save(out)
    out.seek(0)
    return out


def _setup_page(doc) -> None:
    from docx.oxml.ns import qn
    from docx.shared import Cm, Pt

    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.left_margin = Cm(1.6)
    section.right_margin = Cm(1.6)
    section.top_margin = Cm(1.7)
    section.bottom_margin = Cm(1.7)
    section.header_distance = Cm(0.8)
    section.footer_distance = Cm(0.8)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(10)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
    pf = normal.paragraph_format
    pf.space_after = Pt(4)
    pf.space_before = Pt(0)
    pf.line_spacing = 1.12


def _setup_header_footer(doc, *, district: str, year, formed: str) -> None:
    from docx.shared import Pt, RGBColor

    section = doc.sections[0]
    header = section.header
    header.is_linked_to_previous = False
    hp = header.paragraphs[0]
    hp.clear()
    r = hp.add_run(
        f"Информационно-статистический отчёт по муниципалитету  ·  {district}  ·  {year or '—'}"
    )
    r.font.size = Pt(8)
    r.font.color.rgb = RGBColor(0x5D, 0x6D, 0x7E)
    r.font.name = "Calibri"
    _p_border(hp, bottom=True)

    footer = section.footer
    footer.is_linked_to_previous = False
    fp = footer.paragraphs[0]
    fp.clear()
    run = fp.add_run(f"Дата формирования: {formed}  ·  версия {DOC_VERSION}  ·  стр. ")
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(0x5D, 0x6D, 0x7E)
    _add_page_field(fp)


def _add_page_field(paragraph) -> None:
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.shared import Pt, RGBColor

    fld = OxmlElement("w:fldChar")
    fld.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    fld2 = OxmlElement("w:fldChar")
    fld2.set(qn("w:fldCharType"), "end")
    r = paragraph.add_run()
    r._r.append(fld)
    r2 = paragraph.add_run()
    r2._r.append(instr)
    r3 = paragraph.add_run()
    r3._r.append(fld2)
    for rr in (r, r2, r3):
        rr.font.size = Pt(8)
        rr.font.color.rgb = RGBColor(0x5D, 0x6D, 0x7E)


def _title_page(doc, *, district: str, exam: str, year, formed: str) -> None:
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Cm, Pt, RGBColor

    logo = doc.add_table(rows=1, cols=2)
    logo.autofit = False
    c0, c1 = logo.rows[0].cells
    _set_cell_width(c0, Cm(1.4))
    _set_cell_width(c1, Cm(15))
    _shade_cell(c0, NAVY)
    p0 = c0.paragraphs[0]
    p0.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r0 = p0.add_run("ИС")
    r0.bold = True
    r0.font.size = Pt(14)
    r0.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    p1 = c1.paragraphs[0]
    r1 = p1.add_run("УПРАВЛЕНИЕ ОБРАЗОВАНИЯ")
    r1.bold = True
    r1.font.size = Pt(10)
    r1.font.color.rgb = RGBColor(0x1B, 0x4F, 0x72)
    p1b = c1.add_paragraph()
    r1b = p1b.add_run(district)
    r1b.font.size = Pt(12)
    r1b.bold = True
    r1b.font.color.rgb = RGBColor(0x0B, 0x1F, 0x3A)

    for _ in range(4):
        doc.add_paragraph("")

    eye = doc.add_paragraph()
    er = eye.add_run("ОФИЦИАЛЬНЫЙ ИНФОРМАЦИОННО-СТАТИСТИЧЕСКИЙ ОТЧЁТ")
    er.bold = True
    er.font.size = Pt(9)
    er.font.color.rgb = RGBColor(0x5B, 0x9B, 0xD5)

    title = doc.add_paragraph()
    title.paragraph_format.space_before = Pt(8)
    tr = title.add_run("Информационно-статистический\nотчёт по муниципалитету")
    tr.bold = True
    tr.font.size = Pt(26)
    tr.font.color.rgb = RGBColor(0x0B, 0x1F, 0x3A)
    tr.font.name = "Calibri"

    sub = doc.add_paragraph()
    sr = sub.add_run(f"Статистический бюллетень результатов {exam}")
    sr.font.size = Pt(12)
    sr.font.color.rgb = RGBColor(0x5D, 0x6D, 0x7E)

    for _ in range(2):
        doc.add_paragraph("")

    meta = doc.add_table(rows=5, cols=2)
    meta.autofit = False
    pairs = [
        ("Вид экзамена", exam),
        ("Муниципалитет", district),
        ("Отчётный период", str(year or "—")),
        ("Дата формирования", formed),
        ("Версия документа", DOC_VERSION),
    ]
    for i, (k, v) in enumerate(pairs):
        ck, cv = meta.rows[i].cells
        _set_cell_width(ck, Cm(4.5))
        _set_cell_width(cv, Cm(11.5))
        pk = ck.paragraphs[0]
        rk = pk.add_run(k)
        rk.font.size = Pt(9)
        rk.font.color.rgb = RGBColor(0x5D, 0x6D, 0x7E)
        pv = cv.paragraphs[0]
        rv = pv.add_run(v)
        rv.bold = True
        rv.font.size = Pt(12)
        rv.font.color.rgb = RGBColor(0x0B, 0x1F, 0x3A)

    for _ in range(3):
        doc.add_paragraph("")

    conf = doc.add_paragraph()
    cr = conf.add_run("Для служебного пользования · статистический материал муниципального уровня")
    cr.font.size = Pt(9)
    cr.font.color.rgb = RGBColor(0x5D, 0x6D, 0x7E)

    doc.add_page_break()


def _toc(doc) -> None:
    from docx.shared import Pt, RGBColor

    h = doc.add_paragraph()
    r = h.add_run("Содержание")
    r.bold = True
    r.font.size = Pt(16)
    r.font.color.rgb = RGBColor(0x0B, 0x1F, 0x3A)

    items = [
        "1. Паспорт муниципалитета",
        "2. Общая статистика",
        "3. Предметная статистика",
        "4. Рейтинг предметов",
        "5. Статистика образовательных организаций",
        "6. ТОП-10 школ по среднему результату",
        "7. 10 школ с наименьшими результатами",
        "8. Структура результатов",
        "9. Сравнение: муниципалитет и республика",
        "10. Динамика",
        "11. Сводные статистические таблицы",
        "12. Основные выводы",
    ]
    for item in items:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(4)
        run = p.add_run(item)
        run.font.size = Pt(11)
        run.font.color.rgb = RGBColor(0x1A, 0x23, 0x32)

    doc.add_page_break()


# ── sections / notes ──────────────────────────────────────────


def _section(doc, title: str) -> None:
    from docx.oxml import OxmlElement
    from docx.shared import Cm, Pt, RGBColor

    wrap = doc.add_table(rows=1, cols=2)
    wrap.autofit = False
    a, b = wrap.rows[0].cells
    _set_cell_width(a, Cm(0.25))
    _set_cell_width(b, Cm(16.5))
    _shade_cell(a, SKY)
    p = b.paragraphs[0]
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(4)
    pPr = p._p.get_or_add_pPr()
    pPr.append(OxmlElement("w:keepNext"))
    tr = p.add_run(title)
    tr.bold = True
    tr.font.size = Pt(13)
    tr.font.color.rgb = RGBColor(0x0B, 0x1F, 0x3A)
    _p_border(p, bottom=True)


def _caption(doc, text: str) -> None:
    from docx.shared import Pt, RGBColor

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(10)
    r.font.color.rgb = RGBColor(0x1B, 0x4F, 0x72)


def _note(doc, text: str) -> None:
    if not (text or "").strip():
        return
    from docx.shared import Cm, Pt, RGBColor

    t = doc.add_table(rows=1, cols=2)
    t.autofit = False
    a, b = t.rows[0].cells
    _set_cell_width(a, Cm(0.18))
    _set_cell_width(b, Cm(16.5))
    _shade_cell(a, SKY)
    _shade_cell(b, INFO_BG)
    p = b.paragraphs[0]
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    badge = p.add_run("Справка  ")
    badge.bold = True
    badge.font.size = Pt(7)
    badge.font.color.rgb = RGBColor(0x1B, 0x4F, 0x72)
    run = p.add_run(text)
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x1A, 0x23, 0x32)
    doc.add_paragraph("").paragraph_format.space_after = Pt(2)


def _conclusion_item(doc, num: int, text: str) -> None:
    if not (text or "").strip():
        return
    from docx.shared import Cm, Pt, RGBColor

    t = doc.add_table(rows=1, cols=2)
    t.autofit = False
    n, b = t.rows[0].cells
    _set_cell_width(n, Cm(0.9))
    _set_cell_width(b, Cm(15.9))
    _shade_cell(n, BLUE)
    _shade_cell(b, SOFT)
    pn = n.paragraphs[0]
    rn = pn.add_run(f"{num:02d}")
    rn.bold = True
    rn.font.size = Pt(11)
    rn.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    p = b.paragraphs[0]
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run(text)
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x1A, 0x23, 0x32)


# ── KPI ───────────────────────────────────────────────────────


def _kpi_dashboard(doc, cards: list[dict]) -> None:
    from docx.shared import Cm, Pt, RGBColor

    for row_i in range(0, len(cards), 3):
        chunk = cards[row_i : row_i + 3]
        while len(chunk) < 3:
            chunk.append({"icon": "", "label": "", "value": "", "tone": "sky"})
        table = doc.add_table(rows=1, cols=3)
        table.autofit = False
        for col, card in enumerate(chunk):
            cell = table.rows[0].cells[col]
            _set_cell_width(cell, Cm(5.6))
            _shade_cell(cell, SOFT)
            tone = TONE_HEX.get(card.get("tone") or "sky", SKY)
            _cell_left_border(cell, tone)
            p_icon = cell.paragraphs[0]
            p_icon.paragraph_format.space_after = Pt(0)
            p_icon.paragraph_format.space_before = Pt(4)
            if card.get("icon"):
                ri = p_icon.add_run(str(card["icon"]))
                ri.font.size = Pt(9)
                ri.font.color.rgb = RGBColor.from_string(tone)
            p_val = cell.add_paragraph()
            p_val.paragraph_format.space_before = Pt(2)
            p_val.paragraph_format.space_after = Pt(0)
            rv = p_val.add_run(str(card.get("value") or ""))
            rv.bold = True
            rv.font.size = Pt(15)
            rv.font.color.rgb = RGBColor(0x0B, 0x1F, 0x3A)
            p_lab = cell.add_paragraph()
            p_lab.paragraph_format.space_before = Pt(1)
            p_lab.paragraph_format.space_after = Pt(4)
            rl = p_lab.add_run(str(card.get("label") or ""))
            rl.font.size = Pt(8)
            rl.font.color.rgb = RGBColor(0x5D, 0x6D, 0x7E)
        doc.add_paragraph("").paragraph_format.space_after = Pt(2)


# ── tables ────────────────────────────────────────────────────


def _rank_table(doc, rows: list[dict], *, value_key: str, value_header: str, widths_cm: list[float], suffix: str = "") -> None:
    data = []
    for row in rows:
        val = row.get(value_key)
        data.append([str(row.get("rank") or ""), row.get("name") or "—", f"{_s(val)}{suffix}"])
    _data_table(
        doc,
        ["№", "Предмет", value_header],
        data,
        widths_cm=widths_cm,
        num_cols={0, 2},
        top_n=3,
        low_n=0 if value_key == "participants" else 1,
    )


def _top_bottom_table(doc, rows: list[dict], *, avg_label: str, mode: str) -> None:
    data = []
    for row in rows:
        data.append(
            [
                str(row.get("rank") or ""),
                row.get("name") or "—",
                row.get("code") or "—",
                str(row.get("participants") or ""),
                _s(row.get("avg")),
                _s(row.get("pass_rate")),
            ]
        )
    _data_table(
        doc,
        ["№", "ОО", "Код", "Участ.", avg_label, "Усп., %"],
        data,
        widths_cm=[0.9, 8.2, 1.8, 1.5, 2.2, 1.8],
        num_cols={0, 3, 4, 5},
        force_row_fill=TOP_BG if mode == "top" else LOW_BG,
    )


def _distribution_block(doc, rows: list[dict], *, total: int) -> None:
    from docx.shared import Cm, Pt, RGBColor

    if not rows:
        p = doc.add_paragraph()
        r = p.add_run("Данные распределения отсутствуют.")
        r.font.size = Pt(9)
        r.font.color.rgb = RGBColor(0x5D, 0x6D, 0x7E)
        return

    hint = doc.add_paragraph()
    hr = hint.add_run(f"Всего результатов в распределении: {total}")
    hr.font.size = Pt(9)
    hr.font.color.rgb = RGBColor(0x5D, 0x6D, 0x7E)

    table_rows = []
    for row in rows:
        bar = _bar(float(row.get("bar") or 0))
        table_rows.append(
            [
                str(row.get("label") or ""),
                str(row.get("value") or 0),
                f"{_s(row.get('percent'))}%",
                bar,
            ]
        )
    _data_table(
        doc,
        ["Диапазон", "Численность", "Доля", "Инфографика"],
        table_rows,
        widths_cm=[3.2, 2.6, 2.2, 8.8],
        num_cols={1, 2},
    )

    # compact visual bars as separate cards
    for idx, row in enumerate(rows):
        t = doc.add_table(rows=1, cols=2)
        t.autofit = False
        a, b = t.rows[0].cells
        _set_cell_width(a, Cm(3.5))
        _set_cell_width(b, Cm(13.3))
        color = [NAVY, BLUE, SKY, GOOD, WARN][idx % 5]
        _shade_cell(a, color)
        pa = a.paragraphs[0]
        ra = pa.add_run(str(row.get("label") or ""))
        ra.bold = True
        ra.font.size = Pt(9)
        ra.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        _shade_cell(b, SOFT)
        pb = b.paragraphs[0]
        rb = pb.add_run(f"{row.get('value')}  ·  {_s(row.get('percent'))}%   {_bar(float(row.get('bar') or 0))}")
        rb.font.size = Pt(9)
        rb.font.color.rgb = RGBColor(0x1A, 0x23, 0x32)
    doc.add_paragraph("").paragraph_format.space_after = Pt(4)


def _data_table(
    doc,
    headers: list[str],
    rows: list[list[str]],
    *,
    widths_cm: list[float],
    num_cols: set[int] | None = None,
    pass_rate_col: int | None = None,
    pass_tones: list[float] | None = None,
    delta_col: int | None = None,
    delta_tones: list[str] | None = None,
    highlight_numeric_extremes: set[int] | None = None,
    top_n: int = 0,
    low_n: int = 0,
    force_row_fill: str | None = None,
) -> None:
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Cm, Pt, RGBColor

    num_cols = num_cols or set()
    highlight_numeric_extremes = highlight_numeric_extremes or set()
    table = doc.add_table(rows=1, cols=len(headers))
    table.autofit = False
    table.allow_autofit = False

    hdr = table.rows[0]
    _set_row_header(hdr)
    for i, h in enumerate(headers):
        cell = hdr.cells[i]
        _set_cell_width(cell, Cm(widths_cm[i]))
        _shade_cell(cell, NAVY)
        p = cell.paragraphs[0]
        p.paragraph_format.space_before = Pt(3)
        p.paragraph_format.space_after = Pt(3)
        if i in num_cols:
            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        run = p.add_run(h)
        run.bold = True
        run.font.size = Pt(8)
        run.font.color.rgb = RGBColor(0xDC, 0xE8, 0xF3)
        run.font.name = "Calibri"

    extreme_hi: dict[int, float] = {}
    extreme_lo: dict[int, float] = {}
    for col in highlight_numeric_extremes:
        vals = []
        for row in rows:
            try:
                vals.append(float(str(row[col]).replace("%", "").replace("+", "").replace(",", ".")))
            except (TypeError, ValueError, IndexError):
                continue
        if vals:
            extreme_hi[col] = max(vals)
            extreme_lo[col] = min(vals)

    for r_i, row in enumerate(rows):
        cells = table.add_row().cells
        _cant_split_row(table.rows[r_i + 1])
        row_fill = None
        if force_row_fill:
            row_fill = force_row_fill
        elif top_n and r_i < top_n:
            row_fill = TOP_BG
        elif low_n and r_i >= max(0, len(rows) - low_n):
            row_fill = LOW_BG
        elif r_i % 2 == 1:
            row_fill = "F7FAFC"

        for i, val in enumerate(row):
            cell = cells[i]
            _set_cell_width(cell, Cm(widths_cm[i]))
            fill = row_fill
            if pass_rate_col is not None and i == pass_rate_col and pass_tones and r_i < len(pass_tones):
                pr = pass_tones[r_i]
                if pr >= 85:
                    fill = GOOD_BG
                elif pr < 70:
                    fill = RISK_BG
                elif pr < 85:
                    fill = WARN_BG
            if delta_col is not None and i == delta_col and delta_tones and r_i < len(delta_tones):
                tone = delta_tones[r_i]
                if tone == "good":
                    fill = GOOD_BG
                elif tone == "risk":
                    fill = RISK_BG
            if i in extreme_hi:
                try:
                    fv = float(str(val).replace("%", "").replace("+", "").replace(",", "."))
                    if fv == extreme_hi[i]:
                        fill = TOP_BG
                    elif fv == extreme_lo[i]:
                        fill = LOW_BG
                except (TypeError, ValueError):
                    pass
            if fill:
                _shade_cell(cell, fill)
            p = cell.paragraphs[0]
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after = Pt(2)
            if i in num_cols:
                p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            run = p.add_run(str(val))
            run.font.size = Pt(8)
            run.font.name = "Calibri"
            run.font.color.rgb = RGBColor(0x1A, 0x23, 0x32)
            _no_snap_to_grid(p)

    doc.add_paragraph("").paragraph_format.space_after = Pt(6)


# ── helpers ───────────────────────────────────────────────────


def _s(v) -> str:
    if v is None:
        return "—"
    try:
        f = float(v)
    except (TypeError, ValueError):
        return str(v)
    if abs(f - round(f)) < 1e-9:
        return str(int(round(f)))
    return f"{f:.2f}".rstrip("0").rstrip(".")


def _fmt_delta(v) -> str:
    if v is None:
        return "—"
    try:
        f = float(v)
    except (TypeError, ValueError):
        return "—"
    if f > 0:
        return f"+{_s(f)}"
    return _s(f)


def _bar(pct: float) -> str:
    n = max(0, min(20, int(round(float(pct or 0) / 5))))
    return "█" * n + "░" * (20 - n)


def _set_row_header(row) -> None:
    from docx.oxml import OxmlElement

    tr = row._tr
    trPr = tr.get_or_add_trPr()
    trPr.append(OxmlElement("w:tblHeader"))


def _cant_split_row(row) -> None:
    from docx.oxml import OxmlElement

    tr = row._tr
    trPr = tr.get_or_add_trPr()
    trPr.append(OxmlElement("w:cantSplit"))


def _no_snap_to_grid(paragraph) -> None:
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    pPr = paragraph._p.get_or_add_pPr()
    snap = OxmlElement("w:snapToGrid")
    snap.set(qn("w:val"), "0")
    pPr.append(snap)


def _set_cell_width(cell, width) -> None:
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcW = tcPr.find(qn("w:tcW"))
    if tcW is None:
        tcW = OxmlElement("w:tcW")
        tcPr.append(tcW)
    twips = int(getattr(width, "twips", width))
    tcW.set(qn("w:w"), str(twips))
    tcW.set(qn("w:type"), "dxa")
    cell.width = width


def _cell_left_border(cell, color_hex: str) -> None:
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for edge in ("top", "left", "bottom", "right"):
        el = OxmlElement(f"w:{edge}")
        if edge == "left":
            el.set(qn("w:val"), "single")
            el.set(qn("w:sz"), "24")
            el.set(qn("w:color"), color_hex)
        else:
            el.set(qn("w:val"), "nil")
        tcBorders.append(el)
    tcPr.append(tcBorders)


def _p_border(paragraph, *, bottom: bool = False) -> None:
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    pPr = paragraph._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    if bottom:
        el = OxmlElement("w:bottom")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), "6")
        el.set(qn("w:space"), "4")
        el.set(qn("w:color"), LINE)
        pBdr.append(el)
    pPr.append(pBdr)


def _shade_cell(cell, fill_hex: str) -> None:
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill_hex)
    shd.set(qn("w:val"), "clear")
    tcPr.append(shd)
