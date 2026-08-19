"""
Сборка данных и выгрузка отчётов районного (муниципального) кабинета.
"""

from __future__ import annotations

from datetime import date
from io import BytesIO

from django.db.models import Avg, Count, Max, Min, Q

from exams.models import EgePassingThreshold, ExamResult, TaskResult
from exams.passing import is_gve_exam
from organizations.models import District
from analytics.engine.attempts import filter_latest_exam_results, task_results_for_exam_results
from users.district_gigachat_analysis import enrich_district_report_with_ai


def _export_helpers():
    from users.export_reports import (
        MO_SUBJECT_GROUPS,
        _build_mo_subject_filter,
        _is_weak_subject_row,
        _risk_level_ru_from_rates,
        _style_docx_table,
        _topic_for_task,
    )

    return {
        "MO_SUBJECT_GROUPS": MO_SUBJECT_GROUPS,
        "_build_mo_subject_filter": _build_mo_subject_filter,
        "_is_weak_subject_row": _is_weak_subject_row,
        "_risk_level_ru_from_rates": _risk_level_ru_from_rates,
        "_style_docx_table": _style_docx_table,
        "_topic_for_task": _topic_for_task,
    }


def _is_weak_subject_row(*args, **kwargs):
    return _export_helpers()["_is_weak_subject_row"](*args, **kwargs)


def _risk_level_ru_from_rates(*args, **kwargs):
    return _export_helpers()["_risk_level_ru_from_rates"](*args, **kwargs)


def _topic_for_task(*args, **kwargs):
    return _export_helpers()["_topic_for_task"](*args, **kwargs)


def _style_docx_table(*args, **kwargs):
    return _export_helpers()["_style_docx_table"](*args, **kwargs)


def _build_mo_subject_filter(*args, **kwargs):
    return _export_helpers()["_build_mo_subject_filter"](*args, **kwargs)


def _threshold_subject_key(subject_name: str) -> str | None:
    title = (subject_name or "").strip().lower()
    if "рус" in title:
        return "russian"
    if "математика" in title and "проф" in title:
        return "math_profile"
    if "математика" in title and "баз" in title:
        return "math_basic"
    if "обществ" in title:
        return "social"
    if "информат" in title:
        return "informatics"
    if "физик" in title:
        return "physics"
    if "хими" in title:
        return "chemistry"
    if "биолог" in title:
        return "biology"
    if "истори" in title:
        return "history"
    if "литератур" in title:
        return "literature"
    if "географ" in title:
        return "geography"
    if any(lang in title for lang in ("англий", "немец", "француз", "испан", "китай", "иностран")):
        return "foreign_language"
    return None


def _oge_grade_passed(score, passed_flag) -> bool:
    score_value = float(score or 0)
    if 0 < score_value <= 5:
        return score_value >= 3
    return bool(passed_flag)


def _is_passed_ege_row(row: dict, threshold_cache: dict) -> bool:
    if is_gve_exam(exam_code=row.get("exam__code"), subject_name=row.get("exam__subject")):
        return _oge_grade_passed(row.get("score"), row.get("passed"))
    subject_key = _threshold_subject_key(row.get("exam__subject") or "")
    if not subject_key:
        return bool(row.get("passed"))
    threshold_key = (int(row.get("exam__year") or 0), subject_key)
    if threshold_key not in threshold_cache:
        threshold_cache[threshold_key] = (
            EgePassingThreshold.objects.filter(year=threshold_key[0], subject_key=subject_key)
            .only("minimum_score", "minimum_grade")
            .first()
        )
    threshold = threshold_cache[threshold_key]
    if not threshold:
        return bool(row.get("passed"))
    score_value = float(row.get("score") or 0)
    if threshold.minimum_score is not None:
        return score_value >= float(threshold.minimum_score)
    if threshold.minimum_grade is not None:
        return score_value >= float(threshold.minimum_grade)
    return bool(row.get("passed"))


def _district_qs(district_id: int, exam_type: str, year: int | None = None):
    et = (exam_type or "ege").strip().lower()
    if et not in {"ege", "oge"}:
        et = "ege"
    qs = ExamResult.objects.filter(student__school__district_id=district_id, exam__exam_type=et)
    if year:
        qs = qs.filter(exam__year=year)
    return qs, et


def _district_pass_count(qs, exam_type: str) -> int:
    total = qs.count()
    if not total:
        return 0
    if exam_type == "ege":
        rows = list(qs.values("exam__subject", "exam__code", "exam__year", "score", "passed"))
        cache: dict = {}
        return sum(1 for r in rows if _is_passed_ege_row(r, cache))
    return qs.filter(passed=True).count()


def _district_subject_rows(qs, exam_type: str, high_threshold: float) -> list[dict]:
    rows = list(
        qs.values("exam__subject")
        .annotate(
            participants=Count("id"),
            avg=Avg("score"),
            min_v=Min("score"),
            max_v=Max("score"),
            high=Count("id", filter=Q(score__gte=high_threshold)),
        )
        .order_by("exam__subject")
    )
    # min_v / max_v нужны для PDF/XLSX свода в export_reports
    if exam_type == "ege":
        ege_by_subject: dict[str, list] = {}
        for row in qs.values("exam__subject", "exam__code", "exam__year", "score", "passed"):
            subj = row["exam__subject"] or "Предмет не указан"
            ege_by_subject.setdefault(subj, []).append(row)
        cache: dict = {}
        passed_map = {
            subj: sum(1 for r in items if _is_passed_ege_row(r, cache))
            for subj, items in ege_by_subject.items()
        }
    for row in rows:
        subj = row["exam__subject"] or "Предмет не указан"
        p = int(row["participants"] or 0)
        if exam_type == "ege":
            passed = passed_map.get(subj, 0)
        else:
            passed = qs.filter(exam__subject=subj, passed=True).count()
        row["passed"] = passed
        row["pass_rate"] = round((passed / p) * 100, 1) if p else 0.0
        row["quality_rate"] = round((int(row["high"] or 0) / p) * 100, 1) if p else 0.0
        row["failed"] = max(p - passed, 0)
        row["fail_rate"] = round((row["failed"] / p) * 100, 1) if p else 0.0
        row["risk"] = _risk_level_ru_from_rates(float(row["pass_rate"] or 0), low=60, mid=75)
    return rows


def _district_school_rows(qs, exam_type: str) -> list[dict]:
    if exam_type == "ege":
        raw = list(
            qs.values(
                "student__school_id",
                "student__school__code",
                "student__school__name",
                "exam__subject",
                "exam__code",
                "exam__year",
                "score",
                "passed",
            )
        )
        cache: dict = {}
        buckets: dict[int, dict] = {}
        for row in raw:
            sid = int(row["student__school_id"] or 0)
            b = buckets.setdefault(
                sid,
                {
                    "student__school_id": sid,
                    "student__school__code": row.get("student__school__code"),
                    "student__school__name": row.get("student__school__name"),
                    "participants": 0,
                    "sum_score": 0.0,
                    "passed": 0,
                },
            )
            b["participants"] += 1
            b["sum_score"] += float(row.get("score") or 0)
            if _is_passed_ege_row(row, cache):
                b["passed"] += 1
        school_rows = []
        for b in buckets.values():
            p = int(b["participants"] or 0)
            school_rows.append(
                {
                    **b,
                    "avg": round(b["sum_score"] / p, 2) if p else 0.0,
                    "pass_rate": round((int(b["passed"] or 0) / p) * 100, 1) if p else 0.0,
                }
            )
        return sorted(school_rows, key=lambda x: (-float(x.get("avg") or 0), str(x.get("student__school__name") or "")))

    school_rows = list(
        qs.values("student__school_id", "student__school__code", "student__school__name")
        .annotate(participants=Count("id"), avg=Avg("score"), passed=Count("id", filter=Q(passed=True)))
        .order_by("student__school__name")
    )
    for row in school_rows:
        p = int(row["participants"] or 0)
        row["pass_rate"] = round((int(row["passed"] or 0) / p) * 100, 1) if p else 0.0
    return school_rows


def _republic_avg(exam_type: str, year: int) -> float | None:
    metrics = _aggregate_score_metrics(
        ExamResult.objects.filter(exam__exam_type=exam_type, exam__year=year),
        exam_type,
        precise_pass=False,
    )
    return None if not metrics else metrics["avg_score"]


def _aggregate_score_metrics(qs, exam_type: str, *, precise_pass: bool = True) -> dict | None:
    """
    Реальные сводные метрики по queryset ExamResult.
    precise_pass=True — порог ЕГЭ через _district_pass_count;
    False — поле passed протокола (для крупных республиканских срезов).
    """
    total = qs.count()
    if not total:
        return None
    et = (exam_type or "ege").strip().lower()
    quality_threshold = 60 if et == "ege" else 4
    high_threshold = 70 if et == "ege" else 5

    # Один проход агрегатов вместо серии отдельных count().
    agg = qs.aggregate(
        avg_score=Avg("score"),
        quality_count=Count("id", filter=Q(score__gte=quality_threshold)),
        high_count=Count("id", filter=Q(score__gte=high_threshold)),
        passed_flag_count=Count("id", filter=Q(passed=True)),
    )
    avg_score = round(float(agg["avg_score"] or 0), 2)
    quality_count = int(agg["quality_count"] or 0)
    high_count = int(agg["high_count"] or 0)
    if et == "oge" or not precise_pass:
        pass_count = int(agg["passed_flag_count"] or 0)
    else:
        pass_count = _district_pass_count(qs, et)

    return {
        "total": total,
        "participants": total,  # для KPI-срезов достаточно объёма записей; distinct дорог на республике
        "avg_score": avg_score,
        "quality_rate": round((quality_count / total) * 100, 1) if total else 0.0,
        "pass_rate": round((pass_count / total) * 100, 1) if total else 0.0,
        "high_count": high_count,
        "failed_count": max(total - pass_count, 0),
    }


def _aggregate_district_year_metrics(district_id: int, exam_type: str, year: int) -> dict | None:
    """Метрики муниципалитета за год с точным порогом и уникальными участниками."""
    qs = ExamResult.objects.filter(
        student__school__district_id=district_id,
        exam__exam_type=exam_type,
        exam__year=year,
    )
    metrics = _aggregate_score_metrics(qs, exam_type, precise_pass=True)
    if not metrics:
        return None
    metrics["participants"] = qs.values("student_id").distinct().count()
    return metrics


def _republic_metrics(exam_type: str, year: int, subject: str | None = None) -> dict | None:
    qs = ExamResult.objects.filter(exam__exam_type=exam_type, exam__year=year)
    if subject:
        qs = qs.filter(exam__subject=subject)
        # Предметный срез компактнее — считаем порог ЕГЭ точно.
        return _aggregate_score_metrics(qs, exam_type, precise_pass=True)
    return _aggregate_score_metrics(qs, exam_type, precise_pass=False)


def _append_ai_texts(payload: dict, exam_type: str, report_kind: str, stats: dict, **draft_kw) -> dict:
    et = (exam_type or "ege").strip().lower()
    ai = enrich_district_report_with_ai(
        exam_type=exam_type,
        stats=stats,
        report_kind=report_kind,
        draft_insights=draft_kw.get("draft_insights"),
        draft_recommendations=draft_kw.get("draft_recommendations"),
        draft_conclusions=draft_kw.get("draft_conclusions"),
        draft_executive_summary=draft_kw.get("draft_executive_summary"),
        extra_context=draft_kw.get("extra_context", ""),
        payload=payload
        if report_kind
        in {"gia_summary", "info_stat", "analytic_note", "school_comparison", "management"}
        and et in {"ege", "oge"}
        else None,
    )
    payload["ai_insights"] = ai.get("insights", [])
    payload["recommendations"] = ai.get("recommendations", [])
    payload["conclusions"] = ai.get("conclusions", [])
    payload["executive_summary"] = ai.get("executive_summary", [])
    payload["report_sections"] = ai.get("report_sections", [])
    if payload["executive_summary"]:
        payload["summary"] = list(payload["executive_summary"]) + list(payload.get("summary") or [])[:2]
    return payload


def _build_district_gia_summary_core(district_id: int, exam_type: str, year: int | None = None) -> dict:
    """Сводные метрики муниципалитета без AI-обогащения."""
    qs, et = _district_qs(district_id, exam_type, year)
    if not qs.exists():
        label = "ЕГЭ" if et == "ege" else "ОГЭ"
        return {"has_data": False, "message": f"Недостаточно данных для формирования свода результатов ГИА по району ({label})."}

    district = District.objects.filter(id=district_id).only("name", "code").first()
    total = qs.count()
    participants = qs.values("student_id").distinct().count()
    avg_score = round(float(qs.aggregate(v=Avg("score"))["v"] or 0), 2)
    quality_threshold = 60 if et == "ege" else 4
    high_threshold = 70 if et == "ege" else 5
    quality_count = qs.filter(score__gte=quality_threshold).count()
    quality_rate = round((quality_count / total) * 100, 1) if total else 0.0
    pass_count = _district_pass_count(qs, et)
    pass_rate = round((pass_count / total) * 100, 1) if total else 0.0
    high_count = qs.filter(score__gte=high_threshold).count()
    failed_count = max(total - pass_count, 0)

    selected_year = year if year else int(qs.order_by("-exam__year").values_list("exam__year", flat=True).first())
    prev_metrics = _aggregate_district_year_metrics(district_id, et, selected_year - 1)
    prev_avg = prev_metrics["avg_score"] if prev_metrics else None
    avg_delta = round(avg_score - prev_avg, 2) if prev_avg is not None else None
    prev_pass_rate = prev_metrics["pass_rate"] if prev_metrics else None
    prev_quality_rate = prev_metrics["quality_rate"] if prev_metrics else None
    prev_participants = prev_metrics["participants"] if prev_metrics else None

    distribution = (
        [
            ("2", qs.filter(score__lt=3).count()),
            ("3", qs.filter(score__gte=3, score__lt=4).count()),
            ("4", qs.filter(score__gte=4, score__lt=5).count()),
            ("5", qs.filter(score__gte=5).count()),
        ]
        if et == "oge"
        else [
            ("0-35", qs.filter(score__lte=35).count()),
            ("36-60", qs.filter(score__gt=35, score__lte=60).count()),
            ("61-80", qs.filter(score__gt=60, score__lte=80).count()),
            ("81-100", qs.filter(score__gt=80).count()),
        ]
    )

    subject_rows = _district_subject_rows(qs, et, high_threshold)
    school_rows = _district_school_rows(qs, et)
    district_max_score = float(qs.aggregate(v=Max("score"))["v"] or 100)
    weak_subjects = sorted(
        [r for r in subject_rows if _is_weak_subject_row(r, exam_type=et, max_score=district_max_score, pass_rate_threshold=75)],
        key=lambda x: (x["pass_rate"], float(x["avg"] or 0)),
    )
    schools_count = len(school_rows)
    republic_m = _republic_metrics(et, selected_year)
    republic = republic_m["avg_score"] if republic_m else None
    republic_pass_rate = republic_m["pass_rate"] if republic_m else None
    republic_quality_rate = republic_m["quality_rate"] if republic_m else None
    republic_high_count = republic_m["high_count"] if republic_m else None
    republic_failed_count = republic_m["failed_count"] if republic_m else None

    summary = [
        f"Муниципалитет: {district.name if district else 'Район'}. Участников: {participants}; средний результат: {avg_score}; успеваемость: {pass_rate}%.",
        f"ОО в отчёте: {schools_count}. Высокие результаты: {high_count}; неудовлетворительные: {failed_count}.",
    ]
    if republic is not None:
        summary.append(f"Средний балл республики: {republic}; отклонение района: {round(avg_score - republic, 2):+}.")
    if avg_delta is not None:
        summary.append(f"Динамика к {selected_year - 1} году: {'+' if avg_delta > 0 else ''}{avg_delta}.")
    if weak_subjects:
        summary.append("Проблемные предметы: " + ", ".join((r["exam__subject"] or "предмет") for r in weak_subjects[:4]) + ".")

    return {
        "has_data": True,
        "district_name": district.name if district else "",
        "exam_type": et,
        "year": selected_year,
        "generated_at": date.today().strftime("%d.%m.%Y"),
        "participants": participants,
        "total": total,
        "schools_count": schools_count,
        "avg_score": avg_score,
        "quality_rate": quality_rate,
        "pass_rate": pass_rate,
        "high_count": high_count,
        "failed_count": failed_count,
        "avg_delta": avg_delta,
        "prev_avg": prev_avg,
        "prev_pass_rate": prev_pass_rate,
        "prev_quality_rate": prev_quality_rate,
        "prev_participants": prev_participants,
        "republic_avg": republic,
        "republic_pass_rate": republic_pass_rate,
        "republic_quality_rate": republic_quality_rate,
        "republic_high_count": republic_high_count,
        "republic_failed_count": republic_failed_count,
        "distribution": distribution,
        "subject_rows": subject_rows,
        "school_rows": school_rows,
        "weak_subjects": weak_subjects,
        "summary": summary,
    }


def _build_district_gia_summary_payload(district_id: int, exam_type: str, year: int | None = None) -> dict:
    payload = _build_district_gia_summary_core(district_id, exam_type, year)
    if not payload.get("has_data"):
        return payload
    return _append_ai_texts(
        payload,
        payload["exam_type"],
        "gia_summary",
        {
            "year": payload["year"],
            "participants": payload["participants"],
            "schools_count": payload["schools_count"],
            "avg_score": payload["avg_score"],
            "pass_rate": payload["pass_rate"],
            "quality_rate": payload["quality_rate"],
            "republic_avg": payload.get("republic_avg"),
            "avg_delta": payload.get("avg_delta"),
            "weak_subjects": [(r.get("exam__subject") or "") for r in (payload.get("weak_subjects") or [])[:6]],
        },
        draft_executive_summary=payload.get("summary"),
        draft_insights=payload.get("summary"),
        draft_recommendations=[
            "Провести предметный разбор по зонам риска в школах с отставанием.",
            "Организовать муниципальный мониторинг пробных работ по проблемным предметам.",
        ],
    )


def _build_district_info_stat_payload(district_id: int, exam_type: str, year: int | None = None) -> dict:
    """
    Информационно-статистический отчёт муниципалитета.

    Только сводные метрики (без GigaChat): тексты и таблицы строит presentation-слой.
    """
    base = _build_district_gia_summary_core(district_id, exam_type, year)
    if not base.get("has_data"):
        return base
    base["report_title"] = "Информационно-статистический отчёт по муниципалитету"
    risk_schools = [r for r in base["school_rows"] if float(r.get("pass_rate") or 0) < 70][:8]
    base["risk_schools"] = risk_schools
    return base


def _build_district_analytic_note_payload(district_id: int, exam_type: str, year: int | None = None) -> dict:
    """
    Аналитическая справка муниципалитета.

    База — сводные метрики района. В отличие от школьной справки:
    единица анализа — ОО, есть сравнение с республикой и рейтинг школ.
    Тексты экспертных блоков строятся детерминированно в presentation-слое
    (без GigaChat), по той же таксономии факт/вывод/гипотеза/решение.
    """
    # Берём сводные метрики без AI — тексты строит presentation-слой.
    base = _build_district_gia_summary_core(district_id, exam_type, year)
    if not base.get("has_data"):
        return base

    et = base["exam_type"]
    selected_year = int(base["year"])
    school_rows = list(base.get("school_rows") or [])
    risk_schools = [
        r for r in school_rows if float(r.get("pass_rate") or 0) < 50.0
    ]
    watch_schools = [
        r for r in school_rows if 50.0 <= float(r.get("pass_rate") or 0) < 70.0
    ]
    leader_schools = sorted(
        school_rows,
        key=lambda x: (float(x.get("avg") or 0), float(x.get("pass_rate") or 0)),
        reverse=True,
    )[:5]

    dynamics = []
    for dyn_year in [selected_year - 2, selected_year - 1, selected_year]:
        y_qs = ExamResult.objects.filter(
            student__school__district_id=district_id,
            exam__exam_type=et,
            exam__year=dyn_year,
        )
        if not y_qs.exists():
            continue
        results = y_qs.count()
        participants = y_qs.values("student_id").distinct().count()
        passed = _district_pass_count(y_qs, et)
        dynamics.append(
            {
                "year": dyn_year,
                "participants": participants,
                "results": results,
                "avg": round(float(y_qs.aggregate(v=Avg("score"))["v"] or 0), 2),
                "pass_rate": round((passed / results) * 100, 1) if results else 0.0,
            }
        )

    avg_label = "средняя оценка" if et == "oge" else "средний балл"
    weak_subjects = base.get("weak_subjects") or []
    conclusions = [
        f"По муниципалитету {avg_label} {base['avg_score']}, успеваемость {base['pass_rate']}%.",
        f"В отчёте {base.get('schools_count') or 0} ОО; школ риска (усп. <50%): {len(risk_schools)}.",
    ]
    if base.get("republic_avg") is not None:
        conclusions.append(
            f"Средний результат республики: {base['republic_avg']}; "
            f"отклонение муниципалитета: {round(float(base['avg_score']) - float(base['republic_avg']), 2):+}."
        )
    if weak_subjects:
        conclusions.append(
            "Проблемные предметы: "
            + ", ".join((r.get("exam__subject") or "") for r in weak_subjects[:4])
            + "."
        )
    recommendations = [
        "Согласовать с директорами ОО планы коррекции по предметам с низкой успеваемостью.",
        "Провести муниципальные методические сессии и мониторинг пробных работ по зонам риска.",
        "Организовать адресную поддержку школ с успеваемостью ниже 50%.",
    ]

    base["report_title"] = "Аналитическая справка по итогам ГИА (муниципалитет)"
    base["dynamics"] = dynamics
    base["subjects_count"] = len(base.get("subject_rows") or [])
    base["risk_count"] = len(risk_schools)
    base["risk_schools"] = risk_schools
    base["watch_schools"] = watch_schools
    base["leader_schools"] = leader_schools
    base["school_clusters"] = _build_district_school_clusters(school_rows, int(base.get("participants") or 0))
    base["regional_position"] = _build_district_regional_position(
        district_id=district_id,
        exam_type=et,
        year=selected_year,
        own_metrics={
            "avg_score": base.get("avg_score"),
            "quality_rate": base.get("quality_rate"),
            "pass_rate": base.get("pass_rate"),
            "high_count": base.get("high_count"),
            "avg_delta": base.get("avg_delta"),
        },
    )
    base["conclusions"] = conclusions
    base["recommendations"] = recommendations
    base["narrative_mode"] = "rule_based_district"
    return base


def _build_district_school_clusters(school_rows: list[dict], total_participants: int) -> dict:
    """Кластеры ОО только для аналитической справки муниципалитета."""
    clusters = {
        "leaders": [],
        "stable": [],
        "attention": [],
        "risk": [],
    }
    for row in school_rows or []:
        pr = float(row.get("pass_rate") or 0)
        av = float(row.get("avg") or 0)
        item = {
            "name": row.get("student__school__name") or "ОО",
            "code": row.get("student__school__code") or "—",
            "participants": int(row.get("participants") or 0),
            "avg": round(av, 2),
            "pass_rate": pr,
        }
        if pr >= 85 or (pr >= 80 and av >= 55):
            clusters["leaders"].append(item)
        elif pr >= 70:
            clusters["stable"].append(item)
        elif pr >= 50:
            clusters["attention"].append(item)
        else:
            clusters["risk"].append(item)

    def _pack(key: str, title: str, feature: str) -> dict:
        rows = clusters[key]
        parts = sum(int(r.get("participants") or 0) for r in rows)
        avg = round(sum(float(r.get("avg") or 0) for r in rows) / len(rows), 2) if rows else None
        share = round((parts / total_participants) * 100, 1) if total_participants and rows else 0.0
        return {
            "key": key,
            "title": title,
            "count": len(rows),
            "avg": avg,
            "participants": parts,
            "share": share,
            "feature": feature,
            "schools": rows[:8],
        }

    return {
        "leaders": _pack("leaders", "Лидеры", "Высокая успеваемость и устойчивый средний результат."),
        "stable": _pack("stable", "Стабильные", "Результаты в пределах приемлемого муниципального коридора."),
        "attention": _pack("attention", "Требуют внимания", "Успеваемость ниже целевого уровня, нужна адресная поддержка."),
        "risk": _pack("risk", "Зона риска", "Критически низкая успеваемость, приоритет муниципального сопровождения."),
    }


def _build_district_regional_position(
    *,
    district_id: int,
    exam_type: str,
    year: int,
    own_metrics: dict,
) -> dict | None:
    """Место муниципалитета среди других районов региона (только для этой справки)."""
    et = (exam_type or "ege").strip().lower()
    value_fields = (
        "student__school__district_id",
        "exam__subject",
        "exam__code",
        "exam__year",
        "score",
        "passed",
    )
    current_rows = list(
        ExamResult.objects.filter(exam__exam_type=et, exam__year=year)
        .exclude(student__school__district_id__isnull=True)
        .values(*value_fields)
    )
    if not current_rows:
        return None

    high_threshold = 70 if et == "ege" else 5
    quality_threshold = 60 if et == "ege" else 4
    threshold_cache: dict = {}

    by_district: dict[int, list[dict]] = {}
    for row in current_rows:
        did = int(row["student__school__district_id"])
        by_district.setdefault(did, []).append(row)
    if len(by_district) < 2:
        return None

    prev_sums: dict[int, list[float]] = {}
    for row in (
        ExamResult.objects.filter(exam__exam_type=et, exam__year=year - 1)
        .exclude(student__school__district_id__isnull=True)
        .values("student__school__district_id", "score")
        .iterator(chunk_size=5000)
    ):
        did = int(row["student__school__district_id"])
        prev_sums.setdefault(did, []).append(float(row.get("score") or 0))

    rows = []
    for did, items in by_district.items():
        total = len(items)
        if not total:
            continue
        scores = [float(r.get("score") or 0) for r in items]
        avg = round(sum(scores) / total, 2)
        if et == "ege":
            passed = sum(1 for r in items if _is_passed_ege_row(r, threshold_cache))
        else:
            passed = sum(1 for r in items if bool(r.get("passed")))
        pass_rate = round((passed / total) * 100, 1)
        quality_rate = round((sum(1 for s in scores if s >= quality_threshold) / total) * 100, 1)
        high_count = sum(1 for s in scores if s >= high_threshold)
        prev_scores = prev_sums.get(did) or []
        prev_avg = round(sum(prev_scores) / len(prev_scores), 2) if prev_scores else None
        delta = round(avg - prev_avg, 2) if prev_avg is not None else None
        rows.append(
            {
                "district_id": did,
                "avg_score": avg,
                "pass_rate": pass_rate,
                "quality_rate": quality_rate,
                "high_count": high_count,
                "avg_delta": delta,
            }
        )
    if len(rows) < 2:
        return None

    def _rank(metric: str, reverse: bool = True) -> dict | None:
        ranked = sorted(
            [r for r in rows if r.get(metric) is not None],
            key=lambda x: float(x.get(metric) or 0),
            reverse=reverse,
        )
        if not ranked:
            return None
        place = next((i + 1 for i, r in enumerate(ranked) if int(r["district_id"]) == int(district_id)), None)
        if place is None:
            return None
        return {
            "metric": metric,
            "place": place,
            "total": len(ranked),
            "value": next((r.get(metric) for r in ranked if int(r["district_id"]) == int(district_id)), None),
        }

    positions = {
        "avg_score": _rank("avg_score"),
        "quality_rate": _rank("quality_rate"),
        "pass_rate": _rank("pass_rate"),
        "high_count": _rank("high_count"),
        "avg_delta": _rank("avg_delta"),
    }
    available = {k: v for k, v in positions.items() if v}
    if not available:
        return None
    return {
        "has_data": True,
        "districts_total": len(rows),
        "positions": available,
        "own": own_metrics,
    }


def generate_district_analytic_note_docx(district_id: int, exam_type: str, year: int | None = None) -> BytesIO:
    from users.report_ui.district_analytic_note import build_district_analytic_note_presentation
    from users.report_ui.district_analytic_note_docx import render_district_analytic_note_docx

    data = _build_district_analytic_note_payload(district_id, exam_type, year)
    ui = build_district_analytic_note_presentation(data)
    return render_district_analytic_note_docx(data, ui)


def _build_district_subject_note_payload(
    district_id: int,
    exam_type: str,
    subject: str,
    year: int | None = None,
    *,
    with_ai: bool = True,
) -> dict:
    et = (exam_type or "ege").strip().lower()
    subject_name = (subject or "").strip()
    if not subject_name:
        return {"has_data": False, "message": "Выберите предмет для предметной справки."}

    scope_qs = ExamResult.objects.filter(
        student__school__district_id=district_id,
        exam__exam_type=et,
        exam__subject=subject_name,
    )
    if year:
        scope_qs = scope_qs.filter(exam__year=year)
        selected_year = int(year)
    else:
        selected_year = int(scope_qs.order_by("-exam__year").values_list("exam__year", flat=True).first() or 0)
        if not selected_year:
            return {"has_data": False, "message": f"Нет данных по предмету «{subject_name}» в муниципалитете."}
        scope_qs = scope_qs.filter(exam__year=selected_year)

    attempts_qs = scope_qs
    qs = filter_latest_exam_results(attempts_qs)
    if not qs.exists():
        return {"has_data": False, "message": f"Нет данных по предмету «{subject_name}» в муниципалитете."}

    district = District.objects.filter(id=district_id).only("name", "code").first()
    total = qs.count()
    participants = qs.values("student_id").distinct().count()
    avg_score = round(float(qs.aggregate(v=Avg("score"))["v"] or 0), 2)
    max_score = float(qs.aggregate(v=Max("score"))["v"] or 100)
    quality_threshold = 60 if et == "ege" else 4
    high_threshold = 70 if et == "ege" else 5
    quality_count = qs.filter(score__gte=quality_threshold).count()
    quality_rate = round((quality_count / total) * 100, 1) if total else 0.0
    high_count = qs.filter(score__gte=high_threshold).count()
    pass_count = _district_pass_count(qs, et) if et == "ege" else qs.filter(passed=True).count()
    pass_rate = round((pass_count / total) * 100, 1) if total else 0.0

    prev_qs = filter_latest_exam_results(
        ExamResult.objects.filter(
            student__school__district_id=district_id,
            exam__exam_type=et,
            exam__subject=subject_name,
            exam__year=selected_year - 1,
        )
    )
    prev_avg = round(float(prev_qs.aggregate(v=Avg("score"))["v"] or 0), 2) if prev_qs.exists() else None
    avg_delta = round(avg_score - prev_avg, 2) if prev_avg is not None else None

    # Республиканские метрики по тому же предмету и году.
    rep_m = _republic_metrics(et, selected_year, subject=subject_name)
    republic_avg = rep_m["avg_score"] if rep_m else None
    republic_pass_rate = rep_m["pass_rate"] if rep_m else None
    republic_quality_rate = rep_m["quality_rate"] if rep_m else None

    # Место предмета среди предметов муниципалитета по среднему.
    subject_rank = None
    subjects_total = 0
    district_subj_avgs = []
    subject_names = (
        ExamResult.objects.filter(
            student__school__district_id=district_id,
            exam__exam_type=et,
            exam__year=selected_year,
        )
        .values_list("exam__subject", flat=True)
        .distinct()
    )
    for subj in subject_names:
        if not subj:
            continue
        subj_latest = filter_latest_exam_results(
            ExamResult.objects.filter(
                student__school__district_id=district_id,
                exam__exam_type=et,
                exam__subject=subj,
                exam__year=selected_year,
            )
        )
        if not subj_latest.exists():
            continue
        district_subj_avgs.append(
            {"exam__subject": subj, "avg": float(subj_latest.aggregate(v=Avg("score"))["v"] or 0)}
        )
    ranked = sorted(
        [r for r in district_subj_avgs if r.get("exam__subject")],
        key=lambda x: (-float(x.get("avg") or 0), str(x.get("exam__subject") or "")),
    )
    subjects_total = len(ranked)
    for i, row in enumerate(ranked, start=1):
        if row.get("exam__subject") == subject_name:
            subject_rank = i
            break

    if et == "ege":
        school_buckets: dict[str, dict] = {}
        cache: dict = {}
        for row in qs.values(
            "student__school__name",
            "student__school__code",
            "exam__subject",
            "exam__code",
            "exam__year",
            "score",
            "passed",
        ):
            key = row.get("student__school__name") or "—"
            b = school_buckets.setdefault(
                key,
                {
                    "student__school__name": row.get("student__school__name"),
                    "student__school__code": row.get("student__school__code"),
                    "participants": 0,
                    "sum_score": 0.0,
                    "passed": 0,
                    "quality": 0,
                    "high": 0,
                },
            )
            score = float(row.get("score") or 0)
            b["participants"] += 1
            b["sum_score"] += score
            if _is_passed_ege_row(row, cache):
                b["passed"] += 1
            if score >= quality_threshold:
                b["quality"] += 1
            if score >= high_threshold:
                b["high"] += 1
        school_rows = []
        for b in school_buckets.values():
            p = int(b["participants"] or 0)
            school_rows.append(
                {
                    **b,
                    "avg": round(b["sum_score"] / p, 2) if p else 0.0,
                    "pass_rate": round((int(b["passed"] or 0) / p) * 100, 1) if p else 0.0,
                    "quality_rate": round((int(b["quality"] or 0) / p) * 100, 1) if p else 0.0,
                    "high_count": int(b["high"] or 0),
                }
            )
        school_rows.sort(key=lambda x: -float(x.get("avg") or 0))
    else:
        school_rows = list(
            qs.values("student__school__name", "student__school__code")
            .annotate(
                participants=Count("id"),
                avg=Avg("score"),
                passed=Count("id", filter=Q(passed=True)),
                quality=Count("id", filter=Q(score__gte=quality_threshold)),
                high=Count("id", filter=Q(score__gte=high_threshold)),
            )
            .order_by("-avg")
        )
        for row in school_rows:
            p = int(row["participants"] or 0)
            row["pass_rate"] = round((int(row["passed"] or 0) / p) * 100, 1) if p else 0.0
            row["quality_rate"] = round((int(row["quality"] or 0) / p) * 100, 1) if p else 0.0
            row["high_count"] = int(row["high"] or 0)

    task_qs = task_results_for_exam_results(
        TaskResult.objects.filter(
            student__school__district_id=district_id,
            exam__exam_type=et,
            exam__subject=subject_name,
            exam__year=selected_year,
        ),
        attempts_qs,
    )
    task_rows = list(
        task_qs.values("task_number")
        .annotate(total=Count("id"), plus=Count("id", filter=~Q(value__in=["-", "0", ""])))
        .order_by("task_number")
    )
    for row in task_rows:
        t = int(row["total"] or 0)
        plus = int(row["plus"] or 0)
        row["success_rate"] = round((plus / t) * 100, 1) if t else 0.0
        row["topic"] = _topic_for_task(subject_name, int(row["task_number"]), et)
    weak_tasks = [r for r in task_rows if r["success_rate"] < 50]

    from analytics.engine import AnalyticsEngine

    engine_result = AnalyticsEngine().analyze_district_subject(
        district_id, et, subject_name, selected_year
    )
    if engine_result.valid:
        by_number = {task.task_number: task for task in engine_result.tasks}
        for row in task_rows:
            task = by_number.get(int(row["task_number"]))
            if task:
                row["topic"] = task.topic
                row["skill_name"] = task.skill_name

    payload = {
        "has_data": True,
        "report_title": f"Предметная аналитическая справка: {subject_name}",
        "district_name": district.name if district else "",
        "exam_type": et,
        "year": selected_year,
        "generated_at": date.today().strftime("%d.%m.%Y"),
        "aggregate_label": f"итоговые результаты за {selected_year} год",
        "subject": subject_name,
        "total": total,
        "participants": participants,
        "avg_score": avg_score,
        "pass_rate": pass_rate,
        "quality_rate": quality_rate,
        "high_count": high_count,
        "avg_delta": avg_delta,
        "republic_avg": republic_avg,
        "republic_pass_rate": republic_pass_rate,
        "republic_quality_rate": republic_quality_rate,
        "subject_rank": subject_rank,
        "subjects_total": subjects_total,
        "school_rows": school_rows,
        "task_rows": task_rows,
        "weak_tasks": weak_tasks,
    }
    if engine_result.valid:
        payload["subject_engine"] = engine_result
    if not with_ai:
        return payload
    return _append_ai_texts(
        payload,
        et,
        "subject_note",
        {
            "subject": subject_name,
            "year": selected_year,
            "avg_score": avg_score,
            "pass_rate": pass_rate,
            "quality_rate": quality_rate,
            "weak_tasks": [f"№{r['task_number']}" for r in weak_tasks[:8]],
            "schools": len(school_rows),
        },
        extra_context=f"Предмет: {subject_name}",
        draft_insights=[f"Успеваемость по предмету в муниципалитете: {pass_rate}%."],
        draft_recommendations=["Организовать межшкольный разбор слабых заданий на заседании МО."],
    )


def _school_comparison_category(row: dict, district_avg: float, exam_type: str) -> str:
    pr = float(row.get("pass_rate") or 0)
    avg = float(row.get("avg") or 0)
    delta = avg - district_avg
    if exam_type == "oge":
        if pr < 50:
            return "критическая зона"
        if pr < 65:
            return "зона риска"
        if pr < 75:
            return "зона наблюдения"
        if pr >= 85 and delta >= 0.3:
            return "устойчивый лидер"
        if pr >= 80 and avg >= district_avg:
            return "лидер"
        return "стабильная организация"
    if pr < 45 or avg < district_avg - 12:
        return "критическая зона"
    if pr < 60:
        return "зона риска"
    if pr < 70:
        return "зона наблюдения"
    if pr >= 85 and delta >= 5:
        return "устойчивый лидер"
    if pr >= 80 and avg >= district_avg:
        return "лидер"
    return "стабильная организация"


def _build_district_school_comparison_payload(
    district_id: int,
    exam_type: str,
    year: int | None = None,
    *,
    with_ai: bool = True,
) -> dict:
    base = _build_district_gia_summary_core(district_id, exam_type, year)
    if not base.get("has_data"):
        return base

    et = base["exam_type"]
    selected_year = int(base["year"])
    district_avg = float(base["avg_score"])
    quality_threshold = 60 if et == "ege" else 4
    school_rows = list(base.get("school_rows") or [])

    qs = ExamResult.objects.filter(
        student__school__district_id=district_id,
        exam__exam_type=et,
        exam__year=selected_year,
    )
    for row in school_rows:
        sid = row.get("student__school_id")
        p = int(row.get("participants") or 0)
        if sid:
            s_qs = qs.filter(student__school_id=sid)
            high = s_qs.filter(score__gte=quality_threshold).count()
            row["quality_rate"] = round((high / p) * 100, 1) if p else 0.0
        else:
            row["quality_rate"] = 0.0
        row["delta_vs_district"] = round(float(row.get("avg") or 0) - district_avg, 2)
        pr = float(row.get("pass_rate") or 0)
        row["risk"] = _risk_level_ru_from_rates(pr, low=60, mid=75)
        row["category"] = _school_comparison_category(row, district_avg, et)
        row["tier"] = row["category"]

    leaders = [
        r
        for r in sorted(school_rows, key=lambda x: float(x.get("avg") or 0), reverse=True)
        if r.get("category") in {"устойчивый лидер", "лидер"}
    ][:8]
    at_risk = [r for r in school_rows if r.get("category") in {"зона риска", "критическая зона"}]

    dynamics = []
    for dyn_year in [selected_year - 2, selected_year - 1, selected_year]:
        y_qs = ExamResult.objects.filter(
            student__school__district_id=district_id,
            exam__exam_type=et,
            exam__year=dyn_year,
        )
        if not y_qs.exists():
            continue
        p = y_qs.count()
        passed = _district_pass_count(y_qs, et)
        dynamics.append(
            {
                "year": dyn_year,
                "participants": p,
                "avg": round(float(y_qs.aggregate(v=Avg("score"))["v"] or 0), 2),
                "pass_rate": round((passed / p) * 100, 1) if p else 0.0,
            }
        )

    avg_label = "средняя оценка" if et == "oge" else "средний балл"
    summary = [
        f"Средний показатель муниципалитета ({avg_label}): {district_avg}. "
        f"Школ в зоне риска: {len(at_risk)}. Лидеров: {len(leaders)}.",
    ]

    base["report_title"] = "Сравнительный отчёт по образовательным организациям"
    base["district_avg"] = district_avg
    base["school_rows"] = school_rows
    base["leaders"] = leaders
    base["at_risk"] = at_risk
    base["dynamics"] = dynamics
    base["summary"] = summary

    if not with_ai:
        return base

    return _append_ai_texts(
        base,
        et,
        "school_comparison",
        {
            "year": selected_year,
            "district_avg": district_avg,
            "schools_count": len(school_rows),
            "at_risk_count": len(at_risk),
            "leaders": [r.get("student__school__name") for r in leaders],
        },
        draft_executive_summary=summary,
        draft_insights=summary,
        draft_recommendations=[
            "Организовать адресную поддержку школ зоны риска и критической зоны.",
            "Тиражировать практики подготовки школ-лидеров на заседании муниципального МО.",
        ],
    )


def _build_district_mo_payload(
    district_id: int,
    exam_type: str,
    mo_key: str,
    year: int | None = None,
    *,
    with_ai: bool = True,
) -> dict:
    et = (exam_type or "ege").strip().lower()
    helpers = _export_helpers()
    mo_groups = helpers["MO_SUBJECT_GROUPS"]
    meta = mo_groups.get(mo_key)
    if not meta:
        return {"has_data": False, "message": "Неверный профиль методического объединения."}
    mo_title = meta["title"]

    qs = ExamResult.objects.filter(student__school__district_id=district_id, exam__exam_type=et)
    if year:
        qs = qs.filter(exam__year=year)

    # Сопоставление предметов через casefold — SQLite icontains ломает кириллицу.
    from users.export_reports import _match_mo_subjects

    available = list(qs.values_list("exam__subject", flat=True).distinct())
    matched, _ = _match_mo_subjects(mo_key, available)
    if matched is None:
        return {"has_data": False, "message": "Неверный профиль методического объединения."}
    if not matched:
        return {"has_data": False, "message": "Недостаточно данных для отчёта МО по муниципалитету."}
    qs = qs.filter(exam__subject__in=matched)

    total = qs.count()
    avg_score = round(float(qs.aggregate(v=Avg("score"))["v"] or 0), 2)
    pass_count = _district_pass_count(qs, et)
    pass_rate = round((pass_count / total) * 100, 1) if total else 0.0
    max_score = float(qs.aggregate(v=Max("score"))["v"] or 100)
    subject_rows = list(
        qs.values("exam__subject")
        .annotate(participants=Count("id"), avg=Avg("score"), passed=Count("id", filter=Q(passed=True)))
        .order_by("exam__subject")
    )
    for row in subject_rows:
        p = int(row["participants"] or 0)
        row["pass_rate"] = round((int(row["passed"] or 0) / p) * 100, 1) if p else 0.0
    weak_subjects = [r for r in subject_rows if _is_weak_subject_row(r, exam_type=et, max_score=max_score)]
    selected_year = year if year else int(qs.order_by("-exam__year").values_list("exam__year", flat=True).first())

    school_rows = list(
        qs.values("student__school__name")
        .annotate(participants=Count("id"), avg=Avg("score"))
        .order_by("-avg")[:15]
    )

    payload = {
        "has_data": True,
        "mo_title": mo_title,
        "mo_key": mo_key,
        "exam_type": et,
        "year": selected_year,
        "total": total,
        "avg_score": avg_score,
        "pass_rate": pass_rate,
        "subject_rows": subject_rows,
        "school_rows": school_rows,
        "weak_subjects": weak_subjects,
    }
    if not with_ai:
        return payload
    return _append_ai_texts(
        payload,
        et,
        "mo_report",
        {"mo_title": mo_title, "avg_score": avg_score, "pass_rate": pass_rate, "year": selected_year},
        extra_context=f"Методическое объединение: {mo_title}",
        draft_conclusions=[f"{mo_title}: успеваемость {pass_rate}%, средний балл {avg_score}."],
    )


def _build_district_management_payload(
    district_id: int,
    exam_type: str,
    year: int | None = None,
    *,
    with_ai: bool = True,
) -> dict:
    summary = _build_district_gia_summary_core(district_id, exam_type, year)
    if not summary.get("has_data"):
        summary["message"] = summary.get("message", "Недостаточно данных для управленческого отчёта.")
        return summary

    et = summary["exam_type"]
    selected_year = int(summary["year"])
    dynamics = []
    for dyn_year in [selected_year - 2, selected_year - 1, selected_year]:
        y_qs = ExamResult.objects.filter(
            student__school__district_id=district_id,
            exam__exam_type=et,
            exam__year=dyn_year,
        )
        if not y_qs.exists():
            continue
        p = y_qs.count()
        passed = _district_pass_count(y_qs, et)
        dynamics.append(
            {
                "year": dyn_year,
                "participants": p,
                "avg": round(float(y_qs.aggregate(v=Avg("score"))["v"] or 0), 2),
                "pass_rate": round((passed / p) * 100, 1) if p else 0.0,
            }
        )

    at_risk = [r for r in summary.get("school_rows", []) if float(r.get("pass_rate") or 0) < 65]
    task_qs = TaskResult.objects.filter(
        student__school__district_id=district_id,
        exam__exam_type=summary["exam_type"],
        exam__year=summary["year"],
    )
    weak_tasks = []
    if task_qs.exists():
        task_rows = list(
            task_qs.values("exam__subject", "task_number")
            .annotate(total=Count("id"), plus=Count("id", filter=~Q(value__in=["-", "0", ""])))
        )
        for row in task_rows:
            t = int(row["total"] or 0)
            plus = int(row["plus"] or 0)
            rate = round((plus / t) * 100, 1) if t else 0.0
            if rate < 45:
                weak_tasks.append(
                    {
                        "exam__subject": row["exam__subject"],
                        "task_number": row["task_number"],
                        "success_rate": rate,
                        "risk": _risk_level_ru_from_rates(rate),
                    }
                )
        weak_tasks = sorted(weak_tasks, key=lambda x: x["success_rate"])[:15]

    improvement_plan = []
    weak_subjects = list(summary.get("weak_subjects") or [])
    if weak_subjects:
        names = ", ".join((r.get("exam__subject") or "предмет") for r in weak_subjects[:4])
        improvement_plan.append(
            {
                "activity": f"Муниципальный мониторинг пробных экзаменов по предметам риска: {names}",
                "term": "ежемесячно",
                "owner": "отдел образования",
                "expected": "снижение доли неудовлетворительных результатов",
            }
        )
    if at_risk:
        school_names = ", ".join(
            (r.get("student__school__name") or "ОО") for r in at_risk[:4]
        )
        improvement_plan.append(
            {
                "activity": f"Адресная методическая поддержка школ зоны риска ({len(at_risk)} ОО): {school_names}",
                "term": "до следующей волны ГИА",
                "owner": "методисты МО",
                "expected": "рост успеваемости в отстающих ОО",
            }
        )
        improvement_plan.append(
            {
                "activity": "Совещание с руководителями ОО группы риска",
                "term": "в течение 2 недель",
                "owner": "руководитель отдела образования",
                "expected": "утверждённые коррекционные планы",
            }
        )
    if weak_tasks:
        task_sample = ", ".join(
            f"{r.get('exam__subject') or 'предмет'} №{r.get('task_number')}" for r in weak_tasks[:5]
        )
        improvement_plan.append(
            {
                "activity": f"Межшкольный разбор слабых заданий КИМ: {task_sample}",
                "term": "в течение месяца",
                "owner": "руководители РМО",
                "expected": "рост выполнения критических заданий",
            }
        )
    if not improvement_plan:
        improvement_plan.append(
            {
                "activity": "Сохранить текущий мониторинг качества подготовки по итогам ГИА",
                "term": "постоянно",
                "owner": "отдел образования",
                "expected": "устойчивость достигнутых показателей",
            }
        )

    exec_summary = [
        f"Успеваемость муниципалитета: {summary['pass_rate']}%. "
        f"Школ в зоне риска: {len(at_risk)}. Слабых заданий КИМ: {len(weak_tasks)}.",
    ]

    payload = {
        **summary,
        "report_title": "Управленческий отчёт отдела образования муниципалитета",
        "at_risk_schools": at_risk,
        "weak_tasks": weak_tasks,
        "improvement_plan": improvement_plan,
        "dynamics": dynamics,
    }
    if not with_ai:
        payload["executive_summary"] = exec_summary
        payload["summary"] = exec_summary
        return payload
    return _append_ai_texts(
        payload,
        et,
        "management",
        {
            "year": selected_year,
            "at_risk_schools": len(at_risk),
            "weak_tasks_count": len(weak_tasks),
            "pass_rate": summary["pass_rate"],
        },
        draft_executive_summary=exec_summary,
        draft_insights=exec_summary,
        draft_recommendations=[p["activity"] for p in improvement_plan],
    )


def _district_docx_header(doc, data: dict, default_title: str) -> None:
    et_label = "ЕГЭ" if data.get("exam_type") == "ege" else "ОГЭ"
    title = data.get("report_title") or default_title
    doc.add_heading(f"{title} ({et_label})", 0)
    if data.get("district_name"):
        doc.add_paragraph(f"Муниципалитет: {data['district_name']}.")
    if data.get("year"):
        doc.add_paragraph(f"Отчётный год: {data['year']}. Дата: {data.get('generated_at', date.today().strftime('%d.%m.%Y'))}.")


def _docx_add_text_block(doc, heading: str, items: list[str]) -> None:
    if not items:
        return
    doc.add_heading(heading, level=1)
    for line in items[:10]:
        doc.add_paragraph(str(line))


def generate_district_gia_summary_docx(district_id: int, exam_type: str, year: int | None = None) -> BytesIO:
    from users.report_ui.district_gia_summary_docx import render_district_gia_summary_docx

    data = _build_district_gia_summary_payload(district_id, exam_type, year)
    return render_district_gia_summary_docx(data)


def generate_district_info_stat_docx(district_id: int, exam_type: str, year: int | None = None) -> BytesIO:
    from users.report_ui.district_info_stat_docx import render_district_info_stat_docx

    data = _build_district_info_stat_payload(district_id, exam_type, year)
    return render_district_info_stat_docx(data)


def generate_district_subject_note_docx(
    district_id: int,
    exam_type: str,
    subject: str,
    year: int | None = None,
) -> BytesIO:
    from users.report_ui.district_subject_note_docx import render_district_subject_note_docx

    data = _build_district_subject_note_payload(district_id, exam_type, subject, year)
    return render_district_subject_note_docx(data)


def generate_district_school_comparison_docx(district_id: int, exam_type: str, year: int | None = None) -> BytesIO:
    from users.report_ui.district_school_comparison_docx import render_district_school_comparison_docx

    data = _build_district_school_comparison_payload(district_id, exam_type, year)
    return render_district_school_comparison_docx(data)


def generate_district_mo_report_docx(
    district_id: int,
    exam_type: str,
    mo_key: str,
    year: int | None = None,
) -> BytesIO:
    from docx import Document

    data = _build_district_mo_payload(district_id, exam_type, mo_key, year)
    doc = Document()
    if not data["has_data"]:
        doc.add_paragraph(data["message"])
    else:
        doc.add_heading(data.get("mo_title", "Отчёт МО"), 0)
        _docx_add_text_block(doc, "Выводы", data.get("conclusions", []))
        _docx_add_text_block(doc, "Рекомендации", data.get("recommendations", []))
    output = BytesIO()
    doc.save(output)
    output.seek(0)
    return output


def generate_district_management_docx(district_id: int, exam_type: str, year: int | None = None) -> BytesIO:
    from docx import Document

    data = _build_district_management_payload(district_id, exam_type, year)
    doc = Document()
    if not data["has_data"]:
        doc.add_paragraph(data["message"])
    else:
        _district_docx_header(doc, data, "Управленческий отчёт")
        sections = data.get("report_sections") or []
        if sections:
            for section in sections:
                _docx_add_text_block(doc, section.get("title", "Раздел"), section.get("items", []))
        else:
            _docx_add_text_block(doc, "Резюме", data.get("executive_summary", []))
            _docx_add_text_block(doc, "Рекомендации", data.get("recommendations", []))
        if data.get("improvement_plan"):
            doc.add_heading("План мероприятий", level=1)
            t = doc.add_table(rows=1, cols=4)
            h = t.rows[0].cells
            h[0].text, h[1].text, h[2].text, h[3].text = "Мероприятие", "Срок", "Ответственный", "Результат"
            for row in data["improvement_plan"]:
                r = t.add_row().cells
                r[0].text = row["activity"]
                r[1].text = row["term"]
                r[2].text = row["owner"]
                r[3].text = row["expected"]
            _style_docx_table(t, header_rows=1)
    output = BytesIO()
    doc.save(output)
    output.seek(0)
    return output


