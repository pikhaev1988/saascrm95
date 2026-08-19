"""Сравнение ключевых показателей HTML и DOCX с VPRReportFacts."""

from __future__ import annotations

from apps.vpr.validation.consistency import ConsistencyIssue, ConsistencyResult


class CrossFormatConsistencyValidator:
    """
    HTML и DOCX должны показывать те же числа, что VPRReportFacts.
    Не пересчитывает метрики — только сверяет отображение.
    """

    def validate(self, facts, html: str, docx_text: str) -> ConsistencyResult:
        errors: list[ConsistencyIssue] = []
        warnings: list[ConsistencyIssue] = []
        if facts is None:
            errors.append(
                ConsistencyIssue(
                    code="cross_format.no_facts",
                    severity="error",
                    message="VPRReportFacts missing",
                )
            )
            return ConsistencyResult(ok=False, errors=errors, warnings=warnings)

        n = int(facts.participants or 0)
        html_s = html or ""
        docx_s = docx_text or ""
        if n:
            token = str(n)
            if token not in html_s:
                errors.append(
                    ConsistencyIssue(
                        code="cross_format.html_missing",
                        severity="error",
                        message=f"participants={n} not found in HTML",
                        actual=n,
                    )
                )
            if token not in docx_s:
                errors.append(
                    ConsistencyIssue(
                        code="cross_format.docx_missing",
                        severity="error",
                        message=f"participants={n} not found in DOCX",
                        actual=n,
                    )
                )
        for title in ("Группа риска", "Группа высокого уровня"):
            if title in html_s and title not in docx_s:
                errors.append(
                    ConsistencyIssue(
                        code="cross_format.group_title",
                        severity="error",
                        message=f"group title missing in DOCX: {title}",
                    )
                )
        return ConsistencyResult(ok=not errors, errors=errors, warnings=warnings)


def extract_docx_text(buf) -> str:
    from docx import Document

    if hasattr(buf, "seek"):
        buf.seek(0)
    doc = Document(buf)
    parts = [p.text for p in doc.paragraphs]
    for table in doc.tables:
        for row in table.rows:
            parts.extend(cell.text for cell in row.cells)
    if hasattr(buf, "seek"):
        buf.seek(0)
    return "\n".join(parts)
