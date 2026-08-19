"""Выгрузка аналитической справки школы ВПР в Word (ФИОКО, 16 разделов)."""

from __future__ import annotations

from io import BytesIO
from typing import Any

from apps.vpr.labels import label_priority, label_risk, label_school_risk, label_trend
from apps.vpr.school_analysis.schemas import SchoolAnalysisResult


def _fmt(value: Any, suffix: str = "") -> str:
    if value is None or value == "":
        return "—"
    if isinstance(value, float):
        text = f"{value:.2f}".rstrip("0").rstrip(".")
        return f"{text}{suffix}"
    return f"{value}{suffix}"


def _set_cell_shading(cell, fill_hex: str) -> None:
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill_hex)
    cell._tc.get_or_add_tcPr().append(shd)


def _style_table(table, header_rows: int = 1) -> None:
    for style_name in ("Table Grid", "Light Grid Accent 1", "Сетка таблицы"):
        try:
            table.style = style_name
            break
        except (KeyError, ValueError):
            continue
    for ri in range(min(header_rows, len(table.rows))):
        for cell in table.rows[ri].cells:
            _set_cell_shading(cell, "D9E2F3")
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.bold = True


def _add_bullets(doc, items: list[str], *, empty: str = "Нет данных.") -> None:
    cleaned = [str(item).strip() for item in items if str(item or "").strip()]
    if not cleaned:
        doc.add_paragraph(empty)
        return
    for item in cleaned:
        doc.add_paragraph(item, style="List Bullet")


def _add_cycle(
    doc,
    *,
    interpretation: list[str],
    causes: list[str],
    org_decisions: list[str],
    method_decisions: list[str],
    expected_effect: list[str],
) -> None:
    """Обязательный аналитический цикл ФИОКО для каждого раздела."""
    doc.add_heading("Экспертная интерпретация", level=2)
    _add_bullets(doc, interpretation)
    doc.add_heading("Причины выявленных результатов", level=2)
    _add_bullets(doc, causes)
    doc.add_heading("Организационно-управленческие решения", level=2)
    _add_bullets(doc, org_decisions)
    doc.add_heading("Методические решения", level=2)
    _add_bullets(doc, method_decisions)
    doc.add_heading("Ожидаемый результат реализации мероприятий", level=2)
    _add_bullets(doc, expected_effect)


def generate_school_analysis_docx(
    analysis: SchoolAnalysisResult,
    *,
    school_name: str = "",
    academic_year: int | None = None,
    protocol_task_tables: list[dict[str, Any]] | None = None,
) -> BytesIO:
    """Сформировать .docx по SchoolAnalysisResult без пересчётов."""
    from docx import Document

    doc = Document()
    title = "Аналитическая справка по результатам ВПР"
    if academic_year:
        title = f"{title} · {academic_year} учебный год"
    doc.add_heading(title, level=0)
    if school_name:
        doc.add_paragraph(school_name)

    if not analysis.overview.has_data:
        doc.add_paragraph("Нет протоколов ВПР для выбранной школы и года.")
        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer

    ov = analysis.overview
    doc.add_heading("1. Паспорт анализа", level=1)
    overview_rows = [
        ("ОО", school_name or ov.organization_name or "—"),
        ("Учебный год", academic_year or ov.academic_year or "—"),
        ("Протоколов", ov.protocols_count),
        ("Предметов", ov.subjects_count),
        ("Классов", ov.grades_count),
        ("Участников", ov.participants_total),
        ("Средний % выполнения", _fmt(ov.avg_completion_percent, "%")),
        ("Качество знаний", _fmt(ov.avg_quality_percent, "%")),
        ("Успеваемость", _fmt(ov.avg_absolute_percent, "%")),
        ("Среднее число дефицитов", _fmt(ov.avg_deficits_count)),
    ]
    overview_table = doc.add_table(rows=1 + len(overview_rows), cols=2)
    overview_table.rows[0].cells[0].text = "Показатель"
    overview_table.rows[0].cells[1].text = "Значение"
    for idx, (label, value) in enumerate(overview_rows, start=1):
        overview_table.rows[idx].cells[0].text = label
        overview_table.rows[idx].cells[1].text = str(value)
    _style_table(overview_table)
    doc.add_heading("Экспертная оценка общего состояния качества образования", level=2)
    quality = ov.avg_quality_percent or 0
    absolute = ov.avg_absolute_percent or 0
    if quality >= 50 and absolute >= 85:
        doc.add_paragraph(
            "Качество образования характеризуется как устойчивое: образовательные результаты "
            "в целом соответствуют требованиям ВПР, а ВСОКО обеспечивает рабочий уровень "
            "управления качеством."
        )
    elif quality < 35 or absolute < 80:
        doc.add_paragraph(
            "Качество образования требует приоритетного управленческого внимания: "
            "фиксируются риски недостижения планируемых результатов и необходимость "
            "корректировки организационно-управленческих и методических решений."
        )
    else:
        doc.add_paragraph(
            "Качество образования находится на рабочем уровне: достигнуты базовые "
            "образовательные результаты, при этом сохраняется зона роста в части "
            "предметных и метапредметных результатов."
        )

    doc.add_heading("2. Анализ индивидуальных результатов", level=1)
    if analysis.grades:
        table = doc.add_table(rows=1 + len(analysis.grades), cols=6)
        headers = ("Класс", "% выполнения", "Качество", "Риск", "Основные темы", "Основные умения")
        for i, header in enumerate(headers):
            table.rows[0].cells[i].text = header
        for r_idx, row in enumerate(analysis.grades, start=1):
            values = (
                row.parallel,
                _fmt(row.avg_completion_percent, "%"),
                _fmt(row.quality_percent),
                label_risk(row.risk_level),
                "; ".join(row.main_topics[:3]) or "—",
                "; ".join(row.main_skills[:2]) or "—",
            )
            for c_idx, value in enumerate(values):
                table.rows[r_idx].cells[c_idx].text = str(value)
        _style_table(table)
    # Группы участников определяем по текущей бизнес-логике без перерасчётов
    risk_group = [g for g in analysis.grades if (g.avg_completion_percent or 0) < 50]
    stable_group = [g for g in analysis.grades if 50 <= (g.avg_completion_percent or 0) < 60]
    high_group = [g for g in analysis.grades if (g.avg_completion_percent or 0) >= 60]
    potential_group = [g for g in analysis.grades if 55 <= (g.avg_completion_percent or 0) < 60]
    doc.add_heading("Выделенные группы обучающихся", level=2)
    _add_bullets(
        doc,
        [
            f"Группа риска: {len(risk_group)} параллелей с пониженным уровнем результатов.",
            f"Группа стабильных результатов: {len(stable_group)} параллелей.",
            f"Группа высокого уровня: {len(high_group)} параллелей.",
            f"Обучающиеся с положительным потенциалом: {len(potential_group)} параллелей с потенциалом роста.",
        ],
    )
    doc.add_heading("Индивидуальные образовательные маршруты и сопровождение", level=2)
    _add_bullets(
        doc,
        [
            "Сформировать индивидуальные образовательные маршруты для обучающихся группы риска.",
            "Организовать мероприятия по сопровождению обучающихся с устойчивыми дефицитами.",
            "Предусмотреть мероприятия для одарённых детей и обучающихся высокого уровня.",
            "Провести адресную работу с родителями по сопровождению обучающихся группы риска.",
            "Организовать контроль посещаемости дополнительных занятий при признаках риска.",
        ],
    )
    _add_cycle(
        doc,
        interpretation=[
            "Индивидуальные образовательные результаты в школе неоднородны и требуют дифференцированного сопровождения.",
        ],
        causes=[
            "Неравномерность сформированности предметных и метапредметных результатов.",
            "Различия в учебной мотивации и учебной дисциплине обучающихся.",
        ],
        org_decisions=[
            "Утвердить списки группы риска, стабильной и высокой групп на уровне администрации.",
            "Включить мониторинг индивидуальных маршрутов во внутришкольный контроль.",
        ],
        method_decisions=[
            "Организовать дифференцированные консультации по дефицитным содержательным линиям.",
            "Сформировать банк заданий для уровневой работы.",
        ],
        expected_effect=[
            "Снижение доли группы риска и рост доли стабильных/высоких результатов.",
        ],
    )

    doc.add_heading("3. Анализ статистики отметок", level=1)
    if analysis.subjects:
        table = doc.add_table(rows=1 + len(analysis.subjects), cols=7)
        headers = ("Ранг", "Предмет", "% выполнения", "Качество", "Успеваемость", "Дефициты", "Риск")
        for i, header in enumerate(headers):
            table.rows[0].cells[i].text = header
        for r_idx, row in enumerate(analysis.subjects, start=1):
            values = (
                row.rank,
                row.subject,
                _fmt(row.avg_completion_percent, "%"),
                _fmt(row.quality_percent, "%"),
                _fmt(row.absolute_percent, "%"),
                row.deficits_count,
                label_risk(row.risk_level),
            )
            for c_idx, value in enumerate(values):
                table.rows[r_idx].cells[c_idx].text = str(value)
        _style_table(table)
    else:
        doc.add_paragraph("Нет данных по статистике отметок.")
    _add_cycle(
        doc,
        interpretation=[
            "Распределение показателей качества и успеваемости отражает неоднородность образовательных результатов по предметам.",
            "По части предметов наблюдаются признаки системных затруднений в достижении планируемых результатов.",
        ],
        causes=[
            "Недостаточная согласованность текущего контроля и формата ВПР.",
            "Наличие устойчивых образовательных дефицитов по ряду предметов.",
        ],
        org_decisions=[
            "Рассмотреть итоги на административном совещании и методическом совете.",
            "Сформировать предложения по корректировке ВСОКО.",
        ],
        method_decisions=[
            "Скорректировать тематическое планирование и формы текущего контроля.",
            "Организовать методический разбор типичных ошибок на ШМО.",
        ],
        expected_effect=[
            "Рост качества знаний и снижение доли предметов с проблемными показателями.",
        ],
    )

    doc.add_heading("4. Сравнение отметок ВПР и журнала", level=1)
    # На уровне school_analysis нет детального split equal/lower/higher — сохраняем логику без новых расчётов.
    risk_label = label_school_risk(analysis.risk_profile.classification)
    table = doc.add_table(rows=5, cols=2)
    table.rows[0].cells[0].text = "Показатель"
    table.rows[0].cells[1].text = "Значение"
    table.rows[1].cells[0].text = "Степень объективности оценивания (интегрально)"
    table.rows[1].cells[1].text = risk_label
    table.rows[2].cells[0].text = "Качество знаний"
    table.rows[2].cells[1].text = _fmt(ov.avg_quality_percent, "%")
    table.rows[3].cells[0].text = "Успеваемость"
    table.rows[3].cells[1].text = _fmt(ov.avg_absolute_percent, "%")
    table.rows[4].cells[0].text = "Доля группы риска"
    table.rows[4].cells[1].text = _fmt(analysis.risk_profile.risk_group_percent, "%")
    _style_table(table)
    _add_cycle(
        doc,
        interpretation=[
            "Сопоставление внешней и внутренней оценки указывает на необходимость поддержания объективности оценивания в рамках ВСОКО.",
        ],
        causes=[
            "Риски завышения/занижения формируются при недостаточной критериальной согласованности оценивания.",
        ],
        org_decisions=[
            "Скорректировать локальные акты по оцениванию.",
            "Провести внутреннюю экспертизу и перекрёстную проверку.",
            "Включить мероприятия по объективности в план ВСОКО.",
        ],
        method_decisions=[
            "Организовать методические семинары по критериальному оцениванию.",
            "Провести обучение экспертов и взаимопроверку на ШМО.",
        ],
        expected_effect=[
            "Повышение объективности оценивания и снижение рисков расхождений.",
        ],
    )

    doc.add_heading("5. Анализ распределения первичных баллов", level=1)
    if analysis.subjects:
        table = doc.add_table(rows=1 + len(analysis.subjects), cols=3)
        headers = ("Предмет", "Средний % выполнения", "Дефициты")
        for i, header in enumerate(headers):
            table.rows[0].cells[i].text = header
        for r_idx, row in enumerate(analysis.subjects, start=1):
            values = (
                row.subject,
                _fmt(row.avg_completion_percent, "%"),
                row.deficits_count,
            )
            for c_idx, value in enumerate(values):
                table.rows[r_idx].cells[c_idx].text = str(value)
        _style_table(table)
    _add_cycle(
        doc,
        interpretation=[
            "Распределение результатов по предметам демонстрирует зоны концентрации низких и средних баллов.",
            "Выявлены предметы с повышенным риском аномальной концентрации дефицитов.",
        ],
        causes=[
            "Неравномерность освоения содержательных линий и заданий повышенной сложности.",
        ],
        org_decisions=[
            "Усилить контроль объективности процедуры проведения и проверки ВПР.",
            "Проводить внутреннюю выборочную экспертизу работ по риск-предметам.",
        ],
        method_decisions=[
            "Сфокусировать методическую работу на темах с низкими результатами.",
        ],
        expected_effect=[
            "Снижение аномалий и повышение достоверности оценочных процедур.",
        ],
    )

    doc.add_heading("6. Анализ выполнения заданий", level=1)
    doc.add_paragraph("Этапы анализа: задания → темы → содержательные линии → образовательные дефициты.")
    if protocol_task_tables:
        doc.add_paragraph(
            "Ниже приведено количество обучающихся, выполнивших задания правильно "
            "(полный балл), неправильно (0 баллов) и частично — по каждому протоколу."
        )
        for block in protocol_task_tables:
            title_block = str(block.get("title") or "Протокол")
            rows = list(block.get("rows") or [])
            if not rows:
                continue
            doc.add_heading(title_block, level=2)
            table = doc.add_table(rows=1 + len(rows), cols=6)
            headers = (
                "Задание",
                "Тема / умение",
                "% выполнения",
                "Верно (+)",
                "Ошибок (-)",
                "Всего",
            )
            for i, header in enumerate(headers):
                table.rows[0].cells[i].text = header
            for r_idx, row in enumerate(rows, start=1):
                topic = str(row.get("topic") or "—")
                skill = str(row.get("skill") or "").strip()
                if skill:
                    topic = f"{topic} / {skill}"
                pct = row.get("completion_percent")
                correct = int(row.get("correct_count") or 0)
                total = int(row.get("answers_count") or 0)
                incorrect = int(row.get("incorrect_count") if row.get("incorrect_count") is not None else max(0, total - correct))
                values = (
                    f"№{row.get('task_code') or '—'}",
                    topic,
                    _fmt(pct, "%") if pct is not None else "—",
                    correct,
                    incorrect,
                    total,
                )
                for c_idx, value in enumerate(values):
                    table.rows[r_idx].cells[c_idx].text = str(value)
            _style_table(table)
    doc.add_paragraph("Проблемные темы:")
    _add_bullets(doc, [f"{i.name} ({_fmt(i.value, '%')})" for i in analysis.weaknesses.topics])
    doc.add_paragraph("Проблемные умения:")
    _add_bullets(doc, [f"{i.name} ({_fmt(i.value, '%')})" for i in analysis.weaknesses.skills])
    _add_cycle(
        doc,
        interpretation=[
            "Ключевые дефициты фиксируются на уровне тем и умений, что требует перехода от перечня заданий к содержательному анализу.",
            "Количество правильных и неправильных ответов по заданиям протокола позволяет локализовать дефициты до конкретных элементов содержания.",
        ],
        causes=[
            "Недостаточная сформированность предметных и метапредметных действий по дефицитным линиям.",
        ],
        org_decisions=[
            "Утвердить перечень приоритетных содержательных линий для внутришкольного контроля.",
        ],
        method_decisions=[
            "Изменить методику преподавания по дефицитным содержательным линиям.",
            "Включить тематические коррекционные задания в рабочие материалы.",
        ],
        expected_effect=[
            "Повышение уровня освоения содержательных линий и сокращение числа проблемных умений.",
        ],
    )

    doc.add_heading("7. Анализ достижения планируемых результатов", level=1)
    weak_skills = [i.name for i in analysis.weaknesses.skills[:8]]
    strength_skills = [i.name for i in analysis.strengths.skills[:8]]
    table = doc.add_table(rows=1 + max(len(weak_skills), 1), cols=4)
    headers = ("Планируемый результат/умение", "Статус", "Предметные действия", "Метапредметные действия (УУД)")
    for i, header in enumerate(headers):
        table.rows[0].cells[i].text = header
    source = weak_skills if weak_skills else ["Нет достаточных данных для детальной декомпозиции результатов."]
    for r_idx, item in enumerate(source, start=1):
        table.rows[r_idx].cells[0].text = item
        table.rows[r_idx].cells[1].text = "не достигнут" if weak_skills else "частично достигнут"
        table.rows[r_idx].cells[2].text = "Требуется коррекция предметных действий по дефицитной линии."
        table.rows[r_idx].cells[3].text = "Требуется развитие регулятивных и познавательных УУД."
    _style_table(table)
    _add_cycle(
        doc,
        interpretation=[
            "Часть планируемых результатов достигнута частично или не достигнута, что ограничивает качество образования.",
            "Сильные результаты могут использоваться как ресурсные точки опоры.",
        ],
        causes=[
            "Неполная сформированность предметных действий по ряду содержательных линий.",
            "Недостаточное развитие отдельных метапредметных результатов.",
        ],
        org_decisions=[
            "Скорректировать рабочие программы по приоритетным дефицитам.",
            "Внести изменения в календарно-тематическое планирование.",
            "Обновить систему контроля в логике требований ВПР и ФГОС.",
        ],
        method_decisions=[
            "Усилить практико-ориентированные задания и задания на перенос способов действий.",
            "Обеспечить адресную коррекцию недостигнутых результатов.",
        ],
        expected_effect=[
            "Рост доли достигнутых планируемых результатов и повышение качества предметной подготовки.",
        ],
    )

    doc.add_heading("8. Анализ выполнения заданий различными группами участников", level=1)
    _add_bullets(
        doc,
        [
            f"Сильные обучающиеся: опора на сформированные умения ({'; '.join(strength_skills[:3]) or '—'}).",
            f"Средние обучающиеся: зона роста по умениям ({'; '.join(weak_skills[:3]) or '—'}).",
            "Слабые обучающиеся и группа риска: требуются индивидуальные маршруты и дополнительное сопровождение.",
            "Барьерные задания определяются по дефицитным темам и умениям с минимальным процентом выполнения.",
            "Массовые затруднения наблюдаются в содержательных линиях, повторяющихся в слабых результатах.",
        ],
    )
    _add_cycle(
        doc,
        interpretation=[
            "Сравнение групп участников выявляет барьерные задания и дифференцирующие факторы качества подготовки.",
        ],
        causes=[
            "Различия в уровне сформированности базовых и повышенных предметных действий.",
        ],
        org_decisions=[
            "Закрепить групповые стратегии сопровождения в плане внутришкольного контроля.",
        ],
        method_decisions=[
            "Организовать разноуровневую работу: сильные/средние/слабые/группа риска.",
        ],
        expected_effect=[
            "Сокращение разрыва между группами и снижение доли массовых затруднений.",
        ],
    )

    doc.add_heading("9. Образовательные дефициты", level=1)
    prio_table = doc.add_table(rows=1 + len(analysis.deficits.by_priority), cols=2)
    prio_table.rows[0].cells[0].text = "Приоритет"
    prio_table.rows[0].cells[1].text = "Количество"
    for idx, item in enumerate(analysis.deficits.by_priority, start=1):
        prio_table.rows[idx].cells[0].text = label_priority(item.priority)
        prio_table.rows[idx].cells[1].text = str(item.count)
    _style_table(prio_table)

    if analysis.deficits.by_subject:
        doc.add_paragraph("По предметам:")
        subj_table = doc.add_table(rows=1 + len(analysis.deficits.by_subject), cols=6)
        headers = ("Предмет", "Критический", "Высокий", "Средний", "Низкий", "Всего")
        for i, header in enumerate(headers):
            subj_table.rows[0].cells[i].text = header
        for r_idx, row in enumerate(analysis.deficits.by_subject, start=1):
            values = (
                row["name"],
                row["Critical"],
                row["High"],
                row["Medium"],
                row["Low"],
                row["total"],
            )
            for c_idx, value in enumerate(values):
                subj_table.rows[r_idx].cells[c_idx].text = str(value)
        _style_table(subj_table)
    top_def_topics = [x.get("name") for x in analysis.deficits.by_topic[:5] if x.get("name")]
    _add_cycle(
        doc,
        interpretation=[
            "Система образовательных дефицитов определяет ключевые ограничения качества образования.",
            "Приоритетные дефициты требуют управленческих решений на уровне администрации и ШМО.",
        ],
        causes=[
            "Устойчивые пробелы в освоении дефицитных тем и умений.",
        ],
        org_decisions=[
            "Утвердить перечень приоритетных дефицитов и ответственных за их устранение.",
            "Включить мониторинг устранения дефицитов во внутришкольный контроль.",
        ],
        method_decisions=[
            "Скорректировать содержание уроков и текущий контроль по приоритетным дефицитам.",
        ],
        expected_effect=[
            "Снижение числа критических/высоких дефицитов и повышение качества знаний.",
        ],
    )

    doc.add_heading("10. Работа администрации", level=1)
    doc.add_paragraph("Мероприятия для директора:")
    _add_bullets(
        doc,
        [
            "Рассмотрение результатов ВПР на административном совещании.",
            "Корректировка программы развития образовательной организации.",
            "Корректировка ВСОКО по итогам анализа.",
            "Контроль реализации утверждённых мероприятий.",
        ],
    )
    doc.add_paragraph("Мероприятия для заместителя директора:")
    _add_bullets(
        doc,
        [
            "Организация внутришкольного контроля по дефицитным предметным линиям.",
            "Анализ КТП и контроль рабочих программ.",
            "Мониторинг устранения образовательных дефицитов.",
        ],
    )
    _add_cycle(
        doc,
        interpretation=["Административный контур управления является ключевым условием реализации решений по ВПР."],
        causes=["Без закрепления ответственности управленческие решения исполняются фрагментарно."],
        org_decisions=["Утвердить дорожную карту и систему мониторинга исполнения."],
        method_decisions=["Обеспечить методическое сопровождение исполнения решений через ШМО."],
        expected_effect=["Повышение исполнительской дисциплины и управляемости процесса улучшения результатов."],
    )

    doc.add_heading("11. Работа школьных методических объединений", level=1)
    _add_bullets(
        doc,
        [
            "Анализ причин выявленных результатов по предметам.",
            "Проведение открытых уроков по дефицитным темам.",
            "Организация взаимопосещения уроков.",
            "Разработка банка заданий для устранения дефицитов.",
            "Корректировка рабочих программ.",
            "Обсуждение результатов ВПР на заседаниях ШМО.",
        ],
    )
    _add_cycle(
        doc,
        interpretation=["Методические объединения обеспечивают трансляцию аналитики ВПР в практику урока."],
        causes=["Единые методические подходы по дефицитным темам сформированы не в полном объёме."],
        org_decisions=["Включить блок ВПР в обязательную повестку ШМО на учебный год."],
        method_decisions=["Провести серию методических мероприятий по проблемным содержательным линиям."],
        expected_effect=["Рост методической согласованности и улучшение результатов по проблемным темам."],
    )

    doc.add_heading("12. Работа с педагогами", level=1)
    _add_bullets(
        doc,
        [
            "Выявленные профессиональные дефициты педагогов связаны с формированием и оценкой планируемых результатов.",
            "Провести предметные семинары и мастер-классы по дефицитным темам.",
            "Организовать курсы и вебинары по объективности оценивания и методике ВПР.",
            "Обеспечить индивидуальные консультации и наставничество.",
        ],
    )
    _add_cycle(
        doc,
        interpretation=["Профессиональные дефициты педагогов напрямую влияют на качество образовательных результатов."],
        causes=["Недостаточная проработка практик критериального оценивания и предметной методики по дефицитам."],
        org_decisions=["Утвердить индивидуальные траектории профессионального развития педагогов."],
        method_decisions=["Сформировать программу семинаров, мастер-классов и наставничества."],
        expected_effect=["Снижение профессиональных дефицитов и повышение качества преподавания."],
    )

    doc.add_heading("13. Работа с родителями", level=1)
    _add_bullets(
        doc,
        [
            "Индивидуальные консультации с родителями обучающихся группы риска.",
            "Информирование родителей о результатах ВПР и индивидуальных маршрутах.",
            "Совместные мероприятия школы и семьи по сопровождению обучающихся.",
            "Рекомендации по домашнему сопровождению учебной деятельности.",
        ],
    )
    _add_cycle(
        doc,
        interpretation=["Вовлечённость родителей усиливает результативность индивидуального сопровождения."],
        causes=["Недостаточная информированность семей о специфике выявленных дефицитов."],
        org_decisions=["Закрепить ответственность классных руководителей за коммуникацию с семьями группы риска."],
        method_decisions=["Подготовить единые рекомендации для родителей по сопровождению обучающихся."],
        expected_effect=["Рост вовлечённости родителей и стабилизация образовательных результатов группы риска."],
    )

    doc.add_heading("14. Методические рекомендации", level=1)
    recommendations = list(analysis.recommendations.actions[:20])
    if top_def_topics:
        recommendations.insert(
            0,
            "Приоритетное повторение тем: " + "; ".join(top_def_topics) + ".",
        )
    recommendations.extend(
        [
            "Скорректировать формы работы и технологии обучения с акцентом на дефицитные линии.",
            "Скорректировать виды текущего контроля в соответствии с форматом ВПР.",
            "Включить задания ВПР в систему текущего и тематического контроля.",
        ]
    )
    _add_bullets(doc, recommendations, empty="Рекомендации не сформированы.")
    _add_cycle(
        doc,
        interpretation=["Методические рекомендации направлены на адресную коррекцию образовательных дефицитов."],
        causes=["Прежние формы работы не обеспечивают устойчивое достижение планируемых результатов."],
        org_decisions=["Утвердить перечень обязательных методических изменений по предметам."],
        method_decisions=["Организовать внедрение рекомендаций через ШМО и внутришкольный контроль."],
        expected_effect=["Повышение качества знаний и устойчивости достижения планируемых результатов."],
    )

    doc.add_heading("15. План мероприятий", level=1)
    plan_rows = [
        (
            "Рассмотрение результатов ВПР и утверждение решений",
            "Директор, заместитель директора",
            "до 2 недель",
            "Приняты управленческие решения",
            "Наличие протокола/приказа",
        ),
        (
            "Корректировка ВСОКО и локальных актов",
            "Директор, администрация",
            "до 1 месяца",
            "Обновлённая модель внутришкольной оценки",
            "Внедрённые изменения в локальных актах",
        ),
        (
            "Корректировка рабочих программ и КТП",
            "Заместитель директора, руководители ШМО",
            "до 1 месяца",
            "Синхронизация программ с выявленными дефицитами",
            "Доля скорректированных РП/КТП",
        ),
        (
            "Реализация ИОМ и адресного сопровождения группы риска",
            "Учителя-предметники, классные руководители",
            "в течение четверти",
            "Позитивная динамика результатов группы риска",
            "Снижение доли обучающихся группы риска",
        ),
        (
            "Мониторинг устранения дефицитов",
            "Администрация, ШМО",
            "ежемесячно",
            "Контроль исполнения мер и корректировка плана",
            "Динамика показателей по дефицитам",
        ),
    ]
    plan_table = doc.add_table(rows=1 + len(plan_rows), cols=5)
    headers = ("Мероприятие", "Исполнитель", "Срок", "Ожидаемый результат", "Показатель эффективности")
    for i, h in enumerate(headers):
        plan_table.rows[0].cells[i].text = h
    for r_idx, row in enumerate(plan_rows, start=1):
        for c_idx, value in enumerate(row):
            plan_table.rows[r_idx].cells[c_idx].text = str(value)
    _style_table(plan_table)

    doc.add_heading("16. Итоговое экспертное заключение", level=1)
    conclusion = [
        "Общий уровень качества подготовки по ВПР оценивается как рабочий с выраженной неоднородностью результатов между предметами и группами обучающихся.",
        "Системные проблемы связаны с устойчивыми образовательными дефицитами по ряду содержательных линий и необходимостью усиления объективности оценивания.",
        "Основные причины: неравномерность сформированности предметных и метапредметных результатов, а также недостаточная согласованность текущего контроля с требованиями ВПР.",
        "Приоритетные направления вмешательства: адресная работа с группой риска, коррекция рабочих программ и КТП, усиление внутришкольного контроля и методического сопровождения.",
        "Наиболее эффективные управленческие решения включают корректировку ВСОКО, контроль исполнения дорожной карты и системную работу ШМО.",
        "Образовательная деятельность школы должна быть скорректирована в части методики преподавания, системы оценивания и сопровождения обучающихся.",
    ]
    _add_bullets(doc, conclusion)

    doc.add_heading("Приложение. Динамика показателей", level=1)
    if analysis.dynamics.available and analysis.dynamics.points:
        dyn_table = doc.add_table(rows=1 + len(analysis.dynamics.points), cols=6)
        headers = ("Год", "% выполнения", "Качество", "Участников", "Протоколов", "Тренд")
        for i, header in enumerate(headers):
            dyn_table.rows[0].cells[i].text = header
        for r_idx, point in enumerate(analysis.dynamics.points, start=1):
            values = (
                point.academic_year,
                _fmt(point.avg_completion_percent),
                _fmt(point.quality_percent),
                point.participants,
                point.protocols_count,
                label_trend(point.trend),
            )
            for c_idx, value in enumerate(values):
                dyn_table.rows[r_idx].cells[c_idx].text = str(value)
        _style_table(dyn_table)
    else:
        doc.add_paragraph(
            analysis.dynamics.message or "Недостаточно данных для анализа динамики."
        )

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer
