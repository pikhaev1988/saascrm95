"""
VPR STAGE 8 — Final Report Quality Audit (read-only).

Does NOT change methodology, thresholds, Metric Contract, groups, profiles, or peaks.
Produces findings + QUALITY_SCORE matrix for all production protocols.
"""

from __future__ import annotations

import json
import math
import re
from collections import Counter, defaultdict
from dataclasses import asdict, dataclass, field
from io import BytesIO
from pathlib import Path
from typing import Any


FORBIDDEN_PHRASES = [
    "необъективность доказана",
    "объективность нарушена",
    "учитель имеет дефицит",
    "педагог имеет дефицит",
    "учитель не сформировал",
    "учитель не умеет",
    "педагог не сформировал",
    "умение не сформировано",
    "умение отсутствует",
    "полностью не освоено",
    "учитель необъективен",
]

# Soft-forbidden in objectivity / conclusion overclaim contexts
SOFT_FORBIDDEN = [
    "нарушение объективности установлено",
    "доказана необъективность",
]

FIOKO_CLAIM_PATTERNS = [
    r"по\s+фиоко",
    r"фиоко\s+рекоменд",
    r"в\s+соответствии\s+с\s+фиоко",
    r"методологи[яи]\s+фиоко",
    r"в\s+логике\s+фиоко",
    r"классификаци[яи]\s+фиоко",
]

SYSTEM_ANALYTICS_MARKERS = [
    "80/50",
    "preparationprofile",
    "positive_potential",
    "положительный потенциал",
    "группы риска, стабильных",
    "внутренн",
]

LIMITED_MGMT_PATTERNS = [
    r"основной\s+вывод",
    r"управленческ\w+\s+решени",
    r"рекомендуется\s+принять\s+управленческ",
]

CROSS_YEAR_FAKE = [
    r"по\s+сравнению\s+с\s+прошлым\s+годом",
    r"в\s+динамике\s+за\s+\d+\s+лет",
    r"относительно\s+прошлого\s+года",
    r"год\s+назад",
]

OVERCLAIM = [
    (r"возможн\w+\s+дефицит", r"устойчив\w+\s+дефицит"),
    (r"возможн\w+\s+маркер", r"объективность\s+нарушена"),
]


@dataclass
class Finding:
    protocol_id: int
    section: str
    current: str
    problem: str
    evidence: str
    severity: str  # Critical | High | Medium | Low
    proposed_fix: str = ""

    def to_dict(self) -> dict[str, Any]:
        return asdict(self)


@dataclass
class ProtocolAudit:
    protocol_id: int
    subject: str
    parallel: int | None
    year: int | None
    n: int = 0
    catalog_status: str = ""
    grade: str = "B"
    numeric: str = "OK"
    text: str = "OK"
    fioko: str = "OK"
    groups: str = "OK"
    objectivity: str = "OK"
    multi: str = "OK"
    deficits: str = "OK"
    kpi: str = "OK"
    html_docx: str = "OK"
    limited_sample_groups: int = 0
    findings: list[Finding] = field(default_factory=list)
    forbidden_hits: list[dict[str, str]] = field(default_factory=list)
    warning_count: int = 0
    error_count: int = 0

    def to_row(self) -> dict[str, Any]:
        return {
            "protocol_id": self.protocol_id,
            "subject": self.subject,
            "class": self.parallel,
            "Numeric": self.numeric,
            "Text": self.text,
            "FIOKO": self.fioko,
            "Groups": self.groups,
            "Objectivity": self.objectivity,
            "Multi": self.multi,
            "Deficits": self.deficits,
            "KPI": self.kpi,
            "HTML/DOCX": self.html_docx,
            "Grade": self.grade,
            "catalog": self.catalog_status,
            "limited_sample_groups": self.limited_sample_groups,
            "warnings": self.warning_count,
            "errors": self.error_count,
            "findings_count": len(self.findings),
        }


def _safe_float(v: Any) -> float | None:
    try:
        if v is None:
            return None
        x = float(v)
        if math.isnan(x) or math.isinf(x):
            return None
        return x
    except (TypeError, ValueError):
        return None


def _is_bad_number(v: Any) -> bool:
    try:
        if v is None:
            return False
        x = float(v)
        return math.isnan(x) or math.isinf(x)
    except (TypeError, ValueError):
        return True


def _collect_report_texts(report) -> list[tuple[str, str]]:
    """Return list of (section, text)."""
    out: list[tuple[str, str]] = []

    def add(section: str, value: Any) -> None:
        if value is None:
            return
        if isinstance(value, (list, tuple)):
            for item in value:
                add(section, item)
            return
        if isinstance(value, dict):
            for k, v in value.items():
                add(f"{section}.{k}", v)
            return
        text = str(value).strip()
        if text:
            out.append((section, text))

    add("methodology", getattr(report, "methodology_basis", None))
    add("passport_assessment", getattr(report, "passport_assessment", None))
    for attr in (
        "individual_cycle",
        "marks_cycle",
        "journal_cycle",
        "distribution_cycle",
        "tasks_cycle",
        "planned_cycle",
        "groups_cycle",
        "deficits_cycle",
        "causes_cycle",
        "admin_cycle",
        "method_assoc_cycle",
        "teachers_cycle",
        "parents_cycle",
        "method_rec_cycle",
    ):
        cycle = getattr(report, attr, None)
        if cycle is None:
            continue
        for part in ("interpretation", "causes", "org_decisions", "method_decisions", "expected_effect"):
            add(f"{attr}.{part}", getattr(cycle, part, None))

    for g in getattr(report, "individual_groups", None) or []:
        add("groups.characteristic", getattr(g, "characteristic", None))
        add("groups.actions", getattr(g, "actions", None))
    for g in getattr(report, "group_task_insights", None) or []:
        add("groups.task_insight", getattr(g, "explanation", None))
        add("groups.task_evidence", getattr(g, "evidence", None))
    for d in getattr(report, "deficit_items", None) or []:
        add("deficits", getattr(d, "impact_results", None))
        add("deficits", getattr(d, "impact_quality", None))
        add("deficits", getattr(d, "evidence", None))
        add("deficits.mgmt", getattr(d, "management_decisions", None))
    for t in getattr(report, "teacher_deficits", None) or []:
        add("teachers", t)
    add("final_conclusion", getattr(report, "final_conclusion", None))
    add("system_analytics_notes", getattr(report, "system_analytics_notes", None))
    for row in getattr(report, "action_plan", None) or []:
        add("plan.action", getattr(row, "action", None))
        add("plan.problem", getattr(row, "problem", None))
        add("plan.expected", getattr(row, "expected_result", None))
    for rec in getattr(report, "management_recommendations", None) or []:
        if isinstance(rec, dict):
            for k, v in rec.items():
                add(f"mgmt.{k}", v)
        else:
            add("mgmt", rec)
    for note in getattr(report, "distribution_notes", None) or []:
        add("distribution", note)
    for note in getattr(report, "objectivity_notes", None) or []:
        add("objectivity", note)
    return out


def _html_to_text(html: str) -> str:
    text = re.sub(r"(?is)<script.*?>.*?</script>", " ", html)
    text = re.sub(r"(?is)<style.*?>.*?</style>", " ", text)
    text = re.sub(r"(?s)<[^>]+>", " ", text)
    text = re.sub(r"\s+", " ", text)
    return text.strip()


def _docx_to_text(buf: BytesIO) -> str:
    from docx import Document

    doc = Document(buf)
    parts = [p.text for p in doc.paragraphs if p.text and p.text.strip()]
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                t = (cell.text or "").strip()
                if t:
                    parts.append(t)
    return "\n".join(parts)


def _extract_percents(text: str) -> list[float]:
    vals = []
    for m in re.finditer(r"(\d+(?:[.,]\d+)?)\s*%", text):
        try:
            vals.append(float(m.group(1).replace(",", ".")))
        except ValueError:
            continue
    return vals


def _audit_numeric(analysis, report, protocol_id: int) -> list[Finding]:
    findings: list[Finding] = []
    tasks = list(getattr(analysis, "task_rows", None) or getattr(analysis, "tasks", None) or [])
    # Prefer comprehensive task analysis items
    task_analysis = getattr(analysis, "task_analysis", None)
    if task_analysis is not None:
        tasks = list(getattr(task_analysis, "items", None) or tasks)

    for t in tasks:
        n = int(getattr(t, "participants_count", None) or getattr(t, "n", None) or 0)
        full = int(getattr(t, "full_score_count", 0) or 0)
        partial = int(getattr(t, "partial_score_count", 0) or 0)
        zero = int(getattr(t, "zero_score_count", 0) or 0)
        code = str(getattr(t, "task_code", None) or getattr(t, "code", None) or "?")
        if n > 0 and full + partial + zero != n:
            findings.append(
                Finding(
                    protocol_id,
                    "tasks.metric_contract",
                    f"task {code}: full+partial+zero={full + partial + zero}, N={n}",
                    "FULL+PARTIAL+ZERO ≠ N",
                    f"full={full} partial={partial} zero={zero}",
                    "Critical",
                    "Исправить агрегацию Metric Contract для задания",
                )
            )
        for attr in (
            "full_score_rate",
            "partial_score_rate",
            "zero_score_rate",
            "completion_percent",
            "average_percent",
        ):
            v = getattr(t, attr, None)
            if _is_bad_number(v):
                findings.append(
                    Finding(
                        protocol_id,
                        "tasks.numeric",
                        f"task {code}.{attr}={v}",
                        "NaN/inf/invalid number",
                        attr,
                        "Critical",
                        "Убрать невалидные числа из расчёта",
                    )
                )
            else:
                fv = _safe_float(v)
                if fv is not None and (fv < 0 or fv > 100.0001):
                    findings.append(
                        Finding(
                            protocol_id,
                            "tasks.numeric",
                            f"task {code}.{attr}={fv}",
                            "percent out of [0,100]",
                            attr,
                            "High",
                            "Ограничить проценты диапазоном 0–100",
                        )
                    )
        max_score = _safe_float(getattr(t, "max_score", None) or getattr(t, "max_primary", None))
        full_rate = _safe_float(getattr(t, "full_score_rate", None))
        completion = _safe_float(getattr(t, "completion_percent", None))
        if (
            max_score
            and max_score > 1
            and full_rate is not None
            and completion is not None
            and abs(full_rate - completion) < 1e-9
            and int(getattr(t, "partial_score_count", 0) or 0) > 0
        ):
            findings.append(
                Finding(
                    protocol_id,
                    "tasks.metric_contract",
                    f"task {code}: completion={completion} == full_score_rate={full_rate}",
                    "completion_percent равен full_score_rate при partial>0",
                    f"partial={getattr(t, 'partial_score_count', 0)}",
                    "High",
                    "Развести completion и full_score_rate для multi-score",
                )
            )

    summary = getattr(analysis, "summary", None)
    for attr in ("mean_percent", "median_percent", "stdev", "cv", "mean_primary", "median_primary"):
        if summary is None:
            break
        v = getattr(summary, attr, None)
        if _is_bad_number(v):
            findings.append(
                Finding(
                    protocol_id,
                    "summary.numeric",
                    f"{attr}={v}",
                    "NaN/inf in summary",
                    attr,
                    "Critical",
                    "Исправить summary-метрику",
                )
            )
    return findings


def _audit_forbidden(texts: list[tuple[str, str]], protocol_id: int) -> tuple[list[Finding], list[dict]]:
    findings: list[Finding] = []
    hits: list[dict] = []
    for section, text in texts:
        low = text.lower()
        for phrase in FORBIDDEN_PHRASES + SOFT_FORBIDDEN:
            if phrase in low:
                # Multi-score: "умение не сформировано" only High if section suggests deficit absolute claim
                sev = "High"
                if phrase in ("умение не сформировано", "умение отсутствует", "полностью не освоено"):
                    sev = "High"
                if phrase in SOFT_FORBIDDEN or "необъектив" in phrase:
                    sev = "Critical" if "доказа" in phrase or "нарушена" in phrase else "High"
                findings.append(
                    Finding(
                        protocol_id,
                        section,
                        text[:240],
                        f"Forbidden wording: «{phrase}»",
                        phrase,
                        sev,
                        "Заменить на нейтральную формулировку по методологии ФИОКО",
                    )
                )
                hits.append({"section": section, "phrase": phrase, "sentence": text[:300]})
    return findings, hits


def _audit_fioko_attribution(texts: list[tuple[str, str]], report, protocol_id: int) -> list[Finding]:
    findings: list[Finding] = []
    for section, text in texts:
        low = text.lower()
        claims_fioko = any(re.search(p, low) for p in FIOKO_CLAIM_PATTERNS)
        if not claims_fioko:
            continue
        # Mixing SYSTEM_ANALYTICS into FIOKO claim
        systemish = any(m in low for m in SYSTEM_ANALYTICS_MARKERS) or any(
            x in low for x in ("групп риска", "высокого уровня", "стабильных результатов", "cv ", " коэффициент вариации")
        )
        # Allowed: methodology basis / explicit separation notes
        if "внутренн" in low and "system_analytics" in low.replace(" ", "_"):
            continue
        if "дополнительно применена внутренняя" in low:
            continue
        if "методологическая основа анализа" in low and "подтверждённые элементы" in low:
            continue
        if systemish and ("в логике фиоко" in low or "классификаци" in low):
            findings.append(
                Finding(
                    protocol_id,
                    section,
                    text[:240],
                    "SYSTEM_ANALYTICS представлен как FIOKO requirement",
                    text[:180],
                    "High",
                    "Явно разделить FIOKO methodology и SYSTEM_ANALYTICS",
                )
            )
        # 80/50 with FIOKO claim
        if ("80" in text and "50" in text) and claims_fioko and "system" not in low:
            if "групп" in low:
                findings.append(
                    Finding(
                        protocol_id,
                        section,
                        text[:240],
                        "Пороги 80/50 упомянуты в контексте FIOKO",
                        "80/50",
                        "Medium",
                        "Пометить 80/50 как SYSTEM_ANALYTICS",
                    )
                )
    # Check group source_kind — only flag if framed as official FIOKO without negation
    for g in getattr(report, "individual_groups", None) or []:
        sk = str(getattr(g, "source_kind", "") or "")
        char = str(getattr(g, "characteristic", "") or "")
        low = char.lower()
        if sk != "SYSTEM_ANALYTICS":
            continue
        # Explicit safe wording already separates SYSTEM vs FIOKO
        if any(
            x in low
            for x in (
                "system_analytics",
                "не является",
                "не официальн",
                "не требование фиоко",
                "внутренняя аналитическая",
                "дополнительная характеристика",
            )
        ):
            continue
        if "официальн" in low and "фиоко" in low:
            findings.append(
                Finding(
                    protocol_id,
                    "groups.source",
                    char[:200],
                    "SYSTEM_ANALYTICS group framed as official FIOKO",
                    sk,
                    "High",
                    "Уточнить формулировку source_kind",
                )
            )
    return findings


def _audit_groups(report, protocol_id: int) -> tuple[list[Finding], int]:
    findings: list[Finding] = []
    limited = 0
    for g in getattr(report, "individual_groups", None) or []:
        n = int(getattr(g, "count", 0) or 0)
        status = str(getattr(g, "sample_status", "") or "")
        informative = bool(getattr(g, "informative", True))
        title = str(getattr(g, "title", "") or getattr(g, "key", ""))
        char = str(getattr(g, "characteristic", "") or "")
        if n < 10:
            limited += 1
            if status != "LIMITED_SAMPLE" or informative:
                findings.append(
                    Finding(
                        protocol_id,
                        "groups.sample",
                        f"{title} N={n} status={status} informative={informative}",
                        "N<10 без LIMITED_SAMPLE/informative=false",
                        f"N={n}",
                        "High",
                        "Проставить LIMITED_SAMPLE и informative=false",
                    )
                )
            # management conclusion on limited sample
            blob = (char + " " + " ".join(str(a) for a in (getattr(g, "actions", None) or []))).lower()
            if any(re.search(p, blob) for p in LIMITED_MGMT_PATTERNS):
                if "ограничен" not in blob and "недостаточн" not in blob and "диагностическ" not in blob:
                    findings.append(
                        Finding(
                            protocol_id,
                            "groups.mgmt",
                            f"{title} N={n}: {char[:180]}",
                            "N<10 + управленческий вывод без маркировки ограниченной выборки",
                            char[:200],
                            "High",
                            "Смягчить до диагностического вывода / LIMITED_SAMPLE wording",
                        )
                    )
        elif n >= 10 and status == "LIMITED_SAMPLE":
            findings.append(
                Finding(
                    protocol_id,
                    "groups.sample",
                    f"{title} N={n} status={status}",
                    "N>=10 помечен как LIMITED_SAMPLE",
                    f"N={n}",
                    "Medium",
                    "Проверить sample_status",
                )
            )

    for gi in getattr(report, "group_task_insights", None) or []:
        status = str(getattr(gi, "sample_status", "") or "")
        informative = bool(getattr(gi, "informative", True))
        expl = str(getattr(gi, "explanation", "") or "").lower()
        if status == "LIMITED_SAMPLE" or not informative:
            limited += 1
            if any(re.search(p, expl) for p in LIMITED_MGMT_PATTERNS):
                if "ограничен" not in expl and "недостаточн" not in expl:
                    findings.append(
                        Finding(
                            protocol_id,
                            "groups.task_insight",
                            expl[:200],
                            "LIMITED_SAMPLE group insight с управленческим тоном",
                            expl[:200],
                            "High",
                            "Диагностическая формулировка без FIOKO management conclusion",
                        )
                    )
    return findings, limited


def _audit_objectivity(report, analysis, texts: list[tuple[str, str]], protocol_id: int) -> list[Finding]:
    findings: list[Finding] = []
    layer = getattr(analysis, "fioko_2026", None)
    dist = getattr(layer, "distribution", None) if layer else None
    if dist is not None:
        marker = bool(getattr(dist, "possible_objectivity_marker", False))
        bstatus = str(getattr(dist, "boundary_peak_status", "") or "")
        gp = getattr(dist, "general_peak", None)
        gp_peak = bool(getattr(gp, "is_peak", False)) if gp is not None else False
        if marker and bstatus == "NOT_AVAILABLE":
            findings.append(
                Finding(
                    protocol_id,
                    "objectivity.boundary",
                    f"marker={marker} boundary_peak_status={bstatus}",
                    "Objectivity marker при отсутствии официальных границ",
                    bstatus,
                    "High",
                    "Сбрасывать marker при NOT_AVAILABLE",
                )
            )
        if gp_peak and marker and bstatus not in ("HAS_MARKER", "OK"):
            # If only general peak drove marker
            flags = list(getattr(dist, "boundary_peak_flags", None) or [])
            has_boundary = any(str(getattr(f, "status", "")) == "POSSIBLE_MARKER" for f in flags)
            if not has_boundary:
                findings.append(
                    Finding(
                        protocol_id,
                        "objectivity.general_peak",
                        f"general_peak={gp_peak} marker={marker}",
                        "GENERAL_PEAK интерпретирован как objectivity marker",
                        str(getattr(gp, "score", None)),
                        "High",
                        "GENERAL_PEAK только как статистическая особенность",
                    )
                )

    for section, text in texts:
        low = text.lower()
        if "объективность нарушена" in low or "необъективность доказана" in low:
            findings.append(
                Finding(
                    protocol_id,
                    section,
                    text[:240],
                    "Запрещённая формулировка об объективности",
                    text[:180],
                    "Critical",
                    "Заменить на «возможный маркер… требующий дополнительного анализа»",
                )
            )
        if "общий пик" in low and ("нарушен" in low or "необъектив" in low):
            findings.append(
                Finding(
                    protocol_id,
                    section,
                    text[:240],
                    "GENERAL_PEAK связан с нарушением объективности в тексте",
                    text[:180],
                    "High",
                    "Развести GENERAL_PEAK и BOUNDARY_PEAK в тексте",
                )
            )
    return findings


def _audit_multi_score(analysis, texts: list[tuple[str, str]], protocol_id: int) -> list[Finding]:
    findings: list[Finding] = []
    task_analysis = getattr(analysis, "task_analysis", None)
    tasks = list(getattr(task_analysis, "items", None) or []) if task_analysis else []
    suspicious_codes = []
    for t in tasks:
        full = int(getattr(t, "full_score_count", 0) or 0)
        partial = int(getattr(t, "partial_score_count", 0) or 0)
        max_score = _safe_float(getattr(t, "max_score", None)) or 0
        if max_score > 1 and full == 0 and partial > 0:
            code = str(getattr(t, "task_code", None) or getattr(t, "code", "?"))
            suspicious_codes.append(code)

    if not suspicious_codes:
        return findings

    bad = ("умение не сформировано", "умение отсутствует", "полностью не освоено", "не сформировано")
    for section, text in texts:
        low = text.lower()
        if any(b in low for b in bad):
            # If mentions a multi-score task or generic absolute claim
            findings.append(
                Finding(
                    protocol_id,
                    section,
                    text[:240],
                    "Абсолютная формулировка при multi-score full=0/partial>0",
                    f"tasks={suspicious_codes[:8]}",
                    "High",
                    "Учитывать partial credit; не писать «умение отсутствует»",
                )
            )
            break
    return findings


def _audit_deficits(report, protocol_id: int) -> list[Finding]:
    findings: list[Finding] = []
    for d in getattr(report, "deficit_items", None) or []:
        name = str(getattr(d, "name", "") or "")
        evidence_status = str(getattr(d, "evidence_status", "") or "ESTABLISHED")
        evidence = str(getattr(d, "evidence", "") or "").strip()
        linked_tasks = list(getattr(d, "linked_tasks", None) or [])
        linked_results = list(getattr(d, "linked_results", None) or [])
        impact = str(getattr(d, "impact_results", "") or "")
        avg = getattr(d, "average_percent", None)

        if evidence_status == "INSUFFICIENT_DATA":
            # Neutral wording required; categorical claim is High
            low = impact.lower()
            if "недостаточно данных" not in low and any(
                x in low for x in ("выявлен дефицит", "существенно снижает", "дефицит «")
            ):
                findings.append(
                    Finding(
                        protocol_id,
                        "deficits.evidence",
                        name or impact[:120],
                        "INSUFFICIENT_DATA с категоричной формулировкой дефицита",
                        impact[:200],
                        "High",
                        "Нейтральная формулировка / INSUFFICIENT_DATA",
                    )
                )
            continue

        if not evidence and not linked_tasks and not linked_results:
            if avg is None and impact:
                findings.append(
                    Finding(
                        protocol_id,
                        "deficits.evidence",
                        name or impact[:120],
                        "Дефицит без доказательного основания",
                        "no evidence/tasks/results",
                        "High",
                        "Добавить evidence или пометить INSUFFICIENT_DATA",
                    )
                )
            elif avg is None:
                findings.append(
                    Finding(
                        protocol_id,
                        "deficits.evidence",
                        name[:120],
                        "Дефицит без evidence и метрик",
                        name,
                        "Medium",
                        "INSUFFICIENT_DATA при отсутствии данных",
                    )
                )
    return findings


def _audit_teacher_wording(texts: list[tuple[str, str]], protocol_id: int) -> list[Finding]:
    findings: list[Finding] = []
    patterns = [
        r"педагог\s+имеет\s+дефицит",
        r"учитель\s+имеет\s+дефицит",
        r"учитель\s+не\s+сформировал",
        r"педагог\s+не\s+сформировал",
        r"учитель\s+не\s+умеет",
        r"педагог\s+не\s+умеет",
    ]
    for section, text in texts:
        low = text.lower()
        for p in patterns:
            if re.search(p, low):
                findings.append(
                    Finding(
                        protocol_id,
                        section,
                        text[:240],
                        "Автоматический вывод о дефиците педагога по результатам учащихся",
                        p,
                        "High",
                        "Заменить на «зона, требующая методического анализа»",
                    )
                )
    return findings


def _audit_management(report, protocol_id: int) -> list[Finding]:
    findings: list[Finding] = []
    for row in getattr(report, "action_plan", None) or []:
        action = str(getattr(row, "action", "") or "").strip()
        if not action:
            continue
        problem = str(getattr(row, "problem", "") or "").strip()
        executor = str(getattr(row, "executor", "") or "").strip()
        deadline = str(getattr(row, "deadline", "") or "").strip()
        expected = str(getattr(row, "expected_result", "") or "").strip()
        kpi = str(getattr(row, "kpi", "") or getattr(row, "efficiency_indicator", "") or "").strip()
        missing = []
        if not problem:
            missing.append("problem")
        if not executor:
            missing.append("responsible")
        if not deadline:
            missing.append("deadline")
        if not expected:
            missing.append("expected_result")
        # KPI optional in some cases — Medium only if all core missing
        if missing:
            sev = "Medium" if len(missing) <= 2 else "High"
            findings.append(
                Finding(
                    protocol_id,
                    "management.plan",
                    action[:160],
                    f"Неполная управленческая рекомендация: нет {', '.join(missing)}",
                    f"kpi={'yes' if kpi else 'no'}",
                    sev,
                    "Дополнить цепочку FACT→ACTION→RESPONSIBLE→PERIOD→KPI",
                )
            )
        # Generic disconnected
        generic = ("общешкольн", "повысить качество", "усилить контроль", "провести работу")
        low = action.lower()
        if any(g in low for g in generic) and not problem and not kpi:
            findings.append(
                Finding(
                    protocol_id,
                    "management.generic",
                    action[:160],
                    "Generic recommendation без привязки к данным",
                    action[:160],
                    "Medium",
                    "Привязать к конкретному дефициту/метрике",
                )
            )
    return findings


def _audit_catalog_partial(analysis, report, texts: list[tuple[str, str]], protocol_id: int) -> list[Finding]:
    findings: list[Finding] = []
    layer = getattr(analysis, "fioko_2026", None)
    status = str(
        getattr(report, "catalog_mapping_status", None)
        or getattr(layer, "catalog_mapping_status", None)
        or ""
    )
    if status != "PARTIAL":
        return findings

    # What is missing
    cov = getattr(layer, "difficulty_coverage", None) or {}
    if isinstance(cov, dict):
        mapped = int(cov.get("mapped_tasks") or 0)
        total = int(cov.get("total_tasks") or 0)
        if total and mapped < total:
            findings.append(
                Finding(
                    protocol_id,
                    "catalog.partial",
                    f"difficulty mapped {mapped}/{total}",
                    "Catalog PARTIAL: неполное difficulty mapping",
                    f"{mapped}/{total}",
                    "Medium",
                    "Выводы по difficulty помечать LIMITED/NOT_AVAILABLE",
                )
            )

    # Strong conclusions that ignore PARTIAL
    for section, text in texts:
        low = text.lower()
        if status == "PARTIAL" and any(
            x in low for x in ("полный охват каталога", "все планируемые результаты сопоставлены", "catalog complete")
        ):
            findings.append(
                Finding(
                    protocol_id,
                    section,
                    text[:200],
                    "Текст утверждает полноту каталога при PARTIAL",
                    status,
                    "High",
                    "Отразить PARTIAL / INSUFFICIENT_DATA",
                )
            )
    # Planned results claiming certainty without mapping
    for row in getattr(report, "planned_results", None) or []:
        st = str(getattr(row, "status", "") or "")
        expl = str(getattr(row, "explanation", "") or "").lower()
        if st in ("NOT_AVAILABLE", "INSUFFICIENT_DATA", "LIMITED"):
            continue
        if "нет данных" in expl or "не сопостав" in expl:
            continue
        # if tasks_count 0 but strong status
        if int(getattr(row, "tasks_count", 0) or 0) == 0 and st not in ("", "UNKNOWN"):
            findings.append(
                Finding(
                    protocol_id,
                    "catalog.planned",
                    f"{getattr(row, 'result', '')[:80]} status={st}",
                    "Вывод по planned result без mapped tasks при PARTIAL catalog",
                    st,
                    "High",
                    "LIMITED / NOT_AVAILABLE / INSUFFICIENT_DATA",
                )
            )
    return findings


def _audit_cross_year(analysis, texts: list[tuple[str, str]], protocol_id: int) -> list[Finding]:
    findings: list[Finding] = []
    layer = getattr(analysis, "fioko_2026", None)
    cy = getattr(layer, "cross_year", None) if layer else None
    available = False
    if cy is not None:
        status = str(getattr(cy, "status", None) or getattr(cy, "availability", None) or "")
        available = status not in ("", "NOT_AVAILABLE", "UNAVAILABLE", "None")
        # also check if prior exists
        if getattr(cy, "prior_year", None) or getattr(cy, "previous", None):
            available = True
    for section, text in texts:
        low = text.lower()
        if any(re.search(p, low) for p in CROSS_YEAR_FAKE) and not available:
            findings.append(
                Finding(
                    protocol_id,
                    section,
                    text[:240],
                    "Фиктивная межгодовая динамика при отсутствии прошлогодних данных",
                    text[:180],
                    "High",
                    "Использовать NOT_AVAILABLE без сравнения с прошлым годом",
                )
            )
    return findings


def _audit_cross_subject(texts: list[tuple[str, str]], protocol_id: int, subject: str) -> list[Finding]:
    findings: list[Finding] = []
    # crude: comparing named other subjects
    subjects = (
        "математик",
        "русск",
        "биолог",
        "английск",
        "физик",
        "хими",
        "истори",
        "географи",
        "обществознан",
        "литератур",
        "информатик",
    )
    subj_low = (subject or "").lower()
    for section, text in texts:
        low = text.lower()
        if "по сравнению с" in low or "в отличие от" in low or "сопоставим" in low:
            others = [s for s in subjects if s in low and s not in subj_low]
            if others and "несопоставим" not in low and "not_comparable" not in low:
                findings.append(
                    Finding(
                        protocol_id,
                        section,
                        text[:240],
                        "Возможное кросс-предметное сравнение без NOT_COMPARABLE",
                        ",".join(others[:5]),
                        "Medium",
                        "Пометить NOT_COMPARABLE или убрать сравнение",
                    )
                )
    return findings


def _audit_conclusion(report, texts: list[tuple[str, str]], protocol_id: int) -> list[Finding]:
    findings: list[Finding] = []
    conclusions = list(getattr(report, "final_conclusion", None) or [])
    soft_evidence = " ".join(
        t for s, t in texts if "deficits" in s or "interpretation" in s
    ).lower()
    for line in conclusions:
        low = line.lower()
        if "устойчив" in low and "дефицит" in low and "возможн" in soft_evidence and "устойчив" not in soft_evidence:
            findings.append(
                Finding(
                    protocol_id,
                    "conclusion",
                    line[:240],
                    "Заключение сильнее доказательств (устойчивый vs возможный дефицит)",
                    line[:180],
                    "High",
                    "Согласовать силу формулировки с evidence",
                )
            )
        if ("доказан" in low or "установлен" in low) and "объектив" in low:
            findings.append(
                Finding(
                    protocol_id,
                    "conclusion",
                    line[:240],
                    "Заключение утверждает доказанную необъективность",
                    line[:180],
                    "Critical",
                    "Нейтральный язык: возможный маркер",
                )
            )
    return findings


def _audit_html_docx(html_text: str, docx_text: str, report, protocol_id: int) -> list[Finding]:
    findings: list[Finding] = []
    if not html_text or len(html_text) < 100:
        findings.append(
            Finding(protocol_id, "html", "empty/short", "HTML пустой или слишком короткий", str(len(html_text)), "Critical", "Проверить рендер HTML")
        )
    if not docx_text or len(docx_text) < 100:
        findings.append(
            Finding(protocol_id, "docx", "empty/short", "DOCX пустой или слишком короткий", str(len(docx_text)), "Critical", "Проверить генерацию DOCX")
        )
        return findings

    # Key shared tokens
    n = None
    for item in getattr(report, "passport", None) or []:
        label = str(getattr(item, "label", "") or "").lower()
        if "участник" in label or "обучающ" in label or label == "n":
            n = str(getattr(item, "value", "") or "")
            break
    checks = []
    if n:
        checks.append(("N", n))
    mb = str(getattr(report, "methodology_basis", "") or "").strip()
    if mb:
        # compare first 40 chars presence
        checks.append(("methodology", mb[:40]))
    cat = str(getattr(report, "catalog_mapping_status", "") or "")
    if cat:
        checks.append(("catalog", cat))

    html_l = html_text.lower()
    docx_l = docx_text.lower()
    for name, token in checks:
        t = token.lower()
        in_html = t in html_l
        in_docx = t in docx_l
        if in_html != in_docx:
            findings.append(
                Finding(
                    protocol_id,
                    "html_docx",
                    f"{name}={token}",
                    "HTML/DOCX mismatch по ключевому маркеру",
                    f"html={in_html} docx={in_docx}",
                    "High",
                    "Синхронизировать HTML и DOCX",
                )
            )

    # Percent sets — flag only large divergence of unique rounded ints
    hp = {round(x) for x in _extract_percents(html_text) if 0 <= x <= 100}
    dp = {round(x) for x in _extract_percents(docx_text) if 0 <= x <= 100}
    if hp and dp:
        # Allow subset differences due to charts/UI; flag if almost disjoint
        inter = hp & dp
        if len(inter) < max(1, min(len(hp), len(dp)) // 5):
            findings.append(
                Finding(
                    protocol_id,
                    "html_docx.percents",
                    f"html_pcts={len(hp)} docx_pcts={len(dp)} inter={len(inter)}",
                    "Сильное расхождение наборов процентов HTML vs DOCX",
                    f"sample_html={sorted(list(hp))[:8]} sample_docx={sorted(list(dp))[:8]}",
                    "Medium",
                    "Сверить числовые блоки HTML/DOCX",
                )
            )

    # LIMITED_SAMPLE consistency
    lim_html = "limited_sample" in html_l or "ограниченная выборка" in html_l or "недостаточна для информативного" in html_l
    lim_docx = "limited_sample" in docx_l or "ограниченная выборка" in docx_l or "недостаточна для информативного" in docx_l
    has_limited_group = any(
        str(getattr(g, "sample_status", "")) == "LIMITED_SAMPLE"
        for g in (getattr(report, "individual_groups", None) or [])
    )
    if has_limited_group and lim_html != lim_docx:
        findings.append(
            Finding(
                protocol_id,
                "html_docx.limited_sample",
                f"html={lim_html} docx={lim_docx}",
                "LIMITED_SAMPLE маркировка не согласована HTML/DOCX",
                "sample_status=LIMITED_SAMPLE",
                "High",
                "Одинаковая маркировка в HTML и DOCX",
            )
        )
    return findings


def _audit_text_vs_data(report, analysis, texts: list[tuple[str, str]], protocol_id: int) -> list[Finding]:
    """Flag text percentages that don't appear (within rounding) in structured data."""
    findings: list[Finding] = []
    known: set[float] = set()

    def add_num(v: Any) -> None:
        fv = _safe_float(v)
        if fv is None:
            return
        known.add(round(fv, 1))
        known.add(float(round(fv)))
        known.add(round(fv, 0))

    for item in getattr(report, "passport", None) or []:
        add_num(getattr(item, "value", None))
    for g in getattr(report, "individual_groups", None) or []:
        add_num(getattr(g, "percent", None))
        add_num(getattr(g, "count", None))
    task_analysis = getattr(analysis, "task_analysis", None)
    for t in list(getattr(task_analysis, "items", None) or [])[:50]:
        add_num(getattr(t, "completion_percent", None))
        add_num(getattr(t, "full_score_rate", None))
        add_num(getattr(t, "partial_score_rate", None))
    for d in getattr(report, "deficit_items", None) or []:
        add_num(getattr(d, "average_percent", None))
    for row in getattr(report, "planned_results", None) or []:
        add_num(getattr(row, "average_percent", None))
    for row in getattr(report, "mark_rows", None) or []:
        add_num(getattr(row, "percent", None))
        add_num(getattr(row, "share", None))

    if len(known) < 3:
        return findings

    mismatches = 0
    samples = []
    for section, text in texts:
        if section.startswith("plan.") or section.startswith("mgmt"):
            continue
        for p in _extract_percents(text):
            if p > 100 or p < 0:
                continue
            # allow rounding ±1.0
            if any(abs(p - k) <= 1.05 for k in known):
                continue
            # ignore 100/0 boilerplate
            if p in (0.0, 100.0, 50.0, 80.0):
                continue
            mismatches += 1
            if len(samples) < 3:
                samples.append(f"{section}:{p}% :: {text[:100]}")

    # Only flag if several unexplained percents (avoid false positives)
    if mismatches >= 8:
        findings.append(
            Finding(
                protocol_id,
                "text_data",
                f"unmatched_percents≈{mismatches}",
                "Подозрение на text↔data inconsistency (проценты в тексте не найдены в данных)",
                "; ".join(samples),
                "Medium",
                "Сверить округления и источники чисел в тексте",
            )
        )
    elif mismatches >= 5:
        findings.append(
            Finding(
                protocol_id,
                "text_data",
                f"unmatched_percents≈{mismatches}",
                "Подозрение на text↔data inconsistency (проценты в тексте не найдены в данных)",
                "; ".join(samples),
                "Low",
                "Сверить округления и источники чисел в тексте",
            )
        )
    return findings


def _audit_design(html: str, docx_text: str, report, protocol_id: int) -> list[Finding]:
    findings: list[Finding] = []
    # Empty sections / orphan headings — HTML headings followed by empty
    for m in re.finditer(r"<h[1-3][^>]*>(.*?)</h[1-3]>\s*(<p[^>]*>\s*</p>\s*){0,2}(?=<h[1-3]|</section|$)", html, re.I | re.S):
        title = re.sub(r"<[^>]+>", "", m.group(1)).strip()
        chunk = m.group(0)
        body = re.sub(r"<h[1-3][^>]*>.*?</h[1-3]>", "", chunk, flags=re.I | re.S)
        body_text = re.sub(r"<[^>]+>", " ", body)
        if title and len(re.sub(r"\s+", "", body_text)) < 3:
            findings.append(
                Finding(
                    protocol_id,
                    "design.html",
                    title[:120],
                    "Заголовок без содержимого (orphan/empty section)",
                    title[:120],
                    "Low",
                    "Скрыть пустую секцию или добавить контент/заглушку",
                )
            )
            if len(findings) > 8:
                break

    # Very long paragraphs in conclusion
    for line in getattr(report, "final_conclusion", None) or []:
        if len(str(line)) > 1200:
            findings.append(
                Finding(
                    protocol_id,
                    "design.text",
                    str(line)[:120] + "…",
                    "Слишком длинный абзац заключения",
                    f"len={len(str(line))}",
                    "Low",
                    "Разбить на более короткие тезисы",
                )
            )
    return findings


def _grade_from_findings(findings: list[Finding], warning_count: int) -> str:
    sev = {f.severity for f in findings}
    if "Critical" in sev:
        return "D"
    if "High" in sev:
        return "C"
    if findings or warning_count > 0:
        return "B"
    return "A"


def _axis_status(findings: list[Finding], prefixes: tuple[str, ...]) -> str:
    if not prefixes or prefixes == ("",):
        subset = list(findings)
    else:
        subset = [
            f
            for f in findings
            if any(f.section.startswith(p) or (p and p in f.section) for p in prefixes)
        ]
    if not subset:
        return "OK"
    worst = max(subset, key=lambda f: {"Critical": 3, "High": 2, "Medium": 1, "Low": 0}.get(f.severity, 0))
    return worst.severity


def audit_protocol(protocol) -> ProtocolAudit:
    from apps.vpr.comprehensive_analysis.engine import VprComprehensiveAnalysisEngine
    from apps.vpr.overview_docx import generate_overview_report_docx
    from apps.vpr.subject_report import build_subject_report
    from apps.vpr.validation.report_validator import VprReportValidator
    from django.template.loader import render_to_string

    engine = VprComprehensiveAnalysisEngine()
    validator = VprReportValidator()
    analysis = engine.analyze(protocol)
    report = build_subject_report(analysis, protocol, validate=False)
    validation = validator.validate(analysis, report)

    row = ProtocolAudit(
        protocol_id=protocol.id,
        subject=protocol.subject,
        parallel=protocol.parallel,
        year=protocol.academic_year,
        n=int(protocol.participants_count or 0),
        catalog_status=str(
            getattr(report, "catalog_mapping_status", None)
            or getattr(getattr(analysis, "fioko_2026", None), "catalog_mapping_status", None)
            or ""
        ),
        warning_count=len(validation.warnings),
        error_count=len(validation.errors),
    )

    texts = _collect_report_texts(report)
    findings: list[Finding] = []

    findings += _audit_numeric(analysis, report, protocol.id)
    forb, hits = _audit_forbidden(texts, protocol.id)
    findings += forb
    row.forbidden_hits = hits
    findings += _audit_fioko_attribution(texts, report, protocol.id)
    gfind, limited = _audit_groups(report, protocol.id)
    findings += gfind
    row.limited_sample_groups = limited
    findings += _audit_objectivity(report, analysis, texts, protocol.id)
    findings += _audit_multi_score(analysis, texts, protocol.id)
    findings += _audit_deficits(report, protocol.id)
    findings += _audit_teacher_wording(texts, protocol.id)
    findings += _audit_management(report, protocol.id)
    findings += _audit_catalog_partial(analysis, report, texts, protocol.id)
    findings += _audit_cross_year(analysis, texts, protocol.id)
    findings += _audit_cross_subject(texts, protocol.id, protocol.subject or "")
    findings += _audit_conclusion(report, texts, protocol.id)
    findings += _audit_text_vs_data(report, analysis, texts, protocol.id)

    html = render_to_string(
        "vpr/protocol_overview.html",
        {
            "protocol": protocol,
            "analysis": analysis,
            "report": report,
            "report_blocked": False,
            "report_validation": validation.to_dict(),
        },
    )
    html_text = _html_to_text(html)
    buf = generate_overview_report_docx(analysis, protocol, report=report)
    docx_text = _docx_to_text(BytesIO(buf.getvalue()))
    findings += _audit_html_docx(html_text, docx_text, report, protocol.id)
    findings += _audit_design(html, docx_text, report, protocol.id)

    # Validator errors elevate
    if validation.errors:
        for err in validation.errors[:5]:
            findings.append(
                Finding(
                    protocol.id,
                    "validator",
                    str(err)[:240],
                    "Validator error",
                    str(err)[:200],
                    "Critical",
                    "Устранить ошибку validator",
                )
            )

    row.findings = findings
    row.grade = _grade_from_findings(findings, row.warning_count)
    row.numeric = _axis_status(findings, ("tasks", "summary", "text_data"))
    row.text = _axis_status(findings, ("text_data", "forbidden", "teachers", "conclusion"))
    # refine text axis: forbidden hits
    if any(f.problem.startswith("Forbidden") for f in findings):
        row.text = _axis_status([f for f in findings if f.problem.startswith("Forbidden") or "формулиров" in f.problem], ("",)) or row.text
    row.fioko = _axis_status(findings, ("groups.source", "FIOKO", "fioko", "mgmt", "catalog"))
    # better fioko axis
    fioko_f = [f for f in findings if "FIOKO" in f.problem or "SYSTEM_ANALYTICS" in f.problem or "фиоко" in f.problem.lower()]
    row.fioko = _axis_status(fioko_f, ("",)) if fioko_f else "OK"
    row.groups = _axis_status(findings, ("groups",))
    row.objectivity = _axis_status(findings, ("objectivity",))
    row.multi = _axis_status(findings, ("tasks.metric", "multi")) or (
        "High" if any("multi-score" in f.problem.lower() or "Абсолютная" in f.problem for f in findings) else "OK"
    )
    if any("multi-score" in f.problem.lower() or "Абсолютная формулировка" in f.problem for f in findings):
        row.multi = _axis_status([f for f in findings if "multi" in f.problem.lower() or "Абсолютная" in f.problem], ("",))
    row.deficits = _axis_status(findings, ("deficits",))
    row.kpi = _axis_status(findings, ("management",))
    row.html_docx = _axis_status(findings, ("html", "docx", "html_docx", "design"))
    return row


def run_stage8_quality_audit(*, limit: int | None = None, out_dir: str = "apps/vpr/audit") -> dict:
    from apps.vpr.models import VprProtocol

    qs = VprProtocol.objects.all().order_by("id")
    if limit:
        qs = qs[: int(limit)]

    rows: list[ProtocolAudit] = []
    all_findings: list[Finding] = []
    grade_counter: Counter = Counter()
    sev_counter: Counter = Counter()
    forbidden_count = 0
    catalog_partial = 0
    limited_sample_protocols = 0
    html_docx_mismatch = 0
    numeric_mismatch = 0
    fioko_attr_issues = 0

    for protocol in qs.iterator():
        try:
            audited = audit_protocol(protocol)
        except Exception as exc:  # noqa: BLE001
            audited = ProtocolAudit(
                protocol_id=protocol.id,
                subject=getattr(protocol, "subject", "") or "",
                parallel=getattr(protocol, "parallel", None),
                year=getattr(protocol, "academic_year", None),
                grade="D",
                numeric="Critical",
                findings=[
                    Finding(
                        protocol.id,
                        "runner",
                        str(exc)[:300],
                        "Audit runner exception",
                        str(exc)[:300],
                        "Critical",
                        "Исправить падение аудита/отчёта",
                    )
                ],
            )
        rows.append(audited)
        all_findings.extend(audited.findings)
        grade_counter[audited.grade] += 1
        for f in audited.findings:
            sev_counter[f.severity] += 1
        forbidden_count += len(audited.forbidden_hits)
        if audited.catalog_status == "PARTIAL":
            catalog_partial += 1
        if audited.limited_sample_groups:
            limited_sample_protocols += 1
        if audited.html_docx not in ("OK",):
            html_docx_mismatch += 1
        if audited.numeric not in ("OK",):
            numeric_mismatch += 1
        if audited.fioko not in ("OK",):
            fioko_attr_issues += 1

    # Final status
    critical = sev_counter.get("Critical", 0)
    high = sev_counter.get("High", 0)
    if critical:
        status = "QUALITY_AUDIT_BLOCKED"
    elif high:
        status = "QUALITY_AUDIT_REQUIRES_FIXES"
    elif sev_counter.get("Medium", 0) or sev_counter.get("Low", 0) or grade_counter.get("B", 0):
        status = "QUALITY_AUDIT_PASS_WITH_WARNINGS"
    else:
        status = "QUALITY_AUDIT_PASS"

    payload = {
        "TOTAL": len(rows),
        "A": grade_counter.get("A", 0),
        "B": grade_counter.get("B", 0),
        "C": grade_counter.get("C", 0),
        "D": grade_counter.get("D", 0),
        "Critical": critical,
        "High": high,
        "Medium": sev_counter.get("Medium", 0),
        "Low": sev_counter.get("Low", 0),
        "forbidden_wording_count": forbidden_count,
        "catalog_partial_count": catalog_partial,
        "limited_sample_count": limited_sample_protocols,
        "html_docx_mismatch_count": html_docx_mismatch,
        "numeric_mismatch_count": numeric_mismatch,
        "fioko_attribution_issues": fioko_attr_issues,
        "status": status,
        "rows": [r.to_row() for r in rows],
        "findings": [f.to_dict() for f in all_findings],
        "control_cases": {
            "6": next((r.to_row() for r in rows if r.protocol_id == 6), None),
            "11": next((r.to_row() for r in rows if r.protocol_id == 11), None),
            "67": next((r.to_row() for r in rows if r.protocol_id == 67), None),
            "139": next((r.to_row() for r in rows if r.protocol_id == 139), None),
            "140": next((r.to_row() for r in rows if r.protocol_id == 140), None),
        },
    }

    out = Path(out_dir)
    out.mkdir(parents=True, exist_ok=True)
    json_path = out / "VPR_STAGE8_QUALITY_AUDIT.json"
    json_path.write_text(json.dumps(payload, ensure_ascii=False, indent=2), encoding="utf-8")
    md_path = out / "VPR_STAGE8_QUALITY_AUDIT.md"
    md_path.write_text(_render_markdown(payload, all_findings, rows), encoding="utf-8")
    return payload


def _render_markdown(payload: dict, findings: list[Finding], rows: list[ProtocolAudit]) -> str:
    lines: list[str] = []
    lines.append("# VPR STAGE 8 — FINAL REPORT QUALITY AUDIT")
    lines.append("")
    lines.append("Дата: production Beget · audit-only · методология Stage 7.1 не изменялась.")
    lines.append("")
    lines.append("## Summary")
    lines.append("")
    lines.append(f"| Metric | Value |")
    lines.append(f"|--------|-------|")
    for k in (
        "TOTAL",
        "A",
        "B",
        "C",
        "D",
        "Critical",
        "High",
        "Medium",
        "Low",
        "forbidden_wording_count",
        "catalog_partial_count",
        "limited_sample_count",
        "html_docx_mismatch_count",
        "numeric_mismatch_count",
        "fioko_attribution_issues",
        "status",
    ):
        lines.append(f"| {k} | {payload.get(k)} |")
    lines.append("")
    lines.append(f"**Final status: `{payload.get('status')}`**")
    lines.append("")
    lines.append("## Scope")
    lines.append("")
    lines.append("- Все production VPR protocols (ожидается 138).")
    lines.append("- Проверены: numeric, text↔data, FIOKO attribution, groups/LIMITED_SAMPLE,")
    lines.append("  objectivity, multi-score wording, deficits evidence, teacher wording,")
    lines.append("  management/KPI, catalog PARTIAL, cross-year/subject, conclusion, HTML/DOCX, design.")
    lines.append("- Код методологии **не изменялся**.")
    lines.append("")
    lines.append("## Quality matrix")
    lines.append("")
    lines.append(
        "| Protocol | Subject | Class | Numeric | Text | FIOKO | Groups | Objectivity | Multi | Deficits | KPI | HTML/DOCX | Grade |"
    )
    lines.append("|---|---|---|---|---|---|---|---|---|---|---|---|---|")
    for r in rows:
        lines.append(
            f"| {r.protocol_id} | {r.subject} | {r.parallel} | {r.numeric} | {r.text} | {r.fioko} | "
            f"{r.groups} | {r.objectivity} | {r.multi} | {r.deficits} | {r.kpi} | {r.html_docx} | {r.grade} |"
        )
    lines.append("")
    lines.append("## Control cases")
    lines.append("")
    for pid in (6, 11, 67, 139, 140):
        row = next((x for x in rows if x.protocol_id == pid), None)
        if not row:
            lines.append(f"- #{pid}: NOT FOUND")
            continue
        lines.append(
            f"- #{pid} {row.subject} class={row.parallel}: Grade={row.grade}, "
            f"catalog={row.catalog_status}, findings={len(row.findings)}, warnings={row.warning_count}"
        )
    lines.append("")
    lines.append("## Catalog PARTIAL detail")
    lines.append("")
    partial_rows = [r for r in rows if r.catalog_status == "PARTIAL"]
    if not partial_rows:
        lines.append("Нет протоколов с catalog=PARTIAL.")
    else:
        for r in partial_rows:
            lines.append(f"### Protocol {r.protocol_id} — {r.subject}")
            lines.append("")
            cat_f = [f for f in r.findings if f.section.startswith("catalog")]
            if cat_f:
                for f in cat_f:
                    lines.append(f"- [{f.severity}] {f.problem}: {f.evidence}")
            else:
                lines.append("- PARTIAL зафиксирован; отдельных catalog-findings нет (mapping gap отражён статусом).")
            lines.append("")

    lines.append("## Findings (by severity)")
    lines.append("")
    by_sev: dict[str, list[Finding]] = defaultdict(list)
    for f in findings:
        by_sev[f.severity].append(f)
    for sev in ("Critical", "High", "Medium", "Low"):
        items = by_sev.get(sev) or []
        lines.append(f"### {sev} ({len(items)})")
        lines.append("")
        if not items:
            lines.append("_none_")
            lines.append("")
            continue
        # Cap listing for readability but keep all High/Critical
        shown = items if sev in ("Critical", "High") else items[:80]
        for f in shown:
            lines.append(
                f"- **P{f.protocol_id}** / `{f.section}`: {f.problem}  \n"
                f"  current: {f.current[:200]}  \n"
                f"  evidence: {f.evidence[:200]}  \n"
                f"  proposed_fix: {f.proposed_fix}"
            )
        if len(items) > len(shown):
            lines.append(f"- … и ещё {len(items) - len(shown)} (см. JSON)")
        lines.append("")

    lines.append("## Forbidden wording hits")
    lines.append("")
    hits = []
    for r in rows:
        for h in r.forbidden_hits:
            hits.append((r.protocol_id, h))
    lines.append(f"Total hits: {len(hits)}")
    lines.append("")
    for pid, h in hits[:100]:
        lines.append(f"- P{pid} / {h.get('section')}: «{h.get('phrase')}» — {h.get('sentence', '')[:160]}")
    if len(hits) > 100:
        lines.append(f"- … +{len(hits) - 100} more in JSON")
    lines.append("")
    lines.append("## Notes")
    lines.append("")
    lines.append("- STAGE 8 = audit only; исправления не вносились.")
    lines.append("- Grade C = требуется ручная проверка (есть High).")
    lines.append("- Grade D = методологическая/критическая ошибка (есть Critical).")
    lines.append("- Validator warnings сами по себе дают Grade B при отсутствии High/Critical findings.")
    lines.append("")
    lines.append("## STOP")
    lines.append("")
    lines.append("Аудит завершён. Production code не изменять на основании этого файла без отдельного STAGE.")
    lines.append("")
    return "\n".join(lines)


if __name__ == "__main__":
    import os

    import django

    os.environ.setdefault("DJANGO_SETTINGS_MODULE", "analiz_gia.settings")
    django.setup()
    result = run_stage8_quality_audit()
    print(
        "TOTAL={TOTAL} A={A} B={B} C={C} D={D} status={status}".format(**result)
    )
