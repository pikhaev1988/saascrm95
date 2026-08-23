"""Аналитическая справка ВПР (Word) — структура по методологии ФИОКО (16 разделов)."""

from __future__ import annotations

from io import BytesIO

from apps.vpr.overview_charts import (
    chart_groups_distribution,
    chart_marks_distribution,
    chart_objectivity,
    chart_primary_scores,
    chart_task_success,
)
from apps.vpr.school_analysis.docx_export import _style_table
from apps.vpr.subject_report import SubjectReport, build_subject_report


def _cell(row, index: int, text) -> None:
    row.cells[index].text = "" if text is None else str(text)


def _unique(items: list[str] | None, *, exclude: set[str] | None = None) -> list[str]:
    seen: set[str] = set(exclude or ())
    out: list[str] = []
    for raw in items or []:
        text = str(raw or "").strip()
        if not text:
            continue
        key = " ".join(text.split()).lower()
        if key in seen:
            continue
        seen.add(key)
        out.append(text)
    return out


def _add_paragraphs(
    doc,
    paragraphs: list[str],
    *,
    exclude: set[str] | None = None,
    empty: str | None = None,
) -> set[str]:
    used = set(exclude or ())
    texts = _unique(paragraphs, exclude=used)
    if not texts:
        if empty:
            doc.add_paragraph(empty)
        return used
    for text in texts:
        doc.add_paragraph(text)
        used.add(" ".join(text.split()).lower())
    return used


def _add_bullets(
    doc,
    items: list[str],
    *,
    empty: str = "Не выявлено.",
    exclude: set[str] | None = None,
) -> set[str]:
    used = set(exclude or ())
    cleaned = _unique(items, exclude=used)
    if not cleaned:
        doc.add_paragraph(empty)
        return used
    for item in cleaned:
        doc.add_paragraph(item, style="List Bullet")
        used.add(" ".join(item.split()).lower())
    return used


def _add_cycle(doc, cycle, *, exclude: set[str] | None = None) -> set[str]:
    if not cycle:
        return set(exclude or ())
    used = set(exclude or ())

    interpretation = _unique(cycle.interpretation, exclude=used)
    if interpretation:
        doc.add_heading("Экспертная интерпретация", level=2)
        used = _add_paragraphs(doc, interpretation, exclude=used)

    causes = _unique(cycle.causes, exclude=used)
    if causes:
        doc.add_heading("Выявленные причины", level=2)
        used = _add_bullets(doc, causes, exclude=used)

    org = _unique(cycle.org_decisions, exclude=used)
    if org:
        doc.add_heading("Организационно-управленческие решения", level=2)
        used = _add_bullets(doc, org, exclude=used)

    method = _unique(cycle.method_decisions, exclude=used)
    if method:
        doc.add_heading("Методические решения", level=2)
        used = _add_bullets(doc, method, exclude=used)

    effect = _unique(cycle.expected_effect, exclude=used)
    if effect:
        doc.add_heading("Ожидаемый результат реализации мероприятий", level=2)
        used = _add_bullets(doc, effect, exclude=used)

    return used


def _add_chart(doc, image_buf: BytesIO | None, *, width_inches: float = 5.8) -> None:
    if image_buf is None:
        return
    from docx.shared import Inches

    image_buf.seek(0)
    doc.add_picture(image_buf, width=Inches(width_inches))


def generate_overview_report_docx(analysis, protocol, report: SubjectReport | None = None) -> BytesIO:
    from docx import Document

    if report is None:
        report = build_subject_report(analysis, protocol)

    doc = Document()
    doc.add_heading("Анализ результатов ВПР на уровне общеобразовательной организации", level=0)
    doc.add_paragraph(
        f"{report.subject}, {report.parallel} класс, {report.academic_year} учебный год"
    )
    doc.add_paragraph(
        "Аналитическая справка сформирована в логике методических рекомендаций ФИОКО "
        "(16 разделов алгоритма анализа)."
    )
    if report.school_name:
        doc.add_paragraph(report.school_name)
    else:
        org = analysis.organization_name or getattr(protocol, "organization_name", None) or ""
        if org:
            doc.add_paragraph(org)

    # 1
    doc.add_heading("1. Паспорт анализа", level=1)
    doc.add_paragraph(f"Профиль подготовки: {report.quality_level}")
    table = doc.add_table(rows=1 + len(report.passport), cols=2)
    _cell(table.rows[0], 0, "Показатель")
    _cell(table.rows[0], 1, "Значение")
    for i, item in enumerate(report.passport, start=1):
        _cell(table.rows[i], 0, item.label)
        _cell(table.rows[i], 1, item.value)
    _style_table(table)
    if report.passport_assessment:
        doc.add_heading("Экспертная оценка общего состояния качества образования", level=2)
        _add_paragraphs(doc, report.passport_assessment)

    # 2
    doc.add_heading("2. Анализ индивидуальных результатов", level=1)
    shown: set[str] = set()
    for group in report.individual_groups:
        pct = f" ({group.percent}%)" if group.percent else ""
        line = f"{group.title}: {group.count} чел.{pct}. {group.characteristic}"
        doc.add_paragraph(line)
        shown.add(" ".join(line.split()).lower())
        if group.sample_names:
            names = ", ".join(group.sample_names)
            extra = f"Примеры участников: {names}."
            doc.add_paragraph(extra)
            shown.add(" ".join(extra.split()).lower())
    _add_chart(doc, chart_groups_distribution(report.individual_groups), width_inches=5.6)
    _add_cycle(doc, report.individual_cycle, exclude=shown)
    doc.add_heading("Индивидуальные образовательные маршруты", level=2)
    for block in report.iom_blocks:
        doc.add_heading(block.group, level=3)
        doc.add_paragraph(block.focus)
        _add_bullets(doc, block.actions)
    if report.gifted_actions:
        doc.add_heading("Мероприятия для одарённых детей", level=2)
        _add_bullets(doc, report.gifted_actions)
    if report.parent_support_actions:
        doc.add_heading("Работа с родителями (сопровождение)", level=2)
        _add_bullets(doc, report.parent_support_actions)
    if report.attendance_control:
        doc.add_heading("Контроль посещаемости", level=2)
        _add_bullets(doc, report.attendance_control)

    # 3
    doc.add_heading("3. Анализ статистики отметок", level=1)
    if report.marks_rows:
        table = doc.add_table(rows=1 + len(report.marks_rows), cols=3)
        _cell(table.rows[0], 0, "Отметка")
        _cell(table.rows[0], 1, "Количество")
        _cell(table.rows[0], 2, "Доля")
        for i, row in enumerate(report.marks_rows, start=1):
            _cell(table.rows[i], 0, row.mark)
            _cell(table.rows[i], 1, str(row.count))
            _cell(table.rows[i], 2, f"{row.percent}%")
        _style_table(table)
        _add_chart(doc, chart_marks_distribution(report.marks_rows), width_inches=4.8)
    _add_cycle(doc, report.marks_cycle)

    # 4
    doc.add_heading("4. Сравнение отметок ВПР и журнала", level=1)
    doc.add_paragraph(f"Риск необъективности: {report.objectivity_risk}")
    for row in report.objectivity_rows:
        doc.add_paragraph(f"{row['label']}: {row['value']}")
    _add_chart(doc, chart_objectivity(report.objectivity_rows), width_inches=4.6)
    _add_cycle(doc, report.objectivity_cycle)

    # 5
    doc.add_heading("5. Анализ распределения первичных баллов", level=1)
    if report.scores_rows:
        table = doc.add_table(rows=1 + len(report.scores_rows), cols=3)
        _cell(table.rows[0], 0, "Первичный балл")
        _cell(table.rows[0], 1, "Количество")
        _cell(table.rows[0], 2, "Доля")
        for i, row in enumerate(report.scores_rows, start=1):
            _cell(table.rows[i], 0, row.score)
            _cell(table.rows[i], 1, str(row.count))
            _cell(table.rows[i], 2, f"{row.percent}%")
        _style_table(table)
        _add_chart(doc, chart_primary_scores(report.scores_rows), width_inches=5.8)
    _add_cycle(doc, report.scores_cycle)

    # 6
    doc.add_heading("6. Анализ выполнения заданий", level=1)
    task_rows = list(report.task_performance_rows or [])
    if not task_rows:
        from apps.vpr.expert_analysis.fioko_report import TaskPerformanceRow

        for row in getattr(analysis, "task_rows", None) or []:
            code = str(row.get("task_code") or "").strip()
            if not code:
                continue
            correct = int(row.get("plus") or row.get("correct_count") or row.get("full_count") or 0)
            total = int(row.get("total") or row.get("answers_count") or 0)
            incorrect = int(
                row.get("minus")
                if row.get("minus") is not None
                else row.get("incorrect_count")
                if row.get("incorrect_count") is not None
                else max(0, total - correct)
            )
            pct = row.get("success_rate")
            if pct is None:
                pct = row.get("completion_percent")
            if pct is None and total:
                pct = round(100.0 * correct / total, 1)
            task_rows.append(
                TaskPerformanceRow(
                    task_code=code,
                    topic=(row.get("topic") or "").strip(),
                    skill=(row.get("checked_skill") or "").strip(),
                    completion_percent=round(float(pct), 1) if pct is not None else None,
                    correct_count=correct,
                    incorrect_count=incorrect,
                    answers_count=total,
                )
            )

    if task_rows:
        doc.add_paragraph(
            "Количество обучающихся, выполнивших каждое задание верно (полный балл) "
            "и с ошибкой:"
        )
        table = doc.add_table(rows=1 + len(task_rows), cols=6)
        _cell(table.rows[0], 0, "Задание")
        _cell(table.rows[0], 1, "Тема / умение")
        _cell(table.rows[0], 2, "Успешность")
        _cell(table.rows[0], 3, "Верно (+)")
        _cell(table.rows[0], 4, "Ошибок (-)")
        _cell(table.rows[0], 5, "Всего")
        for i, row in enumerate(task_rows, start=1):
            topic_skill = row.topic or "—"
            if row.skill:
                topic_skill = f"{topic_skill} / {row.skill}"
            _cell(table.rows[i], 0, f"№{row.task_code}")
            _cell(table.rows[i], 1, topic_skill)
            _cell(
                table.rows[i],
                2,
                f"{row.completion_percent}%" if row.completion_percent is not None else "—",
            )
            _cell(table.rows[i], 3, str(row.correct_count))
            _cell(table.rows[i], 4, str(row.incorrect_count))
            _cell(table.rows[i], 5, str(row.answers_count))
        _style_table(table)
        _add_chart(doc, chart_task_success(task_rows), width_inches=5.8)

    pipeline_used = _add_paragraphs(doc, report.content_pipeline)
    for line in report.content_lines:
        pct = f"{line.average_percent}%" if line.average_percent is not None else "—"
        doc.add_heading(f"{line.name} — {line.mastery_label} ({pct})", level=3)
        if line.typical_errors:
            doc.add_paragraph("Типичные ошибки / проявления:")
            _add_bullets(doc, line.typical_errors)
        if line.probable_causes:
            doc.add_paragraph("Вероятные причины:")
            _add_bullets(doc, line.probable_causes)
        if line.method_changes:
            doc.add_paragraph("Изменения методики:")
            _add_bullets(doc, line.method_changes)
    _add_cycle(doc, report.content_cycle, exclude=pipeline_used)

    # 7
    doc.add_heading("7. Анализ достижения планируемых результатов", level=1)
    planned_shown: set[str] = set()
    for row in report.planned_results:
        pct = f"{row.average_percent}%" if row.average_percent is not None else "—"
        doc.add_heading(f"{row.result} — {row.status_label} ({pct})", level=3)
        parts = [row.explanation]
        if row.subject_actions:
            parts.append(f"Предметные действия: {row.subject_actions}")
        if row.meta_actions:
            parts.append(f"УУД: {row.meta_actions}")
        if row.content_adjustments:
            parts.append(f"Содержание: {row.content_adjustments}")
        planned_shown = _add_paragraphs(doc, parts, exclude=planned_shown)
    _add_cycle(doc, report.planned_cycle, exclude=planned_shown)

    # 8
    doc.add_heading("8. Анализ выполнения заданий различными группами участников", level=1)
    group_shown: set[str] = set()
    for item in report.group_task_insights:
        doc.add_paragraph(item.title, style="List Bullet")
        group_shown = _add_paragraphs(doc, [item.explanation], exclude=group_shown)
        if item.evidence:
            evidence = "; ".join(item.evidence)
            doc.add_paragraph(evidence)
            group_shown.add(" ".join(evidence.split()).lower())
    _add_cycle(doc, report.group_task_cycle, exclude=group_shown)

    # 9
    doc.add_heading("9. Образовательные дефициты", level=1)
    deficit_shown: set[str] = set()
    for d in report.deficit_items:
        pct = f"{d.average_percent}%" if d.average_percent is not None else "—"
        doc.add_heading(f"{d.name} ({d.kind}, приоритет {d.priority}, {pct})", level=3)
        deficit_shown = _add_paragraphs(
            doc,
            [d.impact_results, d.impact_quality, d.impact_program],
            exclude=deficit_shown,
        )
        if d.management_decisions:
            doc.add_paragraph("Управленческие решения:")
            deficit_shown = _add_bullets(doc, d.management_decisions, exclude=deficit_shown)
    _add_cycle(doc, report.deficits_cycle, exclude=deficit_shown)

    # 10
    doc.add_heading("10. Работа администрации", level=1)
    admin_shown: set[str] = set()
    doc.add_heading("Директор", level=2)
    admin_shown = _add_bullets(doc, report.admin_director, exclude=admin_shown)
    doc.add_heading("Заместитель директора", level=2)
    admin_shown = _add_bullets(doc, report.admin_deputy, exclude=admin_shown)
    _add_cycle(doc, report.admin_cycle, exclude=admin_shown)

    # 11
    doc.add_heading("11. Работа школьных методических объединений", level=1)
    smo_shown = _add_bullets(doc, report.smo_actions)
    _add_cycle(doc, report.smo_cycle, exclude=smo_shown)

    # 12
    doc.add_heading("12. Работа с педагогами", level=1)
    teach_shown: set[str] = set()
    doc.add_heading("Профессиональные дефициты педагогов", level=2)
    teach_shown = _add_bullets(doc, report.teacher_deficits, exclude=teach_shown)
    doc.add_heading("Мероприятия", level=2)
    teach_shown = _add_bullets(doc, report.teacher_actions, exclude=teach_shown)
    _add_cycle(doc, report.teachers_cycle, exclude=teach_shown)

    # 13
    doc.add_heading("13. Работа с родителями", level=1)
    parent_shown = _add_bullets(doc, report.parent_actions)
    _add_cycle(doc, report.parents_cycle, exclude=parent_shown)

    # 14
    doc.add_heading("14. Методические рекомендации", level=1)
    method_shown = _add_bullets(doc, report.method_recommendations)
    _add_cycle(doc, report.method_cycle, exclude=method_shown)

    # 15
    doc.add_heading("15. План мероприятий", level=1)
    if report.action_plan:
        table = doc.add_table(rows=1 + len(report.action_plan), cols=5)
        headers = (
            "Мероприятие",
            "Исполнитель",
            "Срок",
            "Ожидаемый результат",
            "Показатель эффективности",
        )
        for i, h in enumerate(headers):
            _cell(table.rows[0], i, h)
        for i, row in enumerate(report.action_plan, start=1):
            _cell(table.rows[i], 0, row.action)
            _cell(table.rows[i], 1, row.executor)
            _cell(table.rows[i], 2, row.deadline)
            _cell(table.rows[i], 3, row.expected_result)
            _cell(table.rows[i], 4, row.efficiency_indicator)
        _style_table(table)

    # 16
    doc.add_heading("16. Итоговое экспертное заключение", level=1)
    _add_paragraphs(doc, report.final_conclusion)

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer
